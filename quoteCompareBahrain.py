# ============================================================
#  FIXED VERSION — 3 targeted bug fixes applied:
#
#  FIX 1 — _extract_columnar SYSTEM_PROMPT:
#    • Riders/add-ons/sub-sections (TRAVEL RIDER, Artificial body parts, etc.)
#      must NOT become separate plan columns — merge into the parent plan.
#    • "0.2. Plan Name" must be the TOP-LEVEL plan name from the document
#      header/title, NOT a sub-section, rider, or benefit row heading.
#
#  FIX 2 — Raw value cleanup:
#    • LLM instruction: extract ONLY the benefit VALUE, never prepend the
#      row label / heading. The "Label | Value" pattern must be stripped.
#    • Added post-processing helper strip_label_prefix() applied to all
#      raw field values before storing.
#
#  FIX 3 — Normalization / sequential pass inherits same guardrails.
# ============================================================


import argparse
import os, json, pickle, re
from copy import deepcopy
from difflib import SequenceMatcher
from datetime import datetime
from pathlib import Path
from typing import Any
import xml.etree.ElementTree as ET
from dotenv import load_dotenv

load_dotenv()


def clean_excel_string(s):
    """Remove illegal/control characters for Excel compatibility."""
    if not isinstance(s, str):
        return s
    # Remove control characters that openpyxl doesn't allow (ASCII 0-8, 11-12, 14-31, 127)
    # Keep tab (9), newline (10), carriage return (13)
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', s)


def display_benefit_label(label: str) -> str:
    text = str(label or "")
    if text in CATEGORY_LOOKUP:
        return CATEGORY_LOOKUP[text]["name"]
    return text.split(". ", 1)[-1] if ". " in text else text


def repair_json(content: str) -> str:
    """
    Attempt to repair malformed JSON from LLM output.
    Handles: unterminated strings, unescaped quotes, trailing commas, and missing brackets.
    """
    if not content:
        return content
    
    # Strip any leading/trailing whitespace
    content = content.strip()
    
    # Remove markdown code fences if present
    if content.startswith("```json"):
        content = content[7:]
    elif content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()
    
    # Fix common issues:
    # 1. Remove trailing commas before } or ]
    content = re.sub(r',\s*([}\]])', r'\1', content)
    
    # 2. Fix unterminated strings by escaping newlines inside string values
    # Use proper backslash counting to detect real vs escaped quotes
    result = []
    in_string = False
    i = 0
    
    def count_preceding_backslashes(s, pos):
        """Count consecutive backslashes before position pos"""
        count = 0
        pos -= 1
        while pos >= 0 and s[pos] == '\\':
            count += 1
            pos -= 1
        return count
    
    while i < len(content):
        char = content[i]
        
        if char == '"':
            # Check if this quote is escaped (odd number of preceding backslashes)
            num_backslashes = count_preceding_backslashes(content, i)
            if num_backslashes % 2 == 0:
                # Even backslashes (including 0) = real quote, toggles string state
                in_string = not in_string
            result.append(char)
        elif in_string and char == '\n':
            # Newline inside string - escape it
            result.append('\\n')
        elif in_string and char == '\r':
            # Carriage return inside string - escape it  
            result.append('\\r')
        elif in_string and char == '\t':
            # Tab inside string - escape it
            result.append('\\t')
        elif in_string and ord(char) < 32:
            # Other control characters inside string - remove them
            pass
        else:
            result.append(char)
        i += 1
    
    content = ''.join(result)
    
    # 3. If we ended inside a string, close it and try to complete the structure
    if in_string:
        content += '"'
    
    # 4. Try to find the last complete key-value pair and trim incomplete data
    # This handles truncated LLM output better
    lines = content.split('\n')
    valid_lines = []
    brace_stack = []
    
    for line in lines:
        test_content = '\n'.join(valid_lines + [line])
        # Track brace/bracket nesting
        temp_in_string = False
        temp_stack = []
        valid = True
        
        for j, c in enumerate(test_content):
            if c == '"':
                num_bs = count_preceding_backslashes(test_content, j)
                if num_bs % 2 == 0:
                    temp_in_string = not temp_in_string
            elif not temp_in_string:
                if c in '{[':
                    temp_stack.append(c)
                elif c == '}':
                    if temp_stack and temp_stack[-1] == '{':
                        temp_stack.pop()
                    else:
                        valid = False
                        break
                elif c == ']':
                    if temp_stack and temp_stack[-1] == '[':
                        temp_stack.pop()
                    else:
                        valid = False
                        break
        
        if valid:
            valid_lines.append(line)
            brace_stack = temp_stack
    
    content = '\n'.join(valid_lines)
    
    # 5. Remove any trailing incomplete key-value pairs
    # Look for trailing patterns like `"key":` without a value
    content = re.sub(r',\s*"[^"]*":\s*$', '', content)
    content = re.sub(r',\s*$', '', content)
    
    # 6. Close any remaining open structures
    # Close in reverse order (most recently opened first)
    for opener in reversed(brace_stack):
        if opener == '{':
            content += '\n}'
        elif opener == '[':
            content += '\n]'
    
    return content


def extract_partial_json(content: str) -> dict:
    """
    Try to extract valid JSON objects from partially corrupted content.
    Useful when LLM output is truncated mid-way.
    """
    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    elif content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()
    
    # Try to find complete top-level objects
    result = {}
    
    # Find all top-level keys and their objects
    # Pattern: "key": { ... }
    pattern = r'"([^"]+)":\s*\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
    
    for match in re.finditer(pattern, content):
        key = match.group(1)
        obj_content = match.group(2)
        try:
            # Try to parse the inner object
            inner = json.loads('{' + obj_content + '}')
            result[key] = inner
        except:
            # Try a more aggressive cleanup
            try:
                # Remove the last incomplete key-value pair
                cleaned = re.sub(r',\s*"[^"]*":\s*"?[^"]*$', '', obj_content)
                cleaned = re.sub(r',\s*$', '', cleaned)
                inner = json.loads('{' + cleaned + '}')
                result[key] = inner
            except:
                continue
    
    return result


def safe_json_loads(content: str, context: str = "") -> dict:
    """
    Safely parse JSON with automatic repair on failure.
    """
    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        print(f"  ⚠️ JSON parse error{' in ' + context if context else ''}: {e}. Attempting repair...")
        repaired = repair_json(content)
        try:
            result = json.loads(repaired)
            print("  ✓ JSON repaired successfully")
            return result
        except json.JSONDecodeError as e2:
            print(f"  ⚠️ Standard repair failed: {e2}. Trying partial extraction...")
            # Try to extract whatever valid JSON we can
            partial = extract_partial_json(content)
            if partial:
                print(f"  ✓ Partial extraction recovered {len(partial)} top-level objects")
                return partial
            print(f"  ❌ JSON repair failed completely: {e2}")
            print(f"  Content preview (first 500 chars): {content[:500]}")
            raise

# ── ENV ───────────────────────────────────────────────────────────────────────
def _clean_env_value(value: str | None) -> str:
    text = str(value or "").strip().strip('"')
    if not text:
        return ""
    if " #" in text:
        text = text.split(" #", 1)[0].rstrip()
    return text


AZURE_ENDPOINT = _clean_env_value(os.getenv("AZURE_DOC_ENDPOINT", ""))
AZURE_KEY = _clean_env_value(os.getenv("AZURE_DOC_KEY", ""))

openai_endpoint = (
    os.getenv("AZURE_OPENAI_ENDPOINT")
    or os.getenv("AZ_OPENAI_ENDPOINT")
    or ""
)
openai_endpoint = _clean_env_value(openai_endpoint)
OPENAI_API_KEY = (
    os.getenv("AZURE_OPENAI_API_KEY")
    or os.getenv("AZ_OPENAI_KEY")
    or ""
)
OPENAI_API_KEY = _clean_env_value(OPENAI_API_KEY)
openai_deployment = (
    os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT_NAME")
    or os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
    or os.getenv("AZ_OPENAI_DEPLOYMENT")
    or "gpt-5.2"
)
openai_deployment = _clean_env_value(openai_deployment)
openai_standard_endpoint = (
    os.getenv("AZURE_OPENAI_STANDARD_ENDPOINT")
    or os.getenv("AZURE_OPENAI_GPT41_ENDPOINT")
    or os.getenv("AZURE_OPENAI_41_ENDPOINT")
    or os.getenv("AZ_OPENAI_STANDARD_ENDPOINT")
    or openai_endpoint
)
openai_standard_endpoint = _clean_env_value(openai_standard_endpoint)
OPENAI_STANDARD_API_KEY = (
    os.getenv("AZURE_OPENAI_STANDARD_API_KEY")
    or os.getenv("AZURE_OPENAI_GPT41_API_KEY")
    or os.getenv("AZURE_OPENAI_GPT41_KEY")
    or os.getenv("AZURE_OPENAI_41_KEY")
    or os.getenv("AZ_OPENAI_STANDARD_KEY")
    or OPENAI_API_KEY
)
OPENAI_STANDARD_API_KEY = _clean_env_value(OPENAI_STANDARD_API_KEY)
openai_standard_deployment = (
    os.getenv("AZURE_OPENAI_STANDARD_DEPLOYMENT_NAME")
    or os.getenv("AZURE_OPENAI_GPT41_DEPLOYMENT_NAME")
    or os.getenv("AZURE_OPENAI_41_DEPLOYMENT_NAME")
    or os.getenv("AZ_OPENAI_STANDARD_DEPLOYMENT")
    or "gpt-4.1"
)
openai_standard_deployment = _clean_env_value(openai_standard_deployment)
openai_api_version = _clean_env_value(os.getenv("AZURE_OPENAI_API_VERSION", "2024-12-01-preview"))
openai_standard_api_version = (
    os.getenv("AZURE_OPENAI_STANDARD_API_VERSION")
    or os.getenv("AZURE_OPENAI_GPT41_API_VERSION")
    or os.getenv("AZURE_OPENAI_41_API_VERSION")
    or os.getenv("AZ_OPENAI_STANDARD_API_VERSION")
    or openai_api_version
)
openai_standard_api_version = _clean_env_value(openai_standard_api_version)
OPENAI_REASONING_EFFORT = _clean_env_value(os.getenv("AZURE_OPENAI_REASONING_EFFORT", "none"))
DISABLE_SSL_VERIFY = os.getenv("QUOTECOMPARE_DISABLE_SSL_VERIFY", "true").strip().lower() in {"1", "true", "yes", "on"}
HTTP_TIMEOUT_SECONDS = float(os.getenv("QUOTECOMPARE_HTTP_TIMEOUT_SECONDS", "120"))
OUTPUT_DIR = Path(os.getenv("QUOTECOMPARE_OUTPUT_DIR", Path(__file__).resolve().parent / "outputs"))
DEFAULT_JSON_INDENT = 2
INSURER_COMPARISON_CHUNK_SIZE = int(
    os.getenv(
        "QUOTECOMPARE_INSURER_COMPARISON_CHUNK_SIZE",
        os.getenv("QUOTECOMPARE_INSURER_SUMMARY_CHUNK_SIZE", "25"),
    )
)
INSURER_SUMMARY_CHUNK_SIZE = INSURER_COMPARISON_CHUNK_SIZE
PLAN_FIELD_SUMMARY_CHUNK_SIZE = int(os.getenv("QUOTECOMPARE_PLAN_FIELD_SUMMARY_CHUNK_SIZE", "25"))



from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential
from openai import AzureOpenAI, NotFoundError
import httpx
import pandas as pd
import openpyxl
from openpyxl import load_workbook
from openpyxl.styles import (PatternFill, Font, Alignment, Border, Side)
from docx import Document
from docx.shared import Emu, Inches, Mm, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

if not AZURE_ENDPOINT or not AZURE_KEY:
    raise ValueError("Set AZURE_DOC_ENDPOINT and AZURE_DOC_KEY in your .env")
if not openai_endpoint or not OPENAI_API_KEY or not openai_deployment:
    raise ValueError(
        "Set AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_KEY, and GPT-5.2 deployment name in your .env"
    )
if not openai_standard_endpoint or not OPENAI_STANDARD_API_KEY or not openai_standard_deployment:
    raise ValueError(
        "Set AZURE_OPENAI_STANDARD_ENDPOINT, AZURE_OPENAI_STANDARD_API_KEY, and GPT-4.1 deployment name in your .env"
    )

# Initialize Document Intelligence client
doc_client = DocumentIntelligenceClient(AZURE_ENDPOINT, AzureKeyCredential(AZURE_KEY))

http_client = httpx.Client(
    verify=not DISABLE_SSL_VERIFY,
    timeout=HTTP_TIMEOUT_SECONDS,
    trust_env=True,
)

llm_client = AzureOpenAI(
    azure_endpoint=openai_endpoint,
    api_key=OPENAI_API_KEY,
    api_version=openai_api_version,
    http_client=http_client,
)
llm_standard_client = AzureOpenAI(
    azure_endpoint=openai_standard_endpoint,
    api_key=OPENAI_STANDARD_API_KEY,
    api_version=openai_standard_api_version,
    http_client=http_client,
)
print(f"✓ Azure OpenAI premium client initialized for deployment '{openai_deployment}'")
print(f"✓ Azure OpenAI standard client initialized for deployment '{openai_standard_deployment}'")
COST_PER_1K_INPUT = 0.005
COST_PER_1K_OUTPUT = 0.015

PREMIUM_TASK_CLASSES = {
    "extraction",
    "verification",
    "normalization",
    "benefit_summary",
    "insurer_summary",
    "insurer_highlighting",
}

_standard_deployment_missing_warned = False


def _resolve_model_for_task(task_class: str | None) -> tuple[AzureOpenAI, str, str, str]:
    normalized = str(task_class or "general").strip().lower()
    if normalized in PREMIUM_TASK_CLASSES:
        return llm_client, openai_deployment, "premium", openai_api_version
    return llm_standard_client, openai_standard_deployment, "standard", openai_standard_api_version


def llm_chat(
    system_prompt: str,
    user_prompt: str,
    json_mode: bool = True,
    log_label: str = "llm_call",
    log_context: dict | None = None,
    task_class: str | None = None,
    chunk_size: str | int | None = None,
) -> tuple[str, dict]:
    """
    Unified Azure OpenAI call function with task-based model routing.
    Returns: (response_text, usage_dict)
    """
    target_client, model_name, model_bucket, api_version = _resolve_model_for_task(task_class)
    reasoning_effort = OPENAI_REASONING_EFFORT if model_bucket == "premium" else ""

    kwargs = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0,
        "max_completion_tokens": 16384,
    }
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    if reasoning_effort:
        kwargs["reasoning_effort"] = reasoning_effort

    request_log = {
        "logged_at_utc": datetime.utcnow().isoformat() + "Z",
        "source_pdf": _current_pdf,
        "label": log_label,
        "deployment": model_name,
        "model_bucket": model_bucket,
        "task_class": task_class,
        "chunk_size": chunk_size,
        "api_version": api_version,
        "json_mode": json_mode,
        "reasoning_effort": reasoning_effort,
        "request": {
            "system_prompt": system_prompt,
            "user_prompt": user_prompt,
            "kwargs": kwargs,
        },
    }
    if log_context:
        request_log["request"]["context"] = log_context

    def _build_usage(response_obj, selected_model: str, selected_bucket: str) -> dict:
        return {
            "prompt_tokens": response_obj.usage.prompt_tokens,
            "completion_tokens": response_obj.usage.completion_tokens,
            "_logged_at_utc": request_log["logged_at_utc"],
            "_label": log_label,
            "_task_class": task_class,
            "_deployment": selected_model,
            "_model_bucket": selected_bucket,
        }

    def _store_success(response_obj, selected_model: str, selected_bucket: str):
        content = response_obj.choices[0].message.content
        usage = _build_usage(response_obj, selected_model, selected_bucket)
        request_log["response"] = {
            "output": content,
            "content": content,
            "usage": usage,
            "model": getattr(response_obj, "model", selected_model),
            "finish_reason": response_obj.choices[0].finish_reason if response_obj.choices else None,
        }
        log_llm_event(log_label, request_log)
        return content, usage

    try:
        response = target_client.chat.completions.create(**kwargs)
        return _store_success(response, model_name, model_bucket)
    except Exception as exc:
        global _standard_deployment_missing_warned
        if (
            model_bucket == "standard"
            and isinstance(exc, NotFoundError)
            and "DeploymentNotFound" in str(exc)
        ):
            fallback_kwargs = dict(kwargs)
            fallback_kwargs["model"] = openai_deployment
            request_log["fallback"] = {
                "reason": str(exc),
                "from_bucket": model_bucket,
                "from_deployment": model_name,
                "from_endpoint": openai_standard_endpoint,
                "to_bucket": "premium",
                "to_deployment": openai_deployment,
                "to_endpoint": openai_endpoint,
            }
            if not _standard_deployment_missing_warned:
                print(
                    "WARNING: Azure deployment "
                    f"'{model_name}' was not found at '{openai_standard_endpoint}'. "
                    f"Falling back to premium deployment '{openai_deployment}'."
                )
                _standard_deployment_missing_warned = True
            try:
                response = llm_client.chat.completions.create(**fallback_kwargs)
                return _store_success(response, openai_deployment, "premium")
            except Exception as fallback_exc:
                request_log["error"] = {
                    "type": type(fallback_exc).__name__,
                    "message": str(fallback_exc),
                }
                log_llm_event(f"{log_label}_error", request_log)
                raise

        request_log["error"] = {
            "type": type(exc).__name__,
            "message": str(exc),
        }
        log_llm_event(f"{log_label}_error", request_log)
        raise

FORCE_REFRESH = False

class TokenTracker:
    def __init__(self):
        self._data = {}
        self._calls = []

    def record(self, pdf_name: str, usage: dict):
        """Record token usage from an LLM call."""
        if not usage:
            return
        prompt_tokens = int(usage.get("prompt_tokens", 0) or 0)
        completion_tokens = int(usage.get("completion_tokens", 0) or 0)
        entry = self._data.setdefault(pdf_name, {"input": 0, "output": 0, "calls": 0})
        entry["input"]  += prompt_tokens
        entry["output"] += completion_tokens
        entry["calls"]  += 1
        self._calls.append({
            "Sequence": len(self._calls) + 1,
            "PDF / Source": pdf_name,
            "Logged At (UTC)": usage.get("_logged_at_utc", ""),
            "Call Label": usage.get("_label", ""),
            "Task Class": usage.get("_task_class", ""),
            "Deployment": usage.get("_deployment", ""),
            "Model Bucket": usage.get("_model_bucket", ""),
            "Input Tokens": prompt_tokens,
            "Output Tokens": completion_tokens,
            "Total Tokens": prompt_tokens + completion_tokens,
            "Cost (USD)": round(self.cost(prompt_tokens, completion_tokens), 6),
        })

    def cost(self, input_tokens: int, output_tokens: int) -> float:
        return (input_tokens / 1000 * COST_PER_1K_INPUT) + \
               (output_tokens / 1000 * COST_PER_1K_OUTPUT)

    def summary(self) -> dict:
        rows = {}
        total_in, total_out, total_calls = 0, 0, 0
        for pdf, d in self._data.items():
            rows[pdf] = {
                "Input Tokens":  d["input"],
                "Output Tokens": d["output"],
                "Total Tokens":  d["input"] + d["output"],
                "API Calls":     d["calls"],
                "Cost (USD)":    round(self.cost(d["input"], d["output"]), 4)
            }
            total_in    += d["input"]
            total_out   += d["output"]
            total_calls += d["calls"]
        rows["── TOTAL ──"] = {
            "Input Tokens":  total_in,
            "Output Tokens": total_out,
            "Total Tokens":  total_in + total_out,
            "API Calls":     total_calls,
            "Cost (USD)":    round(self.cost(total_in, total_out), 4)
        }
        return rows

    def call_details(self) -> list[dict]:
        return list(self._calls)

tracker = TokenTracker()
_current_pdf = "unknown"
_detected_plan_names = []   # populated by _detect_underwriter, used by _extract_columnar
_llm_log_sequence = 0
_summary_dict_cache: dict[str, dict] = {}
_underwriter_summary_cache: dict[str, dict[str, dict[str, Any]]] = {}
_verification_pass_cache: dict[str, dict] = {}
_normalization_pass_cache: dict[str, dict] = {}
_field_decision_batch_cache: dict[str, dict] = {}


def ensure_output_dir(output_dir: str | Path | None = None) -> Path:
    target = Path(output_dir) if output_dir else OUTPUT_DIR
    target.mkdir(parents=True, exist_ok=True)
    return target


def write_json_file(path: str | Path, payload: Any):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=DEFAULT_JSON_INDENT, ensure_ascii=False)


def read_json_file(path: str | Path) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def slugify_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    return value.strip("._") or "quote_compare"


def next_llm_log_path(label: str = "llm_call") -> Path:
    global _llm_log_sequence
    _llm_log_sequence += 1
    stamp = datetime.utcnow().strftime("%Y%m%dT%H%M%S")
    pdf_slug = slugify_filename(_current_pdf or "unknown_pdf")
    label_slug = slugify_filename(label)
    return ensure_output_dir() / "llm_logs" / f"{stamp}_{_llm_log_sequence:04d}_{pdf_slug}_{label_slug}.json"


def log_llm_event(label: str, payload: dict):
    write_json_file(next_llm_log_path(label), payload)


def serialize_document_intelligence_result(result: Any) -> dict:
    if hasattr(result, "as_dict"):
        return result.as_dict()
    if hasattr(result, "to_dict"):
        return result.to_dict()
    return {
        "content": getattr(result, "content", ""),
        "page_count": len(getattr(result, "pages", []) or []),
        "table_count": len(getattr(result, "tables", []) or []),
        "paragraph_count": len(getattr(result, "paragraphs", []) or []),
    }


def log_document_intelligence_output(
    file_path: str,
    result: Any,
    structured_md: str | None = None,
    di_page_filter: dict[str, Any] | None = None,
):
    output_root = ensure_output_dir() / "document_intelligence"
    output_root.mkdir(parents=True, exist_ok=True)
    stem = slugify_filename(Path(file_path).stem)

    payload = {
        "source_pdf": file_path,
        "logged_at_utc": datetime.utcnow().isoformat() + "Z",
        "di_page_filter": di_page_filter or {},
        "result": serialize_document_intelligence_result(result),
    }
    write_json_file(output_root / f"{stem}_di.json", payload)

    if structured_md is not None:
        md_path = output_root / f"{stem}_structured.md"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(structured_md)


def default_plan_record(plan_name: str = "", underwriter: str = "") -> dict:
    record = {f"{num}. {name}": "" for num, name, _ in CATEGORIES}
    record["0.1. Underwriter"] = underwriter
    record["0.2. Plan Name"] = plan_name
    return record


def score_field_information(value: Any) -> tuple[int, int]:
    text = str(value or "").strip()
    if not text:
        return (0, 0)
    lowered = text.lower()
    if lowered in {"not mentioned", "n/a", "none", "unknown"}:
        return (1, len(text))
    if lowered == "not covered":
        return (2, len(text))
    score = 3
    if any(token in lowered for token in ["covered", "%", "limit", "copay", "co-pay", "coinsurance", "waiting", "bhd", "usd", "aed"]):
        score += 2
    return (score, len(text))


def _coverage_polarity(value: Any) -> str:
    text = str(value or "").strip().lower()
    if not text:
        return ""
    if "not covered" in text:
        return "not_covered"
    if "covered" in text:
        return "covered"
    return ""


MULTI_ROW_FIELDS = {
    "3. Area of Cover",
    "6. Pre-existing and Chronic Conditions",
}

STRICT_SINGLE_FIELDS = {
    "0.1. Underwriter",
    "0.2. Plan Name",
    "2. Annual Limit per person",
    "4. TPA",
    "5. Network",
    "8. Room Type",
    "9. Parent Accommodation for child under 18 years of age",
    "10. Accommodation of an accompanying person",
    "11. Home Nursing",
    "21. New Born Cover",
    "24. Ambulance",
    "25. Psychiatric",
    "26. Organ Transplant",
    "27. Kidney Dialysis Treatment",
    "28. Repatriation",
    "29. Vaccinations",
    "30. Preventive Services",
    "31. Adult Pnuemococcal Conjugate Vaccine",
    "32. Cancer",
    "33. Prosthetics",
    "34. Return Air Fare",
    "35. Work Related Injury",
    "36. Influenza Vaccine",
    "37. HCV Hepatitis C Virus Infection Screening",
    "38. Hepatitis B Virus Screening and treatment",
    "39. UAE Within Network",
    "40. Claims Outside Network Within Country",
    "41. Claims Outside Country",
}

FIELD_CONTEXT_RULES = {
    "3. Area of Cover": {
        "required_any": [
            "country of residence", "principal country", "area where the members are located",
            "emergency medical treatment abroad", "non emergency medical treatment abroad",
            "treatment abroad", "area of cover", "worldwide", "bahrain extended",
        ],
        "forbidden_any": [
            "maternity", "dental", "optical", "consultation", "diagnostic",
            "sports", "sport", "recreation", "recreational", "athletic", "adventure", "extreme", "hazardous",
        ],
    },
    "2. Annual Limit per person": {
        "required_any": [
            "annual", "overall", "aggregate", "sum insured", "policy limit", "annual limit", "per person",
            "indemnity limit", "indemnity", "annual benefit limit", "maximum benefit", "yearly maximum", "yearly max",
        ],
        "forbidden_any": [
            "maternity", "delivery", "pregnancy", "dental", "optical", "consultation", "diagnostic",
            "laboratory", "physician", "medication", "medicine", "newborn", "vaccination", "physiotherapy",
            "cancer", "dialysis", "transplant", "room", "network",
            "worldwide", "area of cover", "country of residence", "principal country",
            "treatment abroad", "emergency medical treatment abroad", "non emergency medical treatment abroad",
            "sports", "sport", "recreation", "recreational", "athletic", "adventure", "extreme", "hazardous",
            "airfare", "air fare", "economy class", "return economy", "return airfare", "pre approval", "pre-approval",
        ],
    },
    "4. TPA": {
        "required_any": ["tpa", "administrator", "claims administrator", "claims processor", "claims handler", "administration company"],
        "forbidden_any": [
            "network list", "provider list", "hospital list", "sports", "recreation",
            "maternity", "consultation", "physician", "doctor", "diagnostic", "laboratory",
            "medication", "pharmacy", "dental", "optical",
        ],
    },
    "5. Network": {
        "required_any": ["network", "provider panel", "hospital network", "provider network", "network name"],
        "forbidden_any": ["copay", "co pay", "co-pay", "coinsurance", "deductible", "waiting period"],
    },
    "6. Pre-existing and Chronic Conditions": {
        "required_any": [
            "pre existing", "pre-existing", "chronic", "ped", "existing illness", "prior condition",
            "pre existing condition", "pre-existing condition", "pre existing chronic", "pre-existing chronic",
        ],
        "forbidden_any": ["maternity", "dental", "optical", "ambulance", "room type", "network"],
    },
    "8. Room Type": {"required_any": ["room", "ward", "accommodation", "semi private", "private room", "shared room"]},
    "12. Out Patient Deductible on Consultation": {
        "required_any": ["out patient", "outpatient", "opd", "consultation", "deductible", "copay", "coinsurance"],
        "forbidden_any": ["in patient", "inpatient", "maternity", "delivery"],
    },
    "13. Physician consultation charges": {
        "required_any": ["physician", "specialist", "doctor", "consultation", "gp", "consultant"],
        "forbidden_any": ["diagnostic", "laboratory", "pharmacy", "maternity", "dental", "optical"],
    },
    "14. Diagnostic Tests & Laboratory Tests": {
        "required_any": ["diagnostic", "laboratory", "lab", "pathology", "radiology", "imaging", "mri", "ct", "x ray", "x-ray"],
        "forbidden_any": [
            "consultation", "physician", "medication", "pharmacy", "maternity", "dental", "optical",
            "hearing", "hearing aid", "hearing aids", "auditory", "cochlear", "vision aid", "vision aids",
            "glasses", "frames", "lens", "lenses", "spectacle", "spectacles", "contact lens", "contact lenses",
        ],
    },
    "15. Prescribed Medication": {
        "required_any": ["medication", "medicine", "drug", "pharmacy", "prescription"],
        "forbidden_any": ["consultation", "diagnostic", "maternity", "dental", "optical"],
    },
    "18. Maternity In Patient Services": {
        "required_any": ["maternity", "pregnancy", "delivery", "childbirth", "caesarean", "c section", "c-section", "normal delivery"],
        "forbidden_any": ["annual limit", "overall limit", "dental", "optical", "cancer"],
    },
    "19. Maternity Out Patient Services": {
        "required_any": ["maternity", "pregnancy", "antenatal", "prenatal", "postnatal", "ob", "gyn"],
        "forbidden_any": ["annual limit", "overall limit", "dental", "optical", "cancer"],
    },
    "20. Life Threatening Maternity Complications": {
        "required_any": ["maternity", "complication", "ectopic", "miscarriage", "pregnancy emergency", "life threatening"],
        "forbidden_any": [
            "annual limit", "overall limit", "dental", "optical", "cancer", "antenatal", "prenatal", "postnatal",
            "routine", "routine maternity", "normal delivery", "consultation", "physician", "out patient", "outpatient",
            "eligible females",
        ],
    },
    "21. New Born Cover": {
        "required_any": ["new born", "newborn", "baby", "infant", "neonatal", "nicu", "day one", "30 days"],
        "forbidden_any": ["maternity annual maximum", "normal delivery", "c section", "antenatal", "prenatal", "postnatal"],
    },
    "22. Dental Benefit": {
        "required_any": ["dental", "tooth", "teeth", "oral", "root canal", "filling", "cleaning"],
        "forbidden_any": ["annual limit", "maternity", "optical", "cancer"],
    },
    "23. Optical Benefit": {
        "required_any": ["optical", "vision", "eye", "frame", "frames", "lens", "lenses", "spectacles", "contact lenses"],
        "forbidden_any": ["annual limit", "maternity", "dental", "cancer"],
    },
    "32. Cancer": {
        "required_any": ["cancer", "oncology", "chemotherapy", "radiation", "tumor", "malignancy", "chemo"],
        "forbidden_any": ["annual limit", "maternity", "dental", "optical"],
    },
    "35. Work Related Injury": {
        "required_any": [
            "work related", "work-related", "work injury", "occupational", "workplace", "employment injury",
            "job related", "job-related", "on duty", "on-duty", "workers compensation", "labor law",
        ],
        "forbidden_any": [
            "maternity", "dental", "optical", "consultation", "diagnostic", "medication", "pharmacy",
            "sports", "recreation", "vaccination",
        ],
    },
    "40. Claims Outside Network Within Country": {
        "required_any": [
            "out of network", "out-of-network", "non network", "non-network", "reimbursement", "oon",
            "outside network within country",
        ],
        "forbidden_any": ["in network", "in-network", "direct billing", "cashless", "network provider"],
    },
}

_field_decision_cache: dict[str, dict] = {}


def _field_behavior(field_key: str) -> str:
    if field_key in STRICT_SINGLE_FIELDS:
        return "strict_single"
    return "contextual"


def _field_allows_multi_value(field_key: str) -> bool:
    return _field_behavior(field_key) != "strict_single"


def _context_terms_present(text: str, terms: list[str]) -> bool:
    return any(term in text for term in terms)


def _context_guard_allows_row(field_key: str, row_label: str, row_section: str = "") -> bool:
    row_context = _normalize_match_text(" ".join(part for part in [row_section, row_label] if part))
    rules = FIELD_CONTEXT_RULES.get(field_key)
    if not rules:
        return True

    required_any = rules.get("required_any", [])
    forbidden_any = rules.get("forbidden_any", [])

    if required_any and not _context_terms_present(row_context, required_any):
        return False
    if forbidden_any and _context_terms_present(row_context, forbidden_any):
        return False
    return True


def _row_has_grouping_context(row: dict[str, Any]) -> bool:
    return bool(str(row.get("parent_label", "")).strip() or row.get("is_grouped"))


def _is_blankish(value: Any) -> bool:
    return str(value or "").strip().lower() in {"", "not mentioned", "n/a", "none", "not found", "not applicable"}


def _looks_like_aggregate_limit_value(value: Any) -> bool:
    text = str(value or "").strip()
    if not text:
        return False

    normalized = _normalize_match_text(text)
    if any(term in normalized for term in [
        "worldwide", "area of cover", "country of residence", "principal country",
        "treatment abroad", "emergency medical treatment abroad", "non emergency medical treatment abroad",
        "sports", "sport", "recreation", "recreational", "athletic", "adventure", "extreme", "hazardous",
        "airfare", "air fare", "economy class", "return economy", "return airfare", "pre approval",
    ]):
        return False

    if any(term in normalized for term in [
        "annual", "overall", "aggregate", "sum insured", "policy limit", "annual limit", "per person",
        "per insured member per year", "per year", "aml", "annual policy limit", "annual upper limit",
        "indemnity limit", "indemnity", "annual benefit limit", "yearly maximum", "yearly max",
    ]):
        return True

    has_currency = bool(re.search(r"\b(?:aed|bhd|usd|sar|qar|omr|kwd|eur|gbp|dhs|dhs\.)\b", normalized))
    has_digits = bool(re.search(r"\d", text))
    if has_currency and has_digits:
        return True

    return bool(re.fullmatch(r"[\d,\.\s]+", text)) and has_digits


def _backfill_annual_limit_from_cached_markdown(file_path: str, plan_map: dict[str, dict[str, Any]]) -> dict[str, dict[str, Any]]:
    if not file_path or not plan_map:
        return plan_map

    structured_path = Path(os.path.splitext(file_path)[0] + "_structured.md")
    if not structured_path.exists():
        return plan_map

    try:
        structured_md = structured_path.read_text(encoding="utf-8")
    except OSError:
        return plan_map

    table_rows = _parse_markdown_table_rows(structured_md)
    if not table_rows:
        return plan_map

    updated_map = deepcopy(plan_map)
    plan_names = list(updated_map.keys())
    field_key = "2. Annual Limit per person"
    single_plan_mode = len(plan_names) == 1

    for plan_idx, plan_name in enumerate(plan_names):
        plan_record = updated_map.get(plan_name, {})
        if str(plan_record.get(field_key, "") or "").strip():
            continue

        candidate_entries = _collect_field_candidate_entries(
            table_rows,
            field_key,
            plan_idx,
            single_plan_mode=single_plan_mode,
        )

        if not candidate_entries:
            continue

        viable_candidates: list[tuple[dict[str, Any], str]] = []
        for entry in candidate_entries:
            candidate_value = str(entry.get("candidate_value", "") or "").strip()
            restored_value = clean_raw_fields({field_key: candidate_value}).get(field_key, "")
            if not restored_value and _looks_like_aggregate_limit_value(candidate_value):
                restored_value = candidate_value
            if restored_value:
                viable_candidates.append((entry, restored_value))

        if not viable_candidates:
            continue

        best_candidate, restored_value = max(
            viable_candidates,
            key=lambda item: (
                score_field_information(str(item[1] or "")),
                len(str(item[1] or "")),
                -(int(item[0].get("page_number") or 0)),
                -(int(item[0].get("row_idx") or 0)),
            ),
        )
        if restored_value:
            plan_record[field_key] = restored_value
            print(f"  ↺ {plan_name}: restored Annual Limit from cached markdown -> {restored_value}")

    return updated_map


def _restore_tpa_from_network_value(network_value: Any) -> str:
    text = str(network_value or "").strip()
    if not text:
        return ""
    lowered = text.lower()
    if "nas" in lowered:
        return "Nas"
    if "nextcare" in lowered:
        return "Nextcare"
    if "mednet" in lowered:
        return "MedNet"
    if "globemed" in lowered or "globe med" in lowered:
        return "GlobeMed"
    if "msh" in lowered:
        return "MSH"
    return ""


def _backfill_tpa_from_cached_markdown(file_path: str, plan_map: dict[str, dict[str, Any]]) -> dict[str, dict[str, Any]]:
    if not file_path or not plan_map:
        return plan_map

    structured_path = Path(os.path.splitext(file_path)[0] + "_structured.md")
    updated_map = deepcopy(plan_map)

    if structured_path.exists():
        try:
            structured_md = structured_path.read_text(encoding="utf-8")
        except OSError:
            structured_md = ""
        table_rows = _parse_markdown_table_rows(structured_md)
    else:
        table_rows = []

    field_key = "4. TPA"
    plan_names = list(updated_map.keys())
    single_plan_mode = len(plan_names) == 1

    for plan_idx, plan_name in enumerate(plan_names):
        plan_record = updated_map.get(plan_name, {})
        if str(plan_record.get(field_key, "") or "").strip():
            continue

        restored_value = ""
        for entry in _collect_field_candidate_entries(
            table_rows,
            field_key,
            plan_idx,
            single_plan_mode=single_plan_mode,
        ):
            if entry:
                restored_value = clean_raw_fields({field_key: entry.get("candidate_value", "")}).get(field_key, "")
                if restored_value:
                    break

        if not restored_value:
            restored_value = _restore_tpa_from_network_value(plan_record.get("5. Network", ""))

        if restored_value:
            plan_record[field_key] = restored_value
            print(f"  ↺ {plan_name}: restored TPA from cached markdown/network -> {restored_value}")

    return updated_map


def _area_of_cover_component_key(text: str) -> str:
    normalized = _normalize_match_text(text)
    if any(term in normalized for term in ["principal country", "country of residence", "area where the members are located"]):
        return "principal"
    if "non emergency medical treatment abroad" in normalized:
        return "non_emergency"
    if any(term in normalized for term in ["emergency medical treatment abroad", "treatment abroad"]):
        return "emergency"
    return ""


def _area_of_cover_completeness_score(value: Any) -> int:
    components = {
        _area_of_cover_component_key(part)
        for part in _split_field_parts(value)
        if _area_of_cover_component_key(part)
    }
    return len(components)


def _preexisting_conditions_completeness_score(value: Any) -> int:
    return len(
        {
            _normalize_match_text(part)
            for part in _split_field_parts(value)
            if _normalize_match_text(part)
        }
    )


def _compose_area_of_cover_value(candidate_entries: list[dict[str, Any]], fallback_value: Any = "") -> str:
    ordered_entries = sorted(
        candidate_entries,
        key=lambda entry: (
            int(entry.get("page_number") or 0),
            int(entry.get("row_idx") or 0),
            str(entry.get("label", "")),
        ),
    )
    component_order = ["principal", "emergency", "non_emergency"]
    grouped_parts: dict[str, str] = {}
    extra_parts: list[str] = []

    for entry in ordered_entries:
        formatted = _format_multi_row_candidate(
            "3. Area of Cover",
            entry.get("label", ""),
            entry.get("candidate_value", ""),
        )
        if not formatted:
            continue
        lowered_formatted = formatted.lower()
        if any(term in lowered_formatted for term in ["sports", "sport", "recreation", "recreational", "athletic", "adventure", "extreme", "hazardous"]):
            continue
        component = _area_of_cover_component_key(formatted)
        if component and component not in grouped_parts:
            grouped_parts[component] = formatted
            continue
        extra_parts.append(formatted)

    ordered_parts = [grouped_parts[key] for key in component_order if key in grouped_parts]
    for extra_part in extra_parts:
        if _normalize_match_text(extra_part) not in {_normalize_match_text(part) for part in ordered_parts}:
            ordered_parts.append(extra_part)

    composed = _merge_distinct_field_parts("", "\n".join(ordered_parts))
    if composed:
        return composed
    return str(fallback_value or "").strip()


def _compose_preexisting_conditions_value(candidate_entries: list[dict[str, Any]], fallback_value: Any = "") -> str:
    ordered_entries = sorted(
        candidate_entries,
        key=lambda entry: (
            int(entry.get("page_number") or 0),
            int(entry.get("row_idx") or 0),
            str(entry.get("label", "")),
        ),
    )
    combined_value = ""
    for entry in ordered_entries:
        combined_value = _merge_distinct_field_parts(
            combined_value,
            _format_multi_row_candidate(
                "6. Pre-existing and Chronic Conditions",
                entry.get("label", ""),
                entry.get("candidate_value", ""),
            ),
        )
    if combined_value:
        return combined_value
    return str(fallback_value or "").strip()


def _compose_multi_row_field_value(field_key: str, candidate_entries: list[dict[str, Any]], fallback_value: Any = "") -> str:
    if field_key == "3. Area of Cover":
        return _compose_area_of_cover_value(candidate_entries, fallback_value=fallback_value)
    if field_key == "6. Pre-existing and Chronic Conditions":
        return _compose_preexisting_conditions_value(candidate_entries, fallback_value=fallback_value)

    combined_value = ""
    ordered_entries = sorted(
        candidate_entries,
        key=lambda entry: (
            int(entry.get("page_number") or 0),
            int(entry.get("row_idx") or 0),
            str(entry.get("label", "")),
        ),
    )
    for entry in ordered_entries:
        combined_value = _merge_distinct_field_parts(
            combined_value,
            _format_multi_row_candidate(field_key, entry.get("label", ""), entry.get("candidate_value", "")),
        )
    return combined_value or str(fallback_value or "").strip()


KNOWN_TPA_VALUES = {
    "nas",
    "nextcare",
    "next care",
    "mednet",
    "globemed",
    "globe med",
    "msh",
    "msa",
    "neuron",
}


def _looks_like_known_tpa_value(value: Any) -> bool:
    normalized = _normalize_match_text(value)
    return normalized in KNOWN_TPA_VALUES


def _normalize_verification_recovered_value(
    field_key: str,
    value: Any,
    candidate_entries: list[dict[str, Any]] | None = None,
    source: str = "llm",
) -> str:
    text = str(value or "").strip()
    if not text:
        return ""

    text = str(clean_raw_fields({field_key: text}).get(field_key, text) or "").strip()
    if not text:
        return ""

    if field_key in MULTI_ROW_FIELDS and candidate_entries:
        text = _compose_multi_row_field_value(field_key, candidate_entries, fallback_value=text)

    normalized = _normalize_match_text(text)
    rules = FIELD_CONTEXT_RULES.get(field_key)
    if not rules:
        if field_key == "4. TPA" and not _looks_like_known_tpa_value(text):
            return "" if source == "llm" else text
        return text

    forbidden_any = rules.get("forbidden_any", [])
    if forbidden_any and _context_terms_present(normalized, forbidden_any):
        return ""

    required_any = rules.get("required_any", [])
    if required_any and not _context_terms_present(normalized, required_any):
        if field_key == "4. TPA" and _looks_like_known_tpa_value(text):
            return text
        if source == "deterministic" and candidate_entries:
            return text
        return ""

    return text


def _is_field_incomplete(field_key: str, value: Any) -> bool:
    if _is_blankish(value):
        return True
    if field_key == "3. Area of Cover":
        return _area_of_cover_completeness_score(value) < 3
    if field_key == "6. Pre-existing and Chronic Conditions":
        return _preexisting_conditions_completeness_score(value) < 2
    return False


def _field_value_preference_score(field_key: str, value: Any) -> tuple[int, int, int]:
    text = str(value or "").strip()
    if field_key == "3. Area of Cover":
        completeness = _area_of_cover_completeness_score(text)
    elif field_key == "6. Pre-existing and Chronic Conditions":
        completeness = _preexisting_conditions_completeness_score(text)
    else:
        completeness = 0
    return (completeness, score_field_information(text), len(text))


def _prefer_field_value(field_key: str, existing: Any, candidate: Any) -> str:
    existing_text = str(existing or "").strip()
    candidate_text = str(candidate or "").strip()
    if not existing_text:
        return candidate_text
    if not candidate_text:
        return existing_text
    if field_key in MULTI_ROW_FIELDS:
        return candidate_text if _field_value_preference_score(field_key, candidate_text) > _field_value_preference_score(field_key, existing_text) else existing_text
    return candidate_text if score_field_information(candidate_text) >= score_field_information(existing_text) else existing_text


def _deterministic_candidate_backfill_value(field_key: str, candidate_entries: list[dict[str, Any]], current_value: Any = "") -> str:
    if not candidate_entries:
        return str(current_value or "").strip()
    ordered_entries = sorted(
        candidate_entries,
        key=lambda entry: (
            int(entry.get("page_number") or 0),
            int(entry.get("row_idx") or 0),
            str(entry.get("label", "")),
        ),
    )
    if field_key in MULTI_ROW_FIELDS:
        return _prefer_field_value(field_key, current_value, _compose_multi_row_field_value(field_key, ordered_entries, fallback_value=current_value))

    combined_value = ""
    for entry in ordered_entries:
        combined_value = _merge_distinct_field_parts(
            combined_value,
            _format_multi_row_candidate(field_key, entry.get("label", ""), entry.get("candidate_value", "")),
        )
    return _prefer_field_value(field_key, current_value, combined_value)


def _split_field_parts(value: Any) -> list[str]:
    text = str(value or "").strip()
    if not text:
        return []
    return [part.strip() for part in re.split(r"\s*(?:;|\n)\s*", text) if part.strip()]


def _strip_aggregate_limit_label_prefix(value: Any) -> str:
    text = str(value or "").strip()
    if not text or ":" not in text:
        return text

    label, remainder = text.split(":", 1)
    label_norm = _normalize_match_text(label)
    remainder = remainder.strip()
    if not remainder:
        return text

    if any(term in label_norm for term in [
        "annual", "overall", "aggregate", "sum insured", "policy limit",
        "annual limit", "maximum limit", "indemnity limit", "indemnity amount", "indemnity",
    ]):
        return remainder
    return text


def _merge_distinct_field_parts(existing: Any, candidate: Any) -> str:
    merged_parts: list[str] = []
    seen_norms: set[str] = set()

    for source in (existing, candidate):
        for part in _split_field_parts(source):
            part_norm = _normalize_match_text(part)
            if not part_norm or part_norm in seen_norms:
                continue
            merged_parts.append(part)
            seen_norms.add(part_norm)

    return "\n".join(merged_parts)


def _format_multi_row_candidate(field_key: str, row_label: str, candidate_value: str) -> str:
    value_text = str(candidate_value or "").strip()
    if not value_text:
        return ""

    label_text = str(row_label or "").strip().strip(":")
    field_name = CATEGORY_LOOKUP.get(field_key, {}).get("name", "")
    if label_text and _normalize_match_text(label_text) != _normalize_match_text(field_name):
        return f"{label_text}: {value_text}"

    return value_text


def merge_field_value(existing: Any, candidate: Any, field_key: str = "") -> Any:
    if candidate is None:
        return existing
    candidate_text = str(candidate).strip()
    if not candidate_text:
        return existing
    if not str(existing or "").strip():
        return candidate_text

    existing_text = str(existing).strip()
    if candidate_text == existing_text:
        return existing_text

    if field_key in MULTI_ROW_FIELDS or (
        _field_allows_multi_value(field_key) and ("\n" in existing_text or "\n" in candidate_text)
    ):
        if field_key == "3. Area of Cover":
            return _prefer_field_value(field_key, existing_text, candidate_text)
        combined_text = _merge_distinct_field_parts(existing_text, candidate_text)
        if combined_text:
            return combined_text

    existing_polarity = _coverage_polarity(existing_text)
    candidate_polarity = _coverage_polarity(candidate_text)
    if existing_polarity and candidate_polarity and existing_polarity != candidate_polarity:
        return candidate_text

    return _prefer_field_value(field_key, existing_text, candidate_text)


def merge_plan_records(existing: dict, incoming: dict, fallback_plan_name: str = "", fallback_underwriter: str = "") -> dict:
    merged = default_plan_record(
        plan_name=existing.get("0.2. Plan Name", fallback_plan_name),
        underwriter=existing.get("0.1. Underwriter", fallback_underwriter),
    )
    for source in (existing, incoming):
        for key, value in source.items():
            resolved_key = _resolve_category_key(key)
            merged[resolved_key] = merge_field_value(merged.get(resolved_key, ""), value, field_key=resolved_key)

    if fallback_underwriter:
        merged["0.1. Underwriter"] = merge_field_value(merged.get("0.1. Underwriter", ""), fallback_underwriter, field_key="0.1. Underwriter")
    if fallback_plan_name:
        merged["0.2. Plan Name"] = merge_field_value(merged.get("0.2. Plan Name", ""), fallback_plan_name, field_key="0.2. Plan Name")
    return clean_raw_fields(merged)


def normalize_plan_map(plan_map: dict, underwriter: str = "") -> dict:
    normalized = {}
    for raw_plan_name, fields in (plan_map or {}).items():
        if not isinstance(fields, dict):
            continue
        plan_name = (fields.get("0.2. Plan Name") or raw_plan_name or "Plan").strip()
        normalized[plan_name] = merge_plan_records(
            normalized.get(plan_name, {}),
            fields,
            fallback_plan_name=plan_name,
            fallback_underwriter=underwriter,
        )
    return normalized


def _is_option_like_plan_name(name: str) -> bool:
    normalized = _normalize_match_text(name)
    if not normalized:
        return False
    return normalized.startswith("option ") or normalized in {"option", "binding quote", "quote option"}


def _is_category_like_plan_name(name: str) -> bool:
    return _normalize_match_text(name).startswith("category ")


def _is_banner_like_plan_name(name: str) -> bool:
    normalized = _normalize_match_text(name)
    return normalized in {"binding quote", "quote summary", "quotation", "quote"}


def _plan_nonblank_field_count(fields: dict[str, Any]) -> int:
    if not isinstance(fields, dict):
        return 0
    return sum(
        1
        for key, value in fields.items()
        if key not in {"0.1. Underwriter", "0.2. Plan Name"} and not _is_blankish(value)
    )


def _choose_canonical_plan_name(plan_names: list[str], plan_map: dict[str, dict]) -> str:
    category_names = [name for name in plan_names if _is_category_like_plan_name(name)]
    if category_names:
        return max(category_names, key=lambda name: (_plan_nonblank_field_count(plan_map.get(name, {})), len(name)))

    non_option_names = [
        name for name in plan_names
        if not _is_option_like_plan_name(name) and not _is_banner_like_plan_name(name)
    ]
    if non_option_names:
        return max(non_option_names, key=lambda name: (_plan_nonblank_field_count(plan_map.get(name, {})), len(name)))

    return max(plan_names, key=lambda name: (_plan_nonblank_field_count(plan_map.get(name, {})), len(name)))


def _collapse_single_tob_alias_plans(
    plan_map: dict[str, dict],
    known_plan_names: list[str] | None = None,
    fallback_underwriter: str = "",
) -> tuple[dict[str, dict], list[str]]:
    normalized_map = normalize_plan_map(plan_map, fallback_underwriter)
    if len(normalized_map) <= 1:
        canonical_names = list(normalized_map.keys())
        if not canonical_names and known_plan_names:
            canonical_names = list(dict.fromkeys(name for name in known_plan_names if str(name).strip()))
        return normalized_map, canonical_names

    plan_names = list(normalized_map.keys())
    category_names = [name for name in plan_names if _is_category_like_plan_name(name)]
    option_like_names = [name for name in plan_names if _is_option_like_plan_name(name) or _is_banner_like_plan_name(name)]
    non_option_names = [
        name for name in plan_names
        if not _is_option_like_plan_name(name) and not _is_banner_like_plan_name(name)
    ]

    should_collapse = False
    if len(category_names) == 1 and option_like_names:
        should_collapse = True
    elif len(non_option_names) == 1 and len(option_like_names) == len(plan_names) - 1:
        should_collapse = True

    if not should_collapse:
        return normalized_map, list(dict.fromkeys(plan_names))

    canonical_name = _choose_canonical_plan_name(plan_names, normalized_map)
    merged_record = default_plan_record(canonical_name, fallback_underwriter)
    for plan_name in plan_names:
        merged_record = merge_plan_records(
            merged_record,
            normalized_map.get(plan_name, {}),
            fallback_plan_name=canonical_name,
            fallback_underwriter=fallback_underwriter,
        )

    collapsed_map = {canonical_name: clean_raw_fields(merged_record)}
    collapsed_names = [canonical_name]
    if known_plan_names:
        for name in known_plan_names:
            if name == canonical_name or name in plan_names:
                continue
            collapsed_names.append(name)
    return collapsed_map, list(dict.fromkeys(collapsed_names))

# ── PAGE COUNT ────────────────────────────────────────────────────────────────
def get_pdf_page_count(file_path: str) -> int:
    if not HAS_PYPDF2:
        return 0
    try:
        with open(file_path, "rb") as f:
            return len(PyPDF2.PdfReader(f).pages)
    except Exception as e:
        print(f"  Warning: Could not read PDF page count: {e}")
        return 0


DI_PAGE_SIGNAL_KEYWORDS = {
    "table of benefits",
    "schedule of benefits",
    "benefits",
    "coverage",
    "copay",
    "co pay",
    "coinsurance",
    "deductible",
    "network",
    "in patient",
    "out patient",
    "maternity",
    "dental",
    "optical",
    "pharmacy",
    "medication",
}


DI_COVER_PAGE_SIGNAL_KEYWORDS = {
    "quotation",
    "quote",
    "insurer",
    "underwriter",
    "category",
    "plan",
    "medical insurance",
    "health insurance",
    "schedule",
    "benefit",
}


def _normalize_di_page_text(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9%]+", " ", str(text or "").lower())).strip()


def _compress_page_numbers(page_numbers: list[int]) -> str:
    if not page_numbers:
        return ""

    ordered = sorted({int(page) for page in page_numbers if int(page) > 0})
    if not ordered:
        return ""

    ranges: list[str] = []
    start = ordered[0]
    end = ordered[0]
    for page in ordered[1:]:
        if page == end + 1:
            end = page
            continue
        ranges.append(f"{start}-{end}" if start != end else str(start))
        start = end = page
    ranges.append(f"{start}-{end}" if start != end else str(start))
    return ",".join(ranges)


def _extract_pdf_text_by_page(file_path: str) -> dict[int, str]:
    if not HAS_PYPDF2:
        return {}

    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            return {
                index + 1: (page.extract_text() or "")
                for index, page in enumerate(reader.pages)
            }
    except Exception as exc:
        print(f"  ⚠️  PDF page text pre-scan failed: {exc}")
        return {}


def _assess_document_intelligence_page_candidate(
    page_text: str,
    page_number: int,
    total_pages: int,
) -> dict[str, Any]:
    normalized_text = _normalize_di_page_text(page_text)
    keyword_hits = sorted(
        keyword for keyword in DI_PAGE_SIGNAL_KEYWORDS
        if keyword in normalized_text
    )
    cover_hits = sorted(
        keyword for keyword in DI_COVER_PAGE_SIGNAL_KEYWORDS
        if keyword in normalized_text
    )
    has_numeric_signal = bool(re.search(r"\b(?:bhd|aed|usd|sar)\b|\d+\s*%", normalized_text))

    should_include = bool(
        "table of benefits" in normalized_text
        or "schedule of benefits" in normalized_text
        or len(keyword_hits) >= 2
        or (keyword_hits and has_numeric_signal)
        or (page_number <= min(2, total_pages) and cover_hits)
    )

    if "table of benefits" in normalized_text or "schedule of benefits" in normalized_text:
        reason = "explicit_tob_heading"
    elif len(keyword_hits) >= 2:
        reason = "multiple_tob_keywords"
    elif keyword_hits and has_numeric_signal:
        reason = "benefit_keyword_with_numeric_signal"
    elif page_number <= min(2, total_pages) and cover_hits:
        reason = "early_cover_signal"
    else:
        reason = "no_strong_di_signal"

    return {
        "page": page_number,
        "should_include": should_include,
        "reason": reason,
        "keyword_hits": keyword_hits[:8],
        "cover_hits": cover_hits[:8],
        "has_numeric_signal": has_numeric_signal,
        "text_length": len(str(page_text or "").strip()),
    }


def _select_document_intelligence_pages(file_path: str, actual_page_count: int) -> dict[str, Any]:
    all_pages = list(range(1, actual_page_count + 1)) if actual_page_count > 0 else []
    default_result = {
        "mode": "full_document",
        "enabled": False,
        "pages_param": None,
        "requested_pages": all_pages,
        "skipped_pages": [],
        "fallback_reason": "",
        "page_assessments": [],
    }

    if actual_page_count <= 0 or not HAS_PYPDF2:
        default_result["fallback_reason"] = "page_prescan_unavailable"
        return default_result

    page_text_by_number = _extract_pdf_text_by_page(file_path)
    if not page_text_by_number:
        default_result["fallback_reason"] = "page_text_prescan_failed"
        return default_result

    page_assessments: list[dict[str, Any]] = []
    requested_pages: list[int] = []
    for page_number in all_pages:
        assessment = _assess_document_intelligence_page_candidate(
            page_text_by_number.get(page_number, ""),
            page_number=page_number,
            total_pages=actual_page_count,
        )
        page_assessments.append(assessment)
        if assessment["should_include"]:
            requested_pages.append(page_number)

    if not requested_pages:
        default_result["fallback_reason"] = "no_di_page_signals_detected"
        default_result["page_assessments"] = page_assessments
        return default_result

    if len(requested_pages) >= actual_page_count:
        default_result["fallback_reason"] = "all_pages_matched_signal"
        default_result["page_assessments"] = page_assessments
        return default_result

    skipped_pages = [page for page in all_pages if page not in set(requested_pages)]
    return {
        "mode": "tob_page_subset",
        "enabled": True,
        "pages_param": _compress_page_numbers(requested_pages),
        "requested_pages": requested_pages,
        "skipped_pages": skipped_pages,
        "fallback_reason": "",
        "page_assessments": page_assessments,
    }


# ── STEP 1: AZURE DOCUMENT INTELLIGENCE ──────────────────────────────────────
def _delete_cache_files(file_path: str):
    for ext in ("_di.pkl", "_di_meta.json", "_structured.md"):
        path = os.path.splitext(file_path)[0] + ext
        if os.path.exists(path):
            os.remove(path)
            print(f"  🗑️  Deleted stale cache: {path}")


def _load_di_cache_metadata(cache_meta_file: str) -> dict[str, Any] | None:
    if not os.path.exists(cache_meta_file):
        return None
    try:
        return read_json_file(cache_meta_file)
    except Exception as exc:
        print(f"  ⚠️  Corrupt DI cache metadata ({exc}) — re-fetching")
        try:
            os.remove(cache_meta_file)
        except OSError:
            pass
        return None


def _expected_di_cache_metadata(actual_page_count: int, di_page_filter: dict[str, Any]) -> dict[str, Any]:
    requested_pages = [int(page) for page in (di_page_filter.get("requested_pages") or [])]
    return {
        "source_page_count": int(actual_page_count or 0),
        "requested_pages": requested_pages,
        "mode": di_page_filter.get("mode", "full_document"),
        "pages_param": di_page_filter.get("pages_param"),
    }


def _try_load_cache(
    cache_file: str,
    cache_meta_file: str,
    actual_page_count: int,
    di_page_filter: dict[str, Any],
):
    if not os.path.exists(cache_file):
        return None

    cache_meta = _load_di_cache_metadata(cache_meta_file)
    expected_meta = _expected_di_cache_metadata(actual_page_count, di_page_filter)
    if cache_meta != expected_meta:
        print("  ⚠️  DI cache metadata mismatch — re-fetching")
        try:
            os.remove(cache_file)
        except OSError:
            pass
        return None

    try:
        with open(cache_file, "rb") as f:
            result = pickle.load(f)
    except Exception as e:
        print(f"  ⚠️  Corrupt cache file ({e}) — deleting and re-fetching")
        os.remove(cache_file)
        return None

    cached_pages = len(result.pages) if result.pages else 0
    print(f"  Cached result has {cached_pages} pages")

    expected_pages = len(expected_meta["requested_pages"]) if expected_meta["requested_pages"] else actual_page_count

    if expected_pages > 0 and cached_pages != expected_pages:
        print(f"  WARNING: Page mismatch (cached={cached_pages}, expected={expected_pages}) — re-fetching")
        os.remove(cache_file)
        try:
            os.remove(cache_meta_file)
        except OSError:
            pass
        return None
    if cached_pages < 1:
        print(f"  WARNING: Cache has 0 pages — re-fetching")
        os.remove(cache_file)
        try:
            os.remove(cache_meta_file)
        except OSError:
            pass
        return None

    return result


def _call_azure_di(file_path: str, pages: str | None = None) -> Any:
    from io import BytesIO
    import urllib.request

    print(f"  Scanning via Azure DI: {file_path}")
    if pages:
        print(f"  Azure DI page filter: {pages}")

    with open(file_path, "rb") as f:
        pdf_bytes = f.read()

    if len(pdf_bytes) == 0:
        raise ValueError(f"File is empty: {file_path}")

    print(f"  File size: {len(pdf_bytes):,} bytes")
    print(f"  AZURE_DOC_ENDPOINT loaded as: '{AZURE_ENDPOINT}'")
    print(f"  AZURE_DOC_KEY loaded as: '{AZURE_KEY[:6]}...{AZURE_KEY[-4:]}' (masked)")

    if not AZURE_ENDPOINT or not AZURE_ENDPOINT.startswith("http"):
        raise RuntimeError(
            "AZURE_DOC_ENDPOINT is empty or invalid.\n"
            "  Check your .env file has: AZURE_DOC_ENDPOINT=https://YOUR-RESOURCE.cognitiveservices.azure.com/"
        )
    if not AZURE_KEY or len(AZURE_KEY) < 10:
        raise RuntimeError(
            "AZURE_DOC_KEY is empty or too short.\n"
            "  Check your .env file has: AZURE_DOC_KEY=<your 32-char key from Azure Portal>"
        )

    endpoint = AZURE_ENDPOINT.rstrip("/")
    api_versions_to_try = ["2024-11-30", "2024-02-29-preview", "2023-07-31"]
    working_api_version = None

    for api_ver in api_versions_to_try:
        test_url = f"{endpoint}/documentintelligence/documentModels?api-version={api_ver}"
        try:
            req = urllib.request.Request(
                test_url,
                headers={"Ocp-Apim-Subscription-Key": AZURE_KEY},
            )
            with urllib.request.urlopen(req, timeout=15) as resp:
                working_api_version = api_ver
                print(f"  ✓ Azure DI credentials OK — API version {api_ver} works (HTTP {resp.status})")
                break
        except urllib.error.HTTPError as he:
            body = he.read().decode("utf-8", errors="replace")[:400]
            if he.code in (401, 403):
                raise RuntimeError(
                    f"Azure DI authentication failed: HTTP {he.code}\n"
                    f"  Endpoint: {test_url}\n"
                    f"  Response: {body}\n"
                    f"  → Your AZURE_DOC_KEY is wrong. Copy it fresh from Azure Portal → "
                    f"Keys and Endpoint page."
                )
            elif he.code == 404:
                print(f"  ⚠️  API version {api_ver} → HTTP 404, trying next version...")
                continue
            elif he.code == 429:
                raise RuntimeError(
                    f"Azure DI rate limit hit (HTTP 429). Wait 60 seconds and retry."
                )
            else:
                print(f"  ⚠️  API version {api_ver} → HTTP {he.code}: {body[:100]}")
                continue
        except Exception as conn_err:
            print(f"  ⚠️  Pre-flight network error: {conn_err} — proceeding anyway")
            break

    if working_api_version is None:
        print(
            f"  ⚠️  Could not verify credentials with any API version.\n"
            f"  Endpoint: {endpoint}\n"
            f"  Proceeding anyway — SDK call may still work."
        )

    e1 = None
    try:
        poller = doc_client.begin_analyze_document(
            model_id="prebuilt-layout",
            body=pdf_bytes,
            content_type="application/pdf",
            pages=pages,
        )
        result = poller.result()
        print(f"  ✓ Strategy 1 succeeded")
        return result
    except Exception as exc:
        e1 = exc
        print(f"  ⚠️  Strategy 1 (bytes + application/pdf) failed: {exc}")

    e2 = None
    try:
        poller = doc_client.begin_analyze_document(
            model_id="prebuilt-layout",
            body=BytesIO(pdf_bytes),
            content_type="application/octet-stream",
            pages=pages,
        )
        result = poller.result()
        print(f"  ✓ Strategy 2 succeeded")
        return result
    except Exception as exc:
        e2 = exc
        print(f"  ⚠️  Strategy 2 (BytesIO + octet-stream) failed: {exc}")

    e3 = None
    try:
        import tempfile, pathlib
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_bytes)
            tmp_path = tmp.name
        from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
        poller = doc_client.begin_analyze_document(
            model_id="prebuilt-layout",
            body=AnalyzeDocumentRequest(url_source=pathlib.Path(tmp_path).as_uri()),
            pages=pages,
        )
        result = poller.result()
        print(f"  ✓ Strategy 3 (URL source) succeeded")
        return result
    except Exception as exc:
        e3 = exc
        print(f"  ⚠️  Strategy 3 (URL source) failed: {exc}")

    raise RuntimeError(
        f"All 3 Azure DI strategies failed for: {file_path}\n"
        f"  Strategy 1 (bytes/pdf):      {e1}\n"
        f"  Strategy 2 (BytesIO/stream): {e2}\n"
        f"  Strategy 3 (URL source):     {e3}\n\n"
        f"Most likely cause: HTTP error from Azure (401/403/404/429).\n"
        f"  1. Verify AZURE_DOC_ENDPOINT ends with a trailing slash '/'\n"
        f"  2. Verify AZURE_DOC_KEY is the correct 32-char key from Azure Portal\n"
        f"  3. Check your Azure DI resource is active and not rate-limited\n"
        f"  4. Confirm SDK version: pip show azure-ai-documentintelligence"
    )


def get_structured_content(file_path: str) -> tuple[Any, dict[str, Any]]:
    cache_file = os.path.splitext(file_path)[0] + "_di.pkl"
    cache_meta_file = os.path.splitext(file_path)[0] + "_di_meta.json"
    actual_page_count = get_pdf_page_count(file_path)
    if actual_page_count > 0:
        print(f"  PDF has {actual_page_count} pages")

    di_page_filter = _select_document_intelligence_pages(file_path, actual_page_count)
    requested_pages = di_page_filter.get("requested_pages") or []
    if di_page_filter.get("enabled"):
        print(
            f"  ✓ DI TOB page filter enabled: {len(requested_pages)}/{actual_page_count} pages "
            f"({di_page_filter.get('pages_param')})"
        )
    elif di_page_filter.get("fallback_reason"):
        print(f"  ℹ️  DI page filter fallback: {di_page_filter['fallback_reason']}")

    cached_result = _try_load_cache(cache_file, cache_meta_file, actual_page_count, di_page_filter)
    if cached_result is not None:
        print(f"  ✓ Using cached DI result ({len(cached_result.pages)} pages)")
        return cached_result, di_page_filter

    result = _call_azure_di(file_path, pages=di_page_filter.get("pages_param"))

    pages_returned = len(result.pages) if result.pages else 0
    print(f"  Azure DI returned {pages_returned} pages")
    expected_pages = len(requested_pages) if di_page_filter.get("enabled") else actual_page_count
    if expected_pages > 0 and pages_returned != expected_pages:
        print(f"  ⚠️  Expected {expected_pages} DI pages but got {pages_returned}")
    print(f"  Tables:     {len(result.tables)     if result.tables     else 0}")
    print(f"  Paragraphs: {len(result.paragraphs) if result.paragraphs else 0}")
    print(f"  Content length: {len(result.content) if result.content else 0} chars")

    if pages_returned > 0:
        with open(cache_file, "wb") as f:
            pickle.dump(result, f)
        write_json_file(cache_meta_file, _expected_di_cache_metadata(actual_page_count, di_page_filter))
        print(f"  ✓ Result cached to {cache_file}")
    else:
        print(f"  ⚠️  Not caching — 0 pages returned")

    return result, di_page_filter


def result_to_structured_markdown(result: Any) -> str:
    sections = []

    for table_idx, table in enumerate(result.tables or []):
        page_num = 1
        if table.bounding_regions:
            page_num = table.bounding_regions[0].page_number
        rows, cols = table.row_count, table.column_count
        grid = [["" for _ in range(cols)] for _ in range(rows)]
        for cell in table.cells:
            grid[cell.row_index][cell.column_index] = cell.content.replace("\n", " ").strip()
        md_lines = []
        for row_idx, row in enumerate(grid):
            md_lines.append("| " + " | ".join(row) + " |")
            if row_idx == 0:
                md_lines.append("|" + "|".join(["---"] * cols) + "|")
        table_md = f"\n<!-- TABLE {table_idx + 1} (Page {page_num}) -->\n" + "\n".join(md_lines)
        offset = table.spans[0].offset if table.spans else 0
        sections.append((page_num, offset, table_md))

    for para in result.paragraphs or []:
        content = para.content.strip()
        if not content:
            continue
        page_num = para.bounding_regions[0].page_number if para.bounding_regions else 1
        offset = para.spans[0].offset if para.spans else 0
        role = getattr(para, "role", None)
        if role == "title":
            content = f"\n# {content}"
        elif role == "sectionHeading":
            content = f"\n## {content}"
        elif role == "pageHeader":
            content = f"\n<!-- HEADER: {content} -->"
        elif role == "pageFooter":
            content = f"\n<!-- FOOTER: {content} -->"
        elif role == "pageNumber":
            continue
        sections.append((page_num, offset, content))

    sections.sort(key=lambda x: (x[0], x[1]))
    output_parts, current_page = [], None
    for page_num, offset, content in sections:
        if page_num != current_page:
            output_parts.append(f"\n\n---\n## PAGE {page_num}\n---\n")
            current_page = page_num
        output_parts.append(content)
    return "\n".join(output_parts)


def clear_cache(file_path: str):
    _delete_cache_files(file_path)


def get_ocr_text(file_path: str) -> tuple[Any, str]:
    md_cache = os.path.splitext(file_path)[0] + "_structured.md"
    result, di_page_filter = get_structured_content(file_path)
    result_page_count = len(getattr(result, "pages", []) or [])

    if os.path.exists(md_cache):
        try:
            with open(md_cache, "r", encoding="utf-8") as f:
                structured_md = f.read()
            cached_md_page_count = len(re.findall(r'^## PAGE \d+$', structured_md, flags=re.MULTILINE))
            if len(structured_md.strip()) < 50:
                print(f"  ⚠️  Cached markdown is too short — regenerating")
                os.remove(md_cache)
                structured_md = None
            elif result_page_count > 0 and cached_md_page_count != result_page_count:
                print(
                    f"  ⚠️  Cached markdown page mismatch "
                    f"(markdown={cached_md_page_count}, DI={result_page_count}) — regenerating"
                )
                os.remove(md_cache)
                structured_md = None
            else:
                print(f"  Loading cached markdown ({len(structured_md):,} chars)")
        except Exception as e:
            print(f"  ⚠️  Markdown cache read error ({e}) — regenerating")
            os.remove(md_cache)
            structured_md = None
    else:
        structured_md = None

    if structured_md is None:
        structured_md = result_to_structured_markdown(result)
        with open(md_cache, "w", encoding="utf-8") as f:
            f.write(structured_md)
        print(f"  Structured markdown saved ({len(structured_md):,} chars) → {md_cache}")

    log_document_intelligence_output(file_path, result, structured_md, di_page_filter=di_page_filter)

    return result, structured_md


# ── CATEGORIES ────────────────────────────────────────────────────────────────
# Comprehensive synonyms for LLM semantic matching - includes copay, coinsurance, waiting periods, conditions
CATEGORIES = [
    ('0.1', 'Underwriter', 'Insurance company name, Insurer, Carrier, Provider, Insurance Provider, Underwriting Company, Insurance Firm'),
    ('0.2', 'Plan Name', 'Plan/category name, Product Name, Scheme Name, Policy Name, Plan Type, Category, Option, Tier, Level'),
    ('2',  'Annual Limit per person', 'Overall annual limit, Sum insured, Maximum benefit, Annual aggregate limit, AAL, Total coverage, Per person limit, Maximum annual benefit, Aggregate annual limit, Policy limit, Indemnity limit, Indemnity amount'),
    ('3',  'Area of Cover', 'Geographic coverage, Territory, Coverage area, Territorial scope, Worldwide, Regional, GCC, UAE only, International coverage, Global coverage, Country coverage'),
    ('4', 'TPA', 'Third party administrator, Claims administrator, TPA name, Administration company, Claims handler company'),
    ('5', 'Network', 'Medical network name, Provider network name, Hospital network name, Network name, Network type, Premium network, Standard network, Enhanced network, Provider panel name'),
    ('6', 'Pre-existing and Chronic Conditions', 'Pre-existing conditions, Chronic conditions, Prior conditions, Existing illness, Chronic disease, Long-term conditions, PED, Pre-existing disease, Chronic illness, Waiting period for pre-existing, Coinsurance for chronic, Copay for chronic'),
    ('7', 'In-Patient Benefits', 'In-patient benefits, Hospitalization, Inpatient care, Hospital stay, Room and board, Medical treatment, Surgery, Inpatient services, Hospital admission, Inpatient coverage, Inpatient copay, Inpatient coinsurance'),
    ('8',  'Room Type', 'Hospital room, Accommodation type, Room and board, Private room, Semi-private, Shared room, Ward, Suite, Single room, Standard room, Deluxe room, Room category'),
    ('9', 'Parent Accommodation for child under 18 years of age', 'Parent accommodation, Guardian stay, Parent room, Accompanying parent, Parent bed, Guardian accommodation, Parent lodging'),
    ('10', 'Accommodation of an accompanying person', 'Companion accommodation, Attendant stay, Family accommodation, Escort stay, Caregiver accommodation, Accompanying person, Attendant lodging'),
    ('11', 'Home Nursing', 'Home nursing, Home healthcare, Skilled nursing, Home care, Domiciliary care, Nursing at home, Home nurse visits, Home nursing copay, Home healthcare limit, Nursing visits'),
    ('12', 'Out Patient Deductible on Consultation', 'OPD deductible, Consultation deductible, Outpatient copay, OP copayment, Consultation copay, Doctor visit copay, Deductible per visit, Co-pay consultation, Coinsurance OPD, OP deductible, Consultation fee, Visit copay, OPD co-payment'),
    ('13', 'Physician consultation charges', 'Doctor consultation, GP consultation, Specialist consultation, Physician fees, Consultation charges, Doctor visit, Medical consultation, Copay for consultation, Coinsurance for doctor, GP visit, Specialist fees, Doctor fees, Consultation limit'),
    ('14', 'Diagnostic Tests & Laboratory Tests', 'Lab tests, Diagnostics, X-ray, MRI, CT scan, Blood tests, Pathology, Radiology, Imaging, Laboratory, Copay for tests, Coinsurance for diagnostics, Diagnostic copay, Lab copay, Investigation charges'),
    ('15', 'Prescribed Medication', 'Medicine, Drugs, Pharmacy, Prescription drugs, Medication coverage, Pharmaceuticals, Drug copay, Medicine copay, Pharmacy copay, Coinsurance for medication, Prescription copay, Medication limit, Drug coverage'),
    ('16', 'Physiotherapy', 'Physical therapy, Physio, Rehabilitation, Rehab, PT sessions, Physiotherapy sessions, Copay for physio, Sessions per year, Physiotherapy limit, Rehab sessions, Physio copay, Physiotherapy coinsurance, Waiting period for physio'),
    ('17', 'Alternative Medical Treatment', 'Alternative medicine, Complementary medicine, Homeopathy, Ayurveda, Acupuncture, Chiropractic, Traditional medicine, AYUSH, Osteopathy, Naturopathy, Chinese medicine, Herbal medicine, Alternative therapy copay'),
    ('18', 'Maternity In Patient Services', 'Maternity inpatient, Delivery, Childbirth, Normal delivery, C-section, Caesarean, Medical necessary C-section, Maternity hospitalization, Labor, Confinement, Maternity waiting period, Maternity copay, Coinsurance maternity, Delivery charges, Birth, Maternity limit, Maternity sublimit, In-patient maternity annual maximum limit, Maternity benefits, Legal abortion, Inpatient coinsurance maternity'),
    ('19', 'Maternity Out Patient Services', 'Maternity outpatient, Prenatal, Antenatal, Postnatal, Pregnancy checkups, OB-GYN visits, Maternity OPD, Pre-natal copay, Antenatal visits, Postnatal care, Pregnancy consultation, Maternity OP copay, Routine out-patient maternity, Eligible females covered'),
    ('20', 'Life Threatening Maternity Complications', 'Maternity complications, Emergency maternity, High-risk pregnancy, Ectopic pregnancy, Miscarriage, Pregnancy emergency, Complicated delivery, Emergency C-section, Maternity emergency, Complicated in-patient maternity'),
    ('21', 'New Born Cover', 'Newborn coverage, Baby cover, Infant coverage, Newborn from day one, Newborn hospital stay, Neonatal, NICU, Newborn waiting period, Baby insurance, Day one cover, Newborn limit, Infant care'),
    ('22', 'Dental Benefit', 'Dental coverage, Dental treatment, Tooth, Teeth, Oral care, Dental checkup, Dental copay, Dental coinsurance, Dental waiting period, Dental sublimit, Dental annual limit, Dental extraction, Filling, Root canal, Dental cleaning'),
    ('23', 'Optical Benefit', 'Optical coverage, Vision, Eye care, Eyeglasses, Spectacles, Contact lenses, Eye exam, Vision copay, Optical sublimit, Optical waiting period, Frames, Lenses, Eye test, Vision screening, Optical limit'),
    ('24', 'Ambulance', 'Ambulance services, Emergency transport, Medical transportation, Air ambulance, Ground ambulance, Evacuation, Emergency ambulance, Road ambulance, Ambulance charges, Ambulance copay'),
    ('25', 'Psychiatric', 'Psychiatric coverage, Mental health, Psychology, Counseling, Therapy, Behavioral health, Psychiatric hospitalization, Mental illness, Depression, Anxiety, Psychiatric copay, Mental health coinsurance, Psychotherapy, Mental health waiting period'),
    ('26', 'Organ Transplant', 'Organ transplant, Transplantation, Organ donor, Kidney transplant, Liver transplant, Heart transplant, Bone marrow transplant, Transplant waiting period, Transplant surgery, Organ recipient, Transplant copay'),
    ('27', 'Kidney Dialysis Treatment', 'Dialysis, Hemodialysis, Kidney dialysis, Renal dialysis, Dialysis sessions, Chronic kidney, Dialysis copay, Dialysis coinsurance, Dialysis limit, Renal treatment'),
    ('28', 'Repatriation', 'Repatriation of remains, Mortal remains, Body repatriation, Repatriation of body, Medical repatriation, Repatriation benefit, Return of remains'),
    ('29', 'Vaccinations', 'Vaccination, Immunization, Vaccines, Shots, Inoculation, Preventive vaccination, Child vaccination, Adult vaccination, Vaccine copay, Immunization coverage, Vaccination limit'),
    ('30', 'Preventive Services', 'Preventive care, Wellness, Health checkup, Annual checkup, Screening, Preventive screening, Health screening, Wellness exam, Preventive copay, Annual physical, Health check, Wellness benefit, Screening tests'),
    ('31', 'Adult Pnuemococcal Conjugate Vaccine', 'Pneumococcal vaccine, Pneumonia vaccine, PCV, PPSV, Pneumococcal vaccination, Pneumonia shot'),
    ('32', 'Cancer', 'Cancer treatment, Oncology, Chemotherapy, Radiation therapy, Cancer screening, Tumor, Malignancy, Cancer copay, Oncology coinsurance, Cancer waiting period, Cancer limit, Radiation, Chemo'),
    ('34', 'Return Air Fare', 'Return airfare, Air ticket, Travel benefit, Flight ticket, Air travel, Medical travel, Return flight, Air fare benefit, Travel reimbursement'),
    ('35', 'Work Related Injury', 'Work injury, Occupational injury, Workplace accident, Workers compensation, Job-related injury, Employment injury, On-duty injury, Work accident'),
    ('36', 'Influenza Vaccine', 'Flu vaccine, Influenza shot, Seasonal flu, Flu shot, Annual flu vaccination, Flu immunization'),
    ('37', 'HCV Hepatitis C Virus Infection Screening', 'Hepatitis C screening, HCV test, Hepatitis C test, HCV screening, Hep C screening'),
    ('38', 'Hepatitis B Virus Screening and treatment', 'Hepatitis B screening, HBV test, Hepatitis B treatment, HBV vaccination, Hep B screening, Hepatitis B vaccine'),
    ('33', 'Prosthetics', 'Prosthetics, Artificial limbs, Prosthesis, Orthopedic devices, Artificial body parts, Implants, Medical devices, Prosthetic limit, Artificial organs, Limb prosthesis'),
    ('39', 'UAE Within Network', 'Claims Inside Network, In-network claims, Network provider claims, Direct billing, Cashless, In-network coinsurance, In-network copay, Network reimbursement, Panel provider, In-network percentage, Network coverage'),
    ('40', 'Claims Outside Network Within Country', 'Out-of-network claims, Non-network claims, Reimbursement claims, Out-of-network copay, Out-of-network coinsurance, Non-panel claims, Non-network reimbursement, OON coverage, Out-of-network percentage'),
    ('41', 'Claims Outside Country', 'International claims, Overseas claims, Foreign claims, Cross-border claims, International reimbursement, Worldwide claims, Abroad claims, International coverage, Overseas treatment'),
]

CATEGORIES_TEXT = "\n".join(
    f'{num}. {name}\n   Find: {hint}'
    for num, name, hint in CATEGORIES
)

CATEGORY_LOOKUP = {
    f"{num}. {name}": {"number": num, "name": name, "hint": hint}
    for num, name, hint in CATEGORIES
}


def _report_ordered_keys() -> list[str]:
    return [f"{num}. {name}" for num, name, _ in CATEGORIES]


def _normalize_match_text(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9]+", " ", str(text or "").lower())).strip()


def _match_tokens(text: str) -> set[str]:
    stopwords = {
        "a", "an", "and", "are", "as", "at", "by", "for", "from", "in", "is", "of", "on", "or",
        "the", "this", "to", "under", "with"
    }
    return {
        token for token in _normalize_match_text(text).split()
        if len(token) > 2 and token not in stopwords
    }


def _resolve_category_key(raw_key: str) -> str:
    raw_key = str(raw_key or "").strip()
    if raw_key in CATEGORY_LOOKUP:
        return raw_key

    raw_norm = _normalize_match_text(raw_key)
    raw_tokens = _match_tokens(raw_key)
    best_key = raw_key
    best_score = 0.0

    for category_key, category in CATEGORY_LOOKUP.items():
        name_norm = _normalize_match_text(category["name"])
        hint_norm = _normalize_match_text(category["hint"])
        name_tokens = _match_tokens(category["name"])
        hint_tokens = _match_tokens(category["hint"])

        overlap_name = len(raw_tokens & name_tokens) / max(len(name_tokens), 1)
        overlap_hint = len(raw_tokens & hint_tokens) / max(min(len(hint_tokens), 10), 1)
        ratio_name = SequenceMatcher(None, raw_norm, name_norm).ratio()
        ratio_hint = SequenceMatcher(None, raw_norm, hint_norm).ratio()
        score = (overlap_name * 0.55) + (overlap_hint * 0.15) + (ratio_name * 0.25) + (ratio_hint * 0.05)

        if raw_norm == name_norm:
            score += 0.3

        if score > best_score:
            best_key = category_key
            best_score = score

    return best_key if best_score >= 0.45 else raw_key


def _parse_markdown_table_rows(page_content: str) -> list[dict[str, list[str] | str]]:
    rows: list[dict[str, list[str] | str]] = []
    pending_label = ""
    inherited_values: list[str] | None = None
    line_number = 0

    for line in str(page_content or "").splitlines():
        line_number += 1
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue

        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if len(cells) < 2:
            continue
        if all(re.fullmatch(r"-+", cell or "") for cell in cells):
            continue

        label = cells[0].strip()
        values = [cell.strip() for cell in cells[1:]]
        if not label:
            continue

        if all(not value for value in values):
            pending_label = " ".join(part for part in [pending_label, label] if part).strip()
            continue

        if label.endswith(";") and any(values):
            inherited_values = values[:]
        elif inherited_values and all(not value for value in values):
            values = inherited_values[:]

        effective_label = " ".join(part for part in [pending_label, label] if part).strip()
        rows.append(
            {
                "label": effective_label or label,
                "raw_label": label,
                "parent_label": pending_label,
                "values": values,
                "is_grouped": bool(pending_label),
                "line_number": str(line_number),
            }
        )
        pending_label = ""

    return rows


def _format_page_attribute_candidates(page_content: str) -> str:
    rows = _parse_markdown_table_rows(page_content)
    labels = []
    seen = set()
    for row in rows:
        label = str(row["label"]).strip()
        if not label:
            continue
        key = _normalize_match_text(label)
        if key in seen:
            continue
        seen.add(key)
        parent_label = str(row.get("parent_label", "")).strip()
        if parent_label:
            labels.append(f"- {parent_label} -> {label}")
        else:
            labels.append(f"- {label}")
    return "\n".join(labels) if labels else "- none"


def _candidate_value_entry(
    row_idx: int,
    row: dict[str, Any],
    plan_idx: int,
    score: float,
    page_number: int | None = None,
) -> dict[str, Any] | None:
    values = row.get("values", [])
    if plan_idx >= len(values):
        return None

    candidate_value = str(values[plan_idx]).strip()
    if not candidate_value:
        return None

    return {
        "row_idx": row_idx,
        "page_number": page_number,
        "row_ref": f"p{page_number}_r{row_idx}" if page_number is not None else f"r{row_idx}",
        "label": str(row.get("label", "")).strip(),
        "parent_label": str(row.get("parent_label", "")).strip(),
        "raw_label": str(row.get("raw_label", "")).strip(),
        "candidate_value": candidate_value,
        "score": score,
        "is_grouped": bool(row.get("is_grouped")),
        "line_number": str(row.get("line_number", "")),
    }


def _single_plan_candidate_value_entry(
    row_idx: int,
    row: dict[str, Any],
    score: float,
    page_number: int | None = None,
) -> dict[str, Any] | None:
    values = [str(value or "").strip() for value in row.get("values", [])]
    non_empty_values = [value for value in values if value]
    if len(non_empty_values) != 1:
        return None

    return {
        "row_idx": row_idx,
        "page_number": page_number,
        "row_ref": f"p{page_number}_r{row_idx}" if page_number is not None else f"r{row_idx}",
        "label": str(row.get("label", "")).strip(),
        "parent_label": str(row.get("parent_label", "")).strip(),
        "raw_label": str(row.get("raw_label", "")).strip(),
        "candidate_value": non_empty_values[0],
        "score": score,
        "is_grouped": bool(row.get("is_grouped")),
        "line_number": str(row.get("line_number", "")),
    }


def _semantic_field_label_match(field_key: str, row_label: str, row_section: str = "") -> bool:
    category = CATEGORY_LOOKUP.get(field_key)
    if not category:
        return False

    row_text = " ".join(part for part in [row_section, row_label] if str(part or "").strip())
    row_norm = _normalize_match_text(row_text)
    if not row_norm:
        return False

    category_name = str(category.get("name", "") or "").strip()
    category_hint = str(category.get("hint", "") or "").strip()
    name_norm = _normalize_match_text(category_name)
    hint_norm = _normalize_match_text(category_hint)
    row_tokens = _match_tokens(row_text)
    name_tokens = _match_tokens(category_name)
    hint_tokens = _match_tokens(category_hint)

    if name_norm and (row_norm == name_norm or name_norm in row_norm or row_norm in name_norm):
        return True

    significant_name_tokens = {token for token in name_tokens if len(token) > 3}
    if significant_name_tokens and len(row_tokens & significant_name_tokens) >= max(1, min(2, len(significant_name_tokens))):
        return True

    significant_hint_tokens = {
        token for token in hint_tokens
        if len(token) > 4 and token not in {"coverage", "benefit", "medical", "services", "service", "charges", "treatment"}
    }
    if significant_hint_tokens and len(row_tokens & significant_hint_tokens) >= 2:
        return True

    return False


def _row_matches_field(field_key: str, row_label: str, row_section: str = "") -> bool:
    if field_key not in CATEGORY_LOOKUP:
        return False

    rules = FIELD_CONTEXT_RULES.get(field_key)
    if rules:
        return _context_guard_allows_row(field_key, row_label, row_section)

    return _semantic_field_label_match(field_key, row_label, row_section)


def _collect_field_candidate_entries(
    table_rows: list[dict[str, Any]],
    field_key: str,
    plan_idx: int,
    single_plan_mode: bool = False,
    page_number: int | None = None,
) -> list[dict[str, Any]]:
    candidate_entries: list[dict[str, Any]] = []
    for row_idx, row in enumerate(table_rows):
        row_label = str(row.get("label", "") or "").strip()
        row_section = str(row.get("parent_label", "") or row.get("section", ""))
        if not row_label or not _row_matches_field(field_key, row_label, row_section):
            continue

        entry = _candidate_value_entry(row_idx, row, plan_idx, 1.0, page_number=page_number)
        if not entry and single_plan_mode:
            entry = _single_plan_candidate_value_entry(row_idx, row, 1.0, page_number=page_number)
        if entry:
            candidate_entries.append(entry)
    return candidate_entries


def _should_call_field_aggregation_llm(field_key: str, current_value: Any, candidate_entries: list[dict[str, Any]]) -> bool:
    if not candidate_entries:
        return False
    if field_key in MULTI_ROW_FIELDS:
        return _field_value_preference_score(field_key, _compose_multi_row_field_value(field_key, candidate_entries, fallback_value=current_value)) > _field_value_preference_score(field_key, current_value)
    if _field_behavior(field_key) == "strict_single":
        return True
    if len(candidate_entries) < 2:
        return False
    if any(entry.get("is_grouped") for entry in candidate_entries):
        return True
    if str(current_value or "").strip():
        return True
    return len(candidate_entries) >= 2


def _select_best_candidate_entries(candidate_entries: list[dict[str, Any]], max_candidates: int = 6) -> list[dict[str, Any]]:
    return sorted(
        candidate_entries,
        key=lambda item: (
            int(item.get("page_number") or 0),
            int(item.get("row_idx") or 0),
            str(item.get("label", "")),
        ),
    )[:max_candidates]


def _normalize_field_decision_result(result: dict[str, Any] | None) -> dict[str, Any]:
    result = result or {}
    decision_mode = str(result.get("decision_mode", "")).strip().lower()
    value = str(result.get("value", "") or "").strip()
    used_row_indexes = result.get("used_row_indexes", [])
    if not isinstance(used_row_indexes, list):
        used_row_indexes = []
    return {
        "decision_mode": decision_mode if decision_mode in {"blank", "single", "multi"} else "blank",
        "value": value,
        "used_row_indexes": [int(idx) for idx in used_row_indexes if isinstance(idx, int) or str(idx).isdigit()],
    }


def _decide_field_values_with_llm(decision_requests: list[dict[str, Any]]) -> dict[str, dict[str, dict[str, Any]]]:
    if not decision_requests:
        return {}

    cache_key = json.dumps(decision_requests, ensure_ascii=False, sort_keys=True, default=str)
    if cache_key in _field_decision_batch_cache:
        return json.loads(json.dumps(_field_decision_batch_cache[cache_key]))

    system_prompt = """You decide final extracted values for ambiguous insurance benefit fields after the full PDF has been processed.

<output_shape>
- Return JSON only.
- Schema:
  {
    "plans": {
      "<plan_name>": {
        "<field_key>": {
          "decision_mode": "blank" | "single" | "multi",
          "value": string,
          "used_row_indexes": number[]
        }
      }
    }
  }
- Return decisions only for the requested plan names and field keys.
</output_shape>

<decision_rules>
- Use only the provided candidate rows for each plan field.
- Respect row grouping and header relationships when deciding whether multiple rows belong together.
- Return `multi` only when multiple candidate rows clearly belong to the same benefit context and should be combined.
- For multi values, return multi-line text with one related item per line.
- Do not mix unrelated benefits just because they are nearby on the page.
- Aggregate fields like annual limit must stay distinct from benefit-specific sublimits such as maternity, dental, or optical limits.
- If one row is clearly the correct value, return `single`.
- If none are valid, return `blank`.
- Preserve exact wording from the candidate rows. Do not summarize or normalize.
</decision_rules>"""

    user_prompt = f"""Decide the best extracted values for all ambiguous plan fields in this PDF.

INPUT JSON:
{json.dumps({"requests": decision_requests}, indent=2, ensure_ascii=False)}

Return JSON only."""

    try:
        content, usage = llm_chat(
            system_prompt,
            user_prompt,
            json_mode=True,
            log_label="field_value_decision",
            task_class="extraction",
            log_context={"decision_count": len(decision_requests)},
        )
        tracker.record(_current_pdf, usage)
        result = safe_json_loads(content, "field value decision batch")
    except Exception as exc:
        print(f"  ⚠️  Batched field aggregation failed: {exc}")
        result = {}

    raw_plans = result.get("plans", result) if isinstance(result, dict) else {}
    normalized_results: dict[str, dict[str, dict[str, Any]]] = {}
    for request in decision_requests:
        plan_name = request.get("plan_name", "")
        field_key = request.get("field_key", "")
        raw_plan_result = raw_plans.get(plan_name, {}) if isinstance(raw_plans, dict) else {}
        raw_field_result = raw_plan_result.get(field_key, {}) if isinstance(raw_plan_result, dict) else {}
        normalized_results.setdefault(plan_name, {})[field_key] = _normalize_field_decision_result(raw_field_result)

    _field_decision_batch_cache[cache_key] = json.loads(json.dumps(normalized_results))
    return normalized_results


def _register_semantic_candidate_entries(
    candidate_state: dict[str, dict[str, dict[str, Any]]],
    plan_name: str,
    field_key: str,
    current_value: Any,
    candidate_entries: list[dict[str, Any]],
):
    if not candidate_entries:
        return

    plan_bucket = candidate_state.setdefault(plan_name, {})
    field_bucket = plan_bucket.setdefault(
        field_key,
        {
            "current_value": str(current_value or "").strip(),
            "candidate_entries": [],
        },
    )

    if str(current_value or "").strip() and not str(field_bucket.get("current_value", "")).strip():
        field_bucket["current_value"] = str(current_value).strip()

    existing_entries = field_bucket["candidate_entries"]
    seen = {
        (
            entry.get("page_number"),
            entry.get("row_idx"),
            _normalize_match_text(entry.get("label", "")),
            str(entry.get("candidate_value", "")).strip(),
        )
        for entry in existing_entries
    }

    for entry in candidate_entries:
        dedupe_key = (
            entry.get("page_number"),
            entry.get("row_idx"),
            _normalize_match_text(entry.get("label", "")),
            str(entry.get("candidate_value", "")).strip(),
        )
        if dedupe_key in seen:
            continue
        existing_entries.append(dict(entry))
        seen.add(dedupe_key)


def _resolve_deferred_semantic_candidates(
    plan_map: dict[str, dict],
    candidate_state: dict[str, dict[str, dict[str, Any]]],
) -> tuple[dict[str, dict], dict[str, dict[str, Any]]]:
    resolved_map = normalize_plan_map(plan_map)
    resolution_audit: dict[str, dict[str, Any]] = {}
    pending_decisions: dict[str, dict[str, dict[str, Any]]] = {}
    decision_requests: list[dict[str, Any]] = []

    for plan_name, field_map in (candidate_state or {}).items():
        plan_record = merge_plan_records(
            resolved_map.get(plan_name, {}),
            {},
            fallback_plan_name=plan_name,
            fallback_underwriter=resolved_map.get(plan_name, {}).get("0.1. Underwriter", ""),
        )

        for field_key, bucket in (field_map or {}).items():
            all_candidates = list(bucket.get("candidate_entries", []))
            if not all_candidates:
                continue

            current_value = plan_record.get(field_key, bucket.get("current_value", ""))
            selected_candidates = _select_best_candidate_entries(all_candidates, max_candidates=10)
            if not _should_call_field_aggregation_llm(field_key, current_value, selected_candidates):
                continue

            category = CATEGORY_LOOKUP.get(field_key, {})
            pending_decisions.setdefault(plan_name, {})[field_key] = {
                "all_candidates": all_candidates,
                "selected_candidates": selected_candidates,
                "current_value": current_value,
                "field_name": category.get("name", field_key),
                "field_behavior": _field_behavior(field_key),
            }
            decision_requests.append(
                {
                    "plan_name": plan_name,
                    "field_key": field_key,
                    "field_name": category.get("name", field_key),
                    "field_behavior": _field_behavior(field_key),
                    "current_value": str(current_value or "").strip(),
                    "candidate_rows": selected_candidates,
                }
            )

    decision_results = _decide_field_values_with_llm(decision_requests) if decision_requests else {}

    for plan_name, field_map in (candidate_state or {}).items():
        plan_record = merge_plan_records(
            resolved_map.get(plan_name, {}),
            {},
            fallback_plan_name=plan_name,
            fallback_underwriter=resolved_map.get(plan_name, {}).get("0.1. Underwriter", ""),
        )
        plan_audit: dict[str, Any] = {}

        for field_key, bucket in (field_map or {}).items():
            pending = pending_decisions.get(plan_name, {}).get(field_key)
            if not pending:
                continue

            all_candidates = pending["all_candidates"]
            selected_candidates = pending["selected_candidates"]
            decision = decision_results.get(plan_name, {}).get(field_key, _normalize_field_decision_result({}))
            decision_value = str(decision.get("value", "")).strip()
            if decision_value:
                final_value = decision_value
                if decision.get("decision_mode") == "single" and len(selected_candidates) == 1:
                    final_value = (
                        _format_multi_row_candidate(
                            field_key,
                            selected_candidates[0].get("label", ""),
                            decision_value,
                        )
                        if field_key in MULTI_ROW_FIELDS else decision_value
                    )
                if field_key == "3. Area of Cover":
                    final_value = _prefer_field_value(
                        field_key,
                        plan_record.get(field_key, ""),
                        _compose_area_of_cover_value(selected_candidates, fallback_value=final_value),
                    )
                plan_record[field_key] = final_value

            plan_audit[field_key] = {
                "candidate_count": len(all_candidates),
                "decision_mode": decision.get("decision_mode", "blank"),
                "final_value": str(plan_record.get(field_key, "") or "").strip(),
                "used_row_indexes": decision.get("used_row_indexes", []),
            }

        resolved_map[plan_name] = clean_raw_fields(plan_record)
        if plan_audit:
            resolution_audit[plan_name] = plan_audit

    return resolved_map, resolution_audit


def _apply_semantic_table_backfill(
    page_content: str,
    plan_names: list[str],
    plan_updates: dict[str, dict],
    fallback_underwriter: str = "",
    candidate_state: dict[str, dict[str, dict[str, Any]]] | None = None,
    page_number: int | None = None,
) -> tuple[dict[str, dict], dict[str, dict[str, int]], dict[str, dict[str, list[dict[str, Any]]]]]:
    table_rows = _parse_markdown_table_rows(page_content)
    if not table_rows or not plan_names:
        return plan_updates, {}, {}

    single_plan_mode = len(plan_names) == 1
    updated_plans: dict[str, dict] = {}
    page_candidate_summary: dict[str, dict[str, int]] = {}
    page_candidate_details: dict[str, dict[str, list[dict[str, Any]]]] = {}
    for plan_idx, plan_name in enumerate(plan_names):
        plan_record = merge_plan_records(
            {},
            plan_updates.get(plan_name, {}),
            fallback_plan_name=plan_name,
            fallback_underwriter=fallback_underwriter,
        )

        for field_key, current_value in list(plan_record.items()):
            if str(current_value or "").strip() and _field_behavior(field_key) == "strict_single":
                continue
            if field_key not in CATEGORY_LOOKUP:
                continue

            candidate_entries = _collect_field_candidate_entries(
                table_rows,
                field_key,
                plan_idx,
                single_plan_mode=single_plan_mode,
                page_number=page_number,
            )

            if not candidate_entries:
                continue

            page_candidate_details.setdefault(plan_name, {})[field_key] = [dict(entry) for entry in candidate_entries]
            page_candidate_summary.setdefault(plan_name, {})[field_key] = len(candidate_entries)

            if field_key == "3. Area of Cover":
                reconstructed_value = _compose_area_of_cover_value(candidate_entries, fallback_value=current_value)
                plan_record[field_key] = _prefer_field_value(field_key, plan_record.get(field_key, ""), reconstructed_value)
                if candidate_state is not None:
                    _register_semantic_candidate_entries(
                        candidate_state,
                        plan_name=plan_name,
                        field_key=field_key,
                        current_value=plan_record.get(field_key, ""),
                        candidate_entries=candidate_entries,
                    )
                continue

            if _should_call_field_aggregation_llm(field_key, current_value, candidate_entries):
                selected_candidates = _select_best_candidate_entries(candidate_entries)
                if candidate_state is not None:
                    _register_semantic_candidate_entries(
                        candidate_state,
                        plan_name=plan_name,
                        field_key=field_key,
                        current_value=current_value,
                        candidate_entries=selected_candidates,
                    )
                candidate_entries = selected_candidates
                if _field_behavior(field_key) == "strict_single":
                    continue

            best_candidate = max(
                candidate_entries,
                key=lambda item: (
                    score_field_information(str(item.get("candidate_value", "") or "")),
                    len(str(item.get("candidate_value", "") or "")),
                    -(int(item.get("page_number") or 0)),
                    -(int(item.get("row_idx") or 0)),
                ),
            )
            best_value = best_candidate["candidate_value"]

            plan_record[field_key] = merge_field_value(plan_record.get(field_key, ""), best_value, field_key=field_key)

        updated_plans[plan_name] = clean_raw_fields(plan_record)

    return updated_plans, page_candidate_summary, page_candidate_details


def _format_semantic_candidate_rows(
    page_content: str,
    category_keys: list[str] | None = None,
    plan_names: list[str] | None = None,
    target_plan_name: str | None = None,
    max_candidates: int = 2,
    min_score: float = 0.35,
) -> str:
    rows = _parse_markdown_table_rows(page_content)
    if not rows:
        return "- none"

    category_keys = category_keys or list(CATEGORY_LOOKUP.keys())
    plan_names = plan_names or []
    snippets: list[str] = []

    for row in rows:
        row_label = str(row.get("label", "")).strip()
        row_section = str(row.get("parent_label", "") or row.get("section", "")).strip()
        row_values = row.get("values", [])

        matches = []
        for field_key in category_keys:
            if _row_matches_field(field_key, row_label, row_section):
                matches.append(field_key)

        if not matches:
            continue

        field_labels = ", ".join(matches[:max_candidates])
        value_preview = []
        for idx, value in enumerate(row_values):
            if idx >= len(plan_names):
                break
            plan_name = plan_names[idx]
            if target_plan_name and plan_name != target_plan_name:
                continue
            if str(value).strip():
                value_preview.append(f"{plan_name}: {str(value).strip()}")
        if not value_preview:
            continue
        snippets.append(f"- {row_label} -> {field_labels} | {'; '.join(value_preview)}")

    return "\n".join(snippets[:25]) if snippets else "- none"


# ── RAW VALUE CLEANUP ─────────────────────────────────────────────────────────
def clean_raw_fields(raw_fields: dict) -> dict:
    """Apply light field-specific cleanup after extraction."""
    cleaned_fields = {}
    for field_key, raw_value in (raw_fields or {}).items():
        value = raw_value.strip() if isinstance(raw_value, str) else raw_value

        if isinstance(value, str) and field_key == "2. Annual Limit per person":
            filtered_parts = []
            for part in _split_field_parts(value):
                part = _strip_aggregate_limit_label_prefix(part)
                lowered_part = part.lower()
                if any(term in lowered_part for term in [
                    "worldwide", "area of cover", "country of residence", "principal country",
                    "treatment abroad", "emergency medical treatment abroad", "non emergency medical treatment abroad",
                    "sports", "sport", "recreation", "recreational", "athletic", "adventure", "extreme", "hazardous",
                    "airfare", "air fare", "economy class", "return economy", "return airfare", "pre approval", "pre-approval",
                ]):
                    continue
                if any(term in lowered_part for term in [
                    "annual", "overall", "aggregate", "sum insured", "policy limit", "annual limit", "per person", "maximum limit",
                    "indemnity limit", "indemnity",
                ]) or _looks_like_aggregate_limit_value(part):
                    filtered_parts.append(part)
            if filtered_parts:
                value = _merge_distinct_field_parts("", "\n".join(filtered_parts))
            elif _looks_like_aggregate_limit_value(value):
                value = _strip_aggregate_limit_label_prefix(value)
            else:
                value = ""

        if isinstance(value, str) and field_key == "5. Network":
            value = re.sub(
                r"\s*\([^)]*(?:copay|co\s*pay|coinsurance|deductible|waiting)[^)]*\)",
                "",
                value,
                flags=re.IGNORECASE,
            ).strip()
            if "," in value and any(term in value.lower() for term in ["copay", "co pay", "coinsurance", "deductible", "waiting"]):
                value = value.split(",", 1)[0].strip()

        if isinstance(value, str) and field_key == "4. TPA":
            value = re.sub(r"^(?:tpa|claims administrator|claims processor|administrator)\s*[:|-]\s*", "", value, flags=re.IGNORECASE).strip()

        if isinstance(value, str) and field_key == "3. Area of Cover":
            filtered_parts = []
            for part in _split_field_parts(value):
                lowered_part = part.lower()
                if any(term in lowered_part for term in ["sports", "sport", "recreation", "recreational", "athletic", "adventure", "extreme", "hazardous"]):
                    continue
                filtered_parts.append(part)
            value = _merge_distinct_field_parts("", "\n".join(filtered_parts)) if filtered_parts else ""

        cleaned_fields[field_key] = value
    return cleaned_fields


# ── LAYOUT DETECTION ─────────────────────────────────────────────────────────
def detect_layout(structured_md: str) -> str:
    system_prompt = """You classify insurance table-of-benefits layouts.

<output_shape>
- Return JSON only.
- Schema: {"layout": "columnar" | "sequential" | "hybrid_inline_plan"}
</output_shape>

<classification_rules>
- columnar: multiple plans appear side by side in the same table row set.
- sequential: each plan/category appears in its own section or page block, one after another.
- hybrid_inline_plan: one plan appears in a true table column or section, while another plan/category on the same page is flattened into inline text or a repeated value block instead of a matching second table column.
- Choose the simplest correct label from the provided text only.
- Do not explain your answer.
</classification_rules>"""
    user_prompt = f"""Determine whether this insurance quote uses a columnar or sequential layout.

DOCUMENT SNIPPET:
{structured_md[:4000]}

Return JSON only."""

    content, usage = llm_chat(system_prompt, user_prompt, json_mode=True, log_label="layout_detection", task_class="layout_detection")
    tracker.record(_current_pdf, usage)
    result = safe_json_loads(content, "layout detection")
    layout = result.get("layout", "columnar").strip().lower()
    if layout not in ("columnar", "sequential", "hybrid_inline_plan"):
        layout = "columnar"
    print(f"  📐 Detected layout: {layout.upper()}")
    return layout


def detect_plan_regions(structured_md: str, layout: str, detected_plan_names: list[str] | None = None) -> list[dict[str, Any]]:
    page_blocks = _split_markdown_by_page(structured_md)
    total_pages = max(page_blocks.keys()) if page_blocks else 1
    page_index_lines = []
    for pn in sorted(page_blocks.keys()):
        snippet = page_blocks[pn][:350].replace("\n", " ").strip()
        page_index_lines.append(f"PAGE {pn}: {snippet}")
    page_index_text = "\n".join(page_index_lines)

    known_plan_hint = ""
    detected_plan_names = [str(name).strip() for name in (detected_plan_names or []) if str(name).strip()]
    if detected_plan_names:
        known_plan_hint = "\nKnown plan/category names already detected:\n" + "\n".join(f"- {name}" for name in detected_plan_names)

    system_prompt = """You identify true insurance plan/category regions in a TOB document.

<task>
- Find every true plan/category represented in the document.
- Return where each plan is represented and how it is structured.
</task>

<output_shape>
- Return JSON only.
- Schema:
  {
    "plans": [
      {
        "plan_name": string,
        "source_type": "table_column" | "section" | "inline_block",
        "start_page": number,
        "end_page": number,
        "anchor_text": string
      }
    ]
  }
</output_shape>

<rules>
- A true plan/category is a primary insured category or product, not an add-on, rider, insurer label, or subsection.
- Use `table_column` when the plan is clearly a table column header.
- Use `section` when the plan occupies a separate page range or section.
- Use `inline_block` when the plan appears as a repeated text block or inline tail rather than a matching table column.
- Return all true plans you can support from the document.
- Preserve exact plan/category names when supported.
- Do not invent plans.
</rules>"""

    user_prompt = f"""Identify all true insurance plans/categories in this document and how each one is represented.

LAYOUT LABEL:
{layout}
{known_plan_hint}

PAGE INDEX:
{page_index_text}

DOCUMENT SNIPPET:
{structured_md[:7000]}

Return JSON only."""

    try:
        content, usage = llm_chat(
            system_prompt,
            user_prompt,
            json_mode=True,
            log_label="plan_region_detection",
            task_class="layout_detection",
        )
        tracker.record(_current_pdf, usage)
        result = safe_json_loads(content, "plan region detection")
    except Exception as exc:
        print(f"  ⚠️  Plan region detection failed: {exc}")
        result = {}

    raw_plans = result.get("plans", []) if isinstance(result, dict) else []
    regions: list[dict[str, Any]] = []
    seen_plan_names: set[str] = set()
    for item in raw_plans:
        if not isinstance(item, dict):
            continue
        plan_name = str(item.get("plan_name", "") or "").strip()
        source_type = str(item.get("source_type", "") or "").strip().lower()
        if not plan_name:
            continue
        if source_type not in {"table_column", "section", "inline_block"}:
            source_type = "section" if layout == "sequential" else "table_column"
        start_page = item.get("start_page", 1)
        end_page = item.get("end_page", total_pages)
        try:
            start_page = int(start_page)
            end_page = int(end_page)
        except Exception:
            start_page = 1
            end_page = total_pages
        region = {
            "plan_name": plan_name,
            "source_type": source_type,
            "start_page": max(1, start_page),
            "end_page": max(max(1, start_page), end_page),
            "anchor_text": str(item.get("anchor_text", "") or "").strip(),
        }
        normalized_name = _normalize_match_text(plan_name)
        if normalized_name in seen_plan_names:
            continue
        seen_plan_names.add(normalized_name)
        regions.append(region)

    for plan_name in detected_plan_names:
        normalized_name = _normalize_match_text(plan_name)
        if normalized_name in seen_plan_names:
            continue
        regions.append(
            {
                "plan_name": plan_name,
                "source_type": "section" if layout == "sequential" else "table_column",
                "start_page": 1,
                "end_page": total_pages,
                "anchor_text": plan_name,
            }
        )

    if regions:
        print(f"  🧭 Plan regions: {[region['plan_name'] for region in regions]}")
    return regions


def _extract_inline_plan_text(page_content: str, anchor_terms: list[str], stop_terms: list[str]) -> str:
    lines = [line.rstrip() for line in str(page_content or "").splitlines()]
    capture_start = -1

    for idx, line in enumerate(lines):
        normalized_line = _normalize_match_text(line)
        if any(term and term in normalized_line for term in anchor_terms):
            capture_start = idx
            break

    if capture_start < 0:
        return ""

    captured_lines: list[str] = []
    for idx in range(capture_start, len(lines)):
        line = lines[idx].strip()
        normalized_line = _normalize_match_text(line)
        if idx > capture_start and normalized_line.startswith("page "):
            break
        if idx > capture_start and any(term and term in normalized_line for term in stop_terms):
            break
        if line:
            captured_lines.append(line)
    return "\n".join(captured_lines).strip()


def _format_page_plan_region_hints(page_content: str, page_number: int, plan_regions: list[dict[str, Any]] | None) -> str:
    plan_regions = plan_regions or []
    page_plan_hints: list[str] = []
    plan_names = [str(region.get("plan_name", "") or "").strip() for region in plan_regions if str(region.get("plan_name", "") or "").strip()]

    for region in plan_regions:
        if region.get("source_type") != "inline_block":
            continue
        if page_number < int(region.get("start_page", 1) or 1) or page_number > int(region.get("end_page", page_number) or page_number):
            continue

        plan_name = str(region.get("plan_name", "") or "").strip()
        anchor_text = str(region.get("anchor_text", "") or "").strip()
        anchor_terms = [
            _normalize_match_text(anchor_text),
            _normalize_match_text(plan_name),
        ]
        stop_terms = [
            _normalize_match_text(other_name)
            for other_name in plan_names
            if _normalize_match_text(other_name) and _normalize_match_text(other_name) != _normalize_match_text(plan_name)
        ]
        snippet = _extract_inline_plan_text(page_content, anchor_terms, stop_terms)
        if not snippet:
            continue
        page_plan_hints.append(f"Plan: {plan_name}\nSource: inline_block\nSnippet:\n{snippet}")

    return "\n\n".join(page_plan_hints) if page_plan_hints else "- none"


TOB_PAGE_SIGNAL_KEYWORDS = {
    "table of benefits",
    "schedule of benefits",
    "benefits",
    "coverage",
    "copay",
    "co pay",
    "coinsurance",
    "deductible",
    "network",
    "in patient",
    "out patient",
    "maternity",
    "dental",
    "optical",
    "pharmacy",
    "medication",
}


def _assess_tob_page_presence(
    page_content: str,
    page_number: int,
    plan_regions: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    usable_rows = [
        row for row in _parse_markdown_table_rows(page_content)
        if str(row.get("label", "") or "").strip()
    ]
    matched_labels: list[str] = []
    matched_field_keys: set[str] = set()

    for row in usable_rows:
        label = str(row.get("label", "") or "").strip()
        row_section = str(row.get("parent_label", "") or row.get("section", ""))
        if not label:
            continue

        row_matches = []
        for field_key in CATEGORY_LOOKUP.keys():
            if _row_matches_field(field_key, label, row_section):
                row_matches.append(field_key)

        if row_matches:
            matched_labels.append(label)
            matched_field_keys.update(row_matches)

    normalized_page = _normalize_match_text(page_content)
    keyword_hits = sorted(
        keyword for keyword in TOB_PAGE_SIGNAL_KEYWORDS
        if keyword in normalized_page
    )
    page_region_count = sum(
        1
        for region in (plan_regions or [])
        if int(region.get("start_page") or 0) <= page_number <= int(region.get("end_page") or 0)
    )

    has_tob_signal = bool(
        len(usable_rows) >= 2
        or matched_labels
        or (page_region_count > 0 and (len(keyword_hits) >= 2 or len(usable_rows) >= 1))
        or ("table of benefits" in normalized_page)
        or ("schedule of benefits" in normalized_page)
    )

    return {
        "should_process": has_tob_signal,
        "usable_row_count": len(usable_rows),
        "matched_row_count": len(matched_labels),
        "matched_field_count": len(matched_field_keys),
        "matched_labels": matched_labels[:8],
        "matched_field_keys": sorted(matched_field_keys)[:12],
        "keyword_hits": keyword_hits[:8],
        "page_region_count": page_region_count,
        "reason": "tob_signal_detected" if has_tob_signal else "no_tob_signal_detected",
    }


# ── STEP 2: EXTRACTION PASS ───────────────────────────────────────────────────
def run_extraction_pass(structured_md: str, fields_to_extract: list[tuple] = None) -> tuple[dict, list[dict]]:
    if fields_to_extract is None:
        fields_to_extract = CATEGORIES

    categories_text = "\n".join(
        f'{num}. {name}\n   Find: {hint}'
        for num, name, hint in fields_to_extract
    )

    cover_text = structured_md[:4000]
    detected_underwriter = _detect_underwriter(cover_text)
    layout = detect_layout(structured_md)
    plan_regions = detect_plan_regions(structured_md, layout, detected_plan_names=_detected_plan_names)

    page_blocks = _split_markdown_by_page(structured_md)
    if not page_blocks:
        page_blocks = {1: structured_md}

    extracted_plans, page_audit = _extract_document_page_by_page(
        page_blocks=page_blocks,
        categories_text=categories_text,
        layout=layout,
        underwriter_hint=detected_underwriter,
        plan_regions=plan_regions,
    )
    return extracted_plans, page_audit


def _extract_document_page_by_page(page_blocks: dict, categories_text: str, layout: str, underwriter_hint: str = "", plan_regions: list[dict[str, Any]] | None = None) -> tuple[dict, list[dict]]:
    global _detected_plan_names

    page_audit = []
    merged_state: dict[str, dict] = {}
    region_plan_names = [str(region.get("plan_name", "") or "").strip() for region in (plan_regions or []) if str(region.get("plan_name", "") or "").strip()]
    known_plan_names = [plan for plan in dict.fromkeys([*region_plan_names, *[plan for plan in _detected_plan_names if str(plan).strip()]]) if str(plan).strip()]
    detected_underwriter = underwriter_hint or ""
    generic_plan_names = {"plan", "category", "n/a", "none", "", "not mentioned", "unknown", "plan name"}

    for page_number in sorted(page_blocks.keys()):
        page_content = page_blocks[page_number]
        tob_presence = _assess_tob_page_presence(
            page_content,
            page_number,
            plan_regions=plan_regions,
        )
        page_attribute_candidates = _format_page_attribute_candidates(page_content)
        existing_state = normalize_plan_map(merged_state, detected_underwriter)
        known_plan_hint = "\n".join(f"- {name}" for name in known_plan_names) if known_plan_names else "- none detected yet"
        page_plan_region_hints = _format_page_plan_region_hints(page_content, page_number, plan_regions)

        if not tob_presence["should_process"]:
            print(f"  ↷ Skipping page {page_number}: no TOB signal detected")
            page_audit.append(
                {
                    "page": page_number,
                    "layout": layout,
                    "skipped": True,
                    "skip_reason": "no_tob_signal_detected",
                    "tob_presence": tob_presence,
                    "running_state_before": existing_state,
                    "running_state_after": existing_state,
                }
            )
            continue

        skeleton_plan_name = known_plan_names[0] if known_plan_names else "<plan_name>"
        output_skeleton = {
            "plans": {
                skeleton_plan_name: default_plan_record(skeleton_plan_name, detected_underwriter)
            }
        }

        system_prompt = f"""You extract structured insurance quote data from one page at a time.

<task>
- Review exactly one page of a table-of-benefits document.
- Maintain a running plan state across pages.
    - Start with the running state provided to you.
    - Update that running state only when the current page gives direct evidence.
    - Return the FULL updated plan state after processing the current page.
</task>

<output_shape>
- Return JSON only.
- Schema:
{json.dumps(output_skeleton, indent=2, ensure_ascii=False)}
- For each plan, return EVERY tracked category key exactly as listed below.
- Use only the provided category keys. No extra keys.
- If a tracked field is still unknown after reviewing the current page, keep the prior value from running state; if there is no prior value, return an empty string for that field.
</output_shape>

<tracked_categories>
{categories_text}
</tracked_categories>

<grounding_rules>
- Work from the current page text only.
- For page 1, initialize plan values directly from page 1 evidence.
- For page 2 onward, keep the prior extracted values unless the current page adds, corrects, or makes them more specific.
- If a field is not mentioned on the current page, preserve the prior value instead of blanking it.
- Read the page data provided in the user prompt carefully before filling any category.
- Map semantically equivalent row labels to the closest tracked field even when the wording differs.
- Check the tracked categories list and decide whether any value from the page data should update a category based on semantic meaning, not exact wording.
- If a page attribute label clearly corresponds to a tracked category by meaning, copy that row's plan value into the tracked category.
- For `3. Area of Cover`, if the page contains principal country, emergency abroad, and non-emergency abroad rows, keep all of them together in the same field using `; ` separators.
- Some benefits may contain multiple directly related sub-values under one benefit header. If the page clearly shows grouped rows that all belong to the same tracked benefit, keep them together in the same field as multi-line text.
- Do not combine nearby values unless they are clearly part of the same benefit context.
- Aggregate fields such as annual limit per person must remain single overall values and must not absorb maternity, dental, optical, consultation, diagnostic, or other benefit-specific sublimits.
- If the user prompt provides semantic candidate rows for a category, treat them as high-priority retrieval hints and verify them against the page content before deciding the final value.
- Use service context to distinguish inpatient vs outpatient evidence; do not populate an outpatient deductible or copay field unless the page explicitly states an outpatient consultation deductible, copay, or coinsurance.
- Do not guess, infer, summarize, normalize, or paraphrase values.
- Preserve exact wording, limits, copays, coinsurance, waiting periods, exclusions, and negative values.
- If the page provides a clearer or more complete value than the running state, replace that field with the improved value.
- If the page does not improve a field, keep the running-state value unchanged.
- Riders and add-ons belong to their parent plan; do not create separate rider-only plans.
- If the page introduces a new true plan, include it.
- Respect the detected document layout: {layout}.
</grounding_rules>

<quality_bar>
- Minimize hallucinations.
- Prefer blank over uncertain.
- Re-scan the page before finalizing to catch missed fields.
- Before finalizing, explicitly check whether any page data in the user prompt semantically matches currently blank categories such as physician, diagnostics, medication, ambulance, nursing, transplant, dialysis, or psychiatric benefits.
</quality_bar>"""

        user_prompt = f"""Update the running insurance extraction with this page.

<page_number>
{page_number}
</page_number>

<known_plan_names>
{known_plan_hint}
</known_plan_names>

<page_plan_region_hints>
These are page-level plan snippets detected from non-tabular or mixed-layout content. Use them when a plan appears as inline text rather than a table column.
{page_plan_region_hints}
</page_plan_region_hints>

<page_attribute_candidates>
These are the row labels explicitly present on this page. Use them as semantic candidates when deciding whether a tracked category should be updated.
{page_attribute_candidates}
</page_attribute_candidates>

<running_state_before>
{json.dumps(existing_state, indent=2, ensure_ascii=False)}
</running_state_before>

<page_content>
{page_content}
</page_content>

Return JSON only with the FULL updated state under the top-level key \"plans\"."""

        content, usage = llm_chat(
            system_prompt,
            user_prompt,
            json_mode=True,
            log_label=f"page_{page_number}_update",
            task_class="extraction",
            log_context={
                "page_number": page_number,
                "known_plan_names": known_plan_names,
                "page_plan_region_hints": page_plan_region_hints,
                "running_state_before": existing_state,
                "page_content": page_content,
                "layout": layout,
            },
        )
        tracker.record(_current_pdf, usage)
        page_result = safe_json_loads(content, f"page {page_number} extraction")

        raw_plan_updates = page_result.get("plans", page_result)
        if any(key in raw_plan_updates for key in ["0.1. Underwriter", "0.2. Plan Name"]):
            fallback_plan_name = known_plan_names[0] if known_plan_names else f"Plan {page_number}"
            raw_plan_updates = {fallback_plan_name: raw_plan_updates}

        normalized_updates = {}
        for raw_plan_name, fields in (raw_plan_updates or {}).items():
            if not isinstance(fields, dict):
                continue
            effective_plan_name = (fields.get("0.2. Plan Name") or raw_plan_name or "").strip()
            if effective_plan_name.lower() in generic_plan_names:
                effective_plan_name = raw_plan_name if raw_plan_name.lower() not in generic_plan_names else ""
            if not effective_plan_name:
                if len(known_plan_names) == 1:
                    effective_plan_name = known_plan_names[0]
                else:
                    effective_plan_name = f"Plan {len(existing_state) + len(normalized_updates) + 1}"

            if effective_plan_name not in known_plan_names:
                known_plan_names.append(effective_plan_name)

            normalized_updates[effective_plan_name] = merge_plan_records(
                {},
                fields,
                fallback_plan_name=effective_plan_name,
                fallback_underwriter=detected_underwriter,
            )

        normalized_updates, known_plan_names = _collapse_single_tob_alias_plans(
            normalized_updates,
            known_plan_names=known_plan_names,
            fallback_underwriter=detected_underwriter,
        )
        page_candidate_summary: dict[str, dict[str, int]] = {}
        page_candidate_details: dict[str, dict[str, list[dict[str, Any]]]] = {}

        all_plan_names = list(dict.fromkeys([*existing_state.keys(), *normalized_updates.keys(), *known_plan_names]))
        for plan_name in all_plan_names:
            fields = normalized_updates.get(plan_name, existing_state.get(plan_name, {}))
            merged_state[plan_name] = merge_plan_records(
                merged_state.get(plan_name, {}),
                fields,
                fallback_plan_name=plan_name,
                fallback_underwriter=detected_underwriter,
            )
            detected_underwriter = merged_state[plan_name].get("0.1. Underwriter", detected_underwriter)

        page_audit.append(
            {
                "page": page_number,
                "layout": layout,
                "tob_presence": tob_presence,
                "running_state_before": existing_state,
                "page_updates": normalized_updates,
                "semantic_candidate_summary": page_candidate_summary,
                "semantic_candidate_details": page_candidate_details,
                "running_state_after": normalize_plan_map(merged_state, detected_underwriter),
            }
        )

    normalized_state = normalize_plan_map(merged_state, detected_underwriter)
    normalized_state, known_plan_names = _collapse_single_tob_alias_plans(
        normalized_state,
        known_plan_names=known_plan_names,
        fallback_underwriter=detected_underwriter,
    )
    if not normalized_state:
        fallback_plan_name = known_plan_names[0] if known_plan_names else "Plan 1"
        normalized_state = {fallback_plan_name: default_plan_record(fallback_plan_name, detected_underwriter)}
    return normalized_state, page_audit

# ─────────────────────────────────────────────────────────────────────────────
# FIX 1 + FIX 2 applied here in the SYSTEM_PROMPT
# ─────────────────────────────────────────────────────────────────────────────
def _extract_columnar(structured_md: str, categories_text: str) -> dict:
        
    """
    Handles side-by-side column layout.
    FIX 1: Riders/add-ons must NOT become separate plan columns.
    FIX 2: Values must never include the row-label prefix.
    """

        # Use plan names pre-detected by _detect_underwriter if available
    global _detected_plan_names
    plan_name_hint = ""
    if _detected_plan_names:
        plan_name_hint = (
            "KNOWN PLAN NAMES — use these EXACTLY for '0.2. Plan Name' and as JSON keys:\n"
            + "\n".join(f"  - {p}" for p in _detected_plan_names)
            + "\n"
        )
    SYSTEM_PROMPT = f"""You are an expert insurance document analyst for health insurance Table of Benefits (TOB).
This document uses a COLUMNAR layout: multiple plans appear as SIDE-BY-SIDE COLUMNS in the same table.

EXTRACTION RULES:
1. Identify every TRUE PLAN column. A TRUE PLAN column is a primary insurance product (e.g., "Individual/Family Medical Plan", "Category A", "Gold Plan").
2. *** CRITICAL — DO NOT treat the following as separate plan columns: ***
   - Riders or add-ons (e.g., "TRAVEL MEDICAL RIDER", "Travel Rider", "Dental Rider")
   - Sub-sections or optional benefit blocks (e.g., "Artificial body parts", "Below limit options", "Maternity add-on")
   - Section headers within a benefit table that appear as a row spanning the table
   If rider/add-on benefits apply to a plan, merge those values INTO that plan's data for the relevant benefit field.
3. If a rider/add-on column provides a value for a benefit that the main plan already covers,
   use the best direct value for the parent plan field rather than building a combined multi-part value.
4. Return a VALID JSON object. Do not include any text outside the JSON.
5. Use semantic matching to find benefit rows.
6. If no value exists, return "". Do not write "Not mentioned".
7. Capture the EXACT wording. Do not summarize or paraphrase.
8. Capture the ENTIRE description — do not stop at first punctuation.
9. If a row says "Not Covered", output "Not Covered". Preserve negative values.
10.Extract each benefit from its own dedicated row. Do not copy values from unrelated sections.
11.For all values, include waiting period,coinsurance, co-payment and all conditions required for each benefit.
12.Choose the single best-supported value for each tracked field from the document.
13.For `3. Area of Cover`, extract only geographic scope or treatment territory. Do not include sports, recreation, activities, or exclusions unrelated to geography.
14.For `5. Network`, extract only the network name or provider panel name. Do not include copay, coinsurance, deductibles, waiting periods, or other conditions.
13.*** PRESERVE EXACT VALUES — DO NOT MODIFY: ***
   Keep ALL values EXACTLY as they appear in the quote/document.
   - Do NOT convert "Not Covered" to "Covered" or vice versa
   - Do NOT change numbers, limits, or percentages
   - Do NOT rephrase or interpret values — copy them verbatim
   - If quote says "750", output "750" (not "BHD 750" unless BHD is in the original)
   - If quote says "Not Covered", output "Not Covered" exactly
   - If quote says "20%", output "20%" exactly
   - Preserve the original text, spelling, and formatting from the document

FIELDS TO EXTRACT:
{categories_text}
"""

    USER_PROMPT = f"""Extract all fields for every TRUE PLAN from this COLUMNAR insurance TOB.
    
{plan_name_hint}Each TRUE PLAN is a primary product column. Use the plan's column header as the JSON key.
Do NOT create separate entries for riders, add-ons, or sub-sections — merge them into the parent plan.

DOCUMENT:
{structured_md}

Return ONLY a JSON object:
{{
  "Individual/Family Medical Plan": {{
    "0.1. Underwriter": "Arabia Insurance",
    "0.2. Plan Name": "Individual/Family Medical Plan",
    "2. Annual Limit per person": "BHD 10,000",
    ...
  }}
}}"""

    content, usage = llm_chat(SYSTEM_PROMPT, USER_PROMPT, json_mode=True, log_label="columnar_extraction", task_class="extraction")
    tracker.record(_current_pdf, usage)
    raw_result = safe_json_loads(content, "_extract_columnar")

    # FIX 2 post-processing: strip label prefixes from all values
    cleaned = {}
    for plan_key, fields in raw_result.items():
        if isinstance(fields, dict):
            cleaned[plan_key] = clean_raw_fields(fields)
        else:
            cleaned[plan_key] = fields

    return cleaned


def _extract_sequential(structured_md: str, categories_text: str) -> dict:
    """
    Handles sequential/vertical layout where each plan is a separate section.
    FIX 1 + FIX 2 applied via prompts and post-processing.
    """
    page_blocks_preview = _split_markdown_by_page(structured_md)
    total_pages = max(page_blocks_preview.keys()) if page_blocks_preview else 1

    page_index_lines = []
    for pn in sorted(page_blocks_preview.keys()):
        snippet = page_blocks_preview[pn][:300].replace("\n", " ").strip()
        page_index_lines.append(f"PAGE {pn}: {snippet}")
    page_index_text = "\n".join(page_index_lines)

    # Seed from pre-detected plan names if available
    global _detected_plan_names
    known_plans_hint = ""
    if _detected_plan_names:
        known_plans_hint = (
            f"\nKNOWN PLAN NAMES already detected: {_detected_plan_names}. "
            f"Use these as the plan_name values in your response.\n"
        )

    SECTION_DETECT_PROMPT = f'''You are analyzing a sequential insurance TOB document.
Each insurance plan/category occupies its OWN set of pages (not side-by-side columns).
Your job: find EVERY distinct TRUE PLAN/category and its page range.
{known_plans_hint}
IMPORTANT: A TRUE PLAN is a primary insurance product (e.g., "CAT A", "Option 1", "Gold Plan").
Do NOT list riders, add-ons, or sub-sections (e.g., "Travel Rider", "Dental Add-on") as separate plans.

Common patterns to look for:
- "Category CAT A" / "Category CAT B" headers at the top of a new section
- "OPTION 1" / "OPTION 2" section headers
- Repeated benefit tables, each starting with a different plan name header
- A single-plan document where the whole document is one plan

CRITICAL: If the document has only 1-2 pages and appears to be a single plan,
return it as one section spanning all pages. Do NOT return an empty sections list.

Return ONLY a JSON object:
{{"sections": [
  {{"plan_name": "Category A", "start_page": 1, "end_page": 2}},
  {{"plan_name": "Category B", "start_page": 3, "end_page": 4}}
]}}

IMPORTANT: Return ALL categories found. NEVER return an empty sections list.'''

    section_user = (
        f"This document has {total_pages} pages. "
        f"Here is a summary of each page's content:\n\n{page_index_text}\n\n"
        f"Full document text:\n\n{structured_md[:4000]}\n\n"
        f"Identify ALL TRUE PLAN/category sections and their page ranges. "
        f"Exclude riders and add-on sections from the list."
    )
    content, usage = llm_chat(SECTION_DETECT_PROMPT, section_user, json_mode=True, log_label="section_detection", task_class="extraction")
    tracker.record(_current_pdf, usage)
    sections_data = safe_json_loads(content, "section detection")
    sections = sections_data.get("sections", [])
    section_summary = [(s['plan_name'], f"pp.{s['start_page']}-{s['end_page']}") for s in sections]
 #   print(f"  📋 Sequential sections detected: {section_summary}")

    if not sections:
        print("  ⚠️  LLM detected no sections — attempting regex fallback")
        import re
        sections = []
        for pn in sorted(page_blocks_preview.keys()):
            page_text = page_blocks_preview[pn]
            match = re.search(r'Category\s+(CAT\s*[A-Z]|Option\s*\d+|Plan\s*\w+)', page_text, re.IGNORECASE)
            if match:
                plan_id = match.group(1).strip()
                if not sections or sections[-1]["plan_name"] != plan_id:
                    if sections:
                        sections[-1]["end_page"] = pn - 1
                    sections.append({"plan_name": plan_id, "start_page": pn, "end_page": total_pages})
        if sections:
            print(f"  ✓ Fallback found: {[s['plan_name'] for s in sections]}")
        else:
            print("  ⚠️  No sections found via fallback — treating as single plan")
            plan_label = _detected_plan_names[0] if _detected_plan_names else "Plan"
            sections = [{"plan_name": plan_label, "start_page": 1, "end_page": total_pages}]

    all_plans = {}
    page_blocks = _split_markdown_by_page(structured_md)

    # Use more text for underwriter detection, and don't overwrite _detected_plan_names
    cover_text = structured_md
    underwriter = _detect_underwriter(cover_text)
    # Reset plan names after sequential's own detection — sequential manages its
    # own plan names via sections, we don't want to pollute the global for other PDFs
    _detected_plan_names = []

    for section in sections:
        plan_name = section.get("plan_name", "Unknown Plan")
        start_p   = section.get("start_page", 1)
        end_p     = section.get("end_page", 999)

        section_md_parts = []
        for page_num in range(start_p, end_p + 1):
            if page_num in page_blocks:
                section_md_parts.append(f"\n\n---\n## PAGE {page_num}\n---\n")
                section_md_parts.append(page_blocks[page_num])
        section_md = "\n".join(section_md_parts)

        if not section_md.strip():
            print(f"  ⚠️  No content found for {plan_name} (pages {start_p}-{end_p})")
            continue

        print(f"  → Extracting {plan_name} (pages {start_p}-{end_p}, {len(section_md):,} chars)")

        EXTRACT_SYSTEM = f"""You are extracting benefits from a SINGLE plan section of a sequential insurance TOB.
This section covers only: {plan_name}
The underwriter is: {underwriter}

EXTRACTION RULES:
1. This section has ONE plan only — extract its TWO-COLUMN table (benefit name | value).
2. "0.2. Plan Name" MUST be set to: "{plan_name}" (the primary product name, NOT a rider or sub-section name).
3. "0.1. Underwriter" MUST be set to: "{underwriter}"
4. *** CRITICAL — Extract ONLY the benefit VALUE. Do NOT include the row label or heading. ***
   The table has two columns: [Benefit Name] | [Value].
   You must ONLY extract the VALUE column. Never copy the Benefit Name into the value.
   WRONG: "Physician Consultation | Covered with 5% Copay"  ← row label included, WRONG
   CORRECT: "Covered with 5% Copay"                         ← value only, CORRECT
   WRONG: "Accommodation Type | Private"                    ← row label included, WRONG
   CORRECT: "Private"                                       ← value only, CORRECT
   WRONG: "Network | Premium"                               ← row label included, WRONG
   CORRECT: "Premium"                                       ← value only, CORRECT
5. Do NOT create separate entries for riders or add-ons — their benefits belong to this plan.
6. For benefits with conditions, use: [Status] | [Limit] | [Condition]
7. If a value is blank/not present, return "".
8. Capture the ENTIRE description for each benefit.
9. Choose the single best-supported value for each tracked field from this plan section.
10. For `3. Area of Cover`, extract only geographic scope or treatment territory. Do not include sports, recreation, activities, or exclusions unrelated to geography.
11. For `5. Network`, extract only the network name or provider panel name. Do not include copay, coinsurance, deductibles, waiting periods, or other conditions.
10.*** PRESERVE EXACT VALUES — DO NOT MODIFY: ***
   Keep ALL values EXACTLY as they appear in the quote/document.
   - Do NOT convert "Not Covered" to "Covered" or vice versa
   - Do NOT change numbers, limits, or percentages
   - Do NOT rephrase or interpret values — copy them verbatim
   - Preserve the original text from the document exactly as written

FIELDS TO EXTRACT:
{categories_text}
"""

        EXTRACT_USER = f"""Extract all the following fields from this single plan section.

PLAN: {plan_name}
SECTION CONTENT:
{section_md}

Return ONLY a JSON object (flat dict of field→value, value = benefit VALUE only, no row labels):
{{
  "0.1. Underwriter": "{underwriter}",
  "0.2. Plan Name": "{plan_name}",
  "2. Annual Limit per person": "BHD 10,000",
  ...
}}"""

        content, usage = llm_chat(
            EXTRACT_SYSTEM,
            EXTRACT_USER,
            json_mode=True,
            log_label=f"sequential_extract_{slugify_filename(plan_name)}",
            task_class="extraction",
        )
        tracker.record(_current_pdf, usage)
        plan_data = safe_json_loads(content, f"extract {plan_name}")

        if len(plan_data) == 1:
            only_key = list(plan_data.keys())[0]
            if isinstance(plan_data[only_key], dict):
                plan_data = plan_data[only_key]

        plan_data["0.2. Plan Name"] = plan_name
        plan_data["0.1. Underwriter"] = underwriter or plan_data.get("0.1. Underwriter", "")

        plan_data = clean_raw_fields(plan_data)

        all_plans[plan_name] = plan_data
        print(f"  ✓ {plan_name}: {sum(1 for v in plan_data.values() if v and str(v).strip())} non-empty fields")

    
    return all_plans


def _split_markdown_by_page(structured_md: str) -> dict:
    import re
    page_blocks = {}
    parts = re.split(r'---\n## PAGE (\d+)\n---', structured_md)
    i = 1
    while i < len(parts) - 1:
        try:
            page_num = int(parts[i])
            page_content = parts[i + 1]
            page_blocks[page_num] = page_content
        except (ValueError, IndexError):
            pass
        i += 2
    return page_blocks


def _detect_underwriter(cover_text: str) -> str:
    """Quick LLM call to extract underwriter name and plan names from cover/first pages."""
    system_msg = """You identify the underwriter and true plan names from an insurance quote.

<output_shape>
- Return JSON only.
- Schema: {"underwriter": string, "plan_names": string[]}
</output_shape>

<grounding_rules>
- Use only names supported by the provided text.
- Prefer explicit underwriter or insurer names.
- Prefer full plan names over generic labels.
- If the document uses categories like A/B/C, expand them to Category A / Category B when that wording is supported.
- If the document has only one table of benefits and shows both an option banner like Option 1 and a category label like Category A, treat them as the same plan and return only the category-style plan name.
- Do not return rider names, add-ons, section headers, bare letters, or blank strings as plan names.
- If a value is missing, return an empty string or an empty list rather than guessing.
</grounding_rules>"""
    try:
        content, usage = llm_chat(
            system_msg,
            "Document text:\n" + cover_text[:3000],
            json_mode=True,
            log_label="underwriter_detection",
            task_class="extraction",
        )
        tracker.record(_current_pdf, usage)
        result = safe_json_loads(content, "underwriter detection")
        uw = result.get("underwriter", "")
        plan_names = result.get("plan_names", [])
        # Store plan names globally so _extract_columnar can use them
        global _detected_plan_names
        _detected_plan_names = plan_names
        print(f"  ✓ Underwriter: {uw} | Plans detected: {plan_names}")
        return uw
    except Exception as e:
        print(f"  ⚠️  Underwriter detection failed: {e}")
        return ""


# ── STEP 2: VERIFICATION PASS ────────────────────────────────────────────────
def _missing_field_keys_for_verification(plan_fields: dict) -> list[str]:
    return [
        key for key, value in (plan_fields or {}).items()
        if _is_field_incomplete(key, value)
    ]


def _collect_document_level_candidates(
    structured_md: str,
    plan_names: list[str],
    target_plan_name: str,
    field_key: str,
    min_score: float = 0.45,
) -> list[dict[str, Any]]:
    if target_plan_name not in plan_names:
        return []

    target_idx = plan_names.index(target_plan_name)
    single_plan_mode = len(plan_names) == 1
    page_blocks = _split_markdown_by_page(structured_md)
    if not page_blocks:
        page_blocks = {1: structured_md}

    candidate_entries: list[dict[str, Any]] = []
    for page_number in sorted(page_blocks.keys()):
        candidate_entries.extend(
            _collect_field_candidate_entries(
                _parse_markdown_table_rows(page_blocks[page_number]),
                field_key,
                target_idx,
                single_plan_mode=single_plan_mode,
                page_number=page_number,
            )
        )

    if field_key == "3. Area of Cover":
        return candidate_entries
    return _select_best_candidate_entries(candidate_entries, max_candidates=10)


def run_document_evidence_enrichment(structured_md: str, extracted_plans: dict[str, dict]) -> dict[str, dict]:
    print("\n── Document Evidence Enrichment Pass ──")
    print("  ↷ Skipped: deterministic semantic enrichment disabled; using page update + verification only")
    return {
        name: fields
        for name, fields in (extracted_plans or {}).items()
        if isinstance(fields, dict)
    }


def run_post_extraction_verification(structured_md: str, extracted_plans: dict[str, dict]) -> dict[str, dict]:
    print("\n── Post-Extraction Verification Pass ──")
    extracted_plans = {name: fields for name, fields in (extracted_plans or {}).items() if isinstance(fields, dict)}
    plan_requests: dict[str, dict[str, Any]] = {}

    for plan_name, raw_fields in extracted_plans.items():
        missing_keys = _missing_field_keys_for_verification(raw_fields)
        if missing_keys:
            print(f"  → Queuing plan for one PDF-level verification call: {plan_name} ({len(missing_keys)} missing fields)")
            plan_requests[plan_name] = {
                "missing_keys": missing_keys,
                "current_values": {key: raw_fields.get(key, "") for key in missing_keys},
                "semantic_candidate_rows": _format_semantic_candidate_rows(
                    structured_md,
                    category_keys=missing_keys,
                    plan_names=list(extracted_plans.keys()),
                    target_plan_name=plan_name,
                ),
            }
        else:
            print(f"  ✓ {plan_name}: verification skipped (no missing fields)")

    if not plan_requests:
        return extracted_plans

    cache_key = json.dumps(
        {
            "structured_md_hash": hash(structured_md),
            "plans": {
                plan_name: {
                    "missing_keys": request["missing_keys"],
                    "current_values": request["current_values"],
                }
                for plan_name, request in plan_requests.items()
            },
        },
        sort_keys=True,
        ensure_ascii=False,
        default=str,
    )
    if cache_key in _verification_pass_cache:
        print("  Verification pass: using cached PDF-level result.")
        return dict(_verification_pass_cache[cache_key])

    key_to_category = {f"{num}. {name}": {"number": num, "name": name, "hint": hint} for num, name, hint in CATEGORIES}
    plans_payload = {}
    for plan_name, request in plan_requests.items():
        plans_payload[plan_name] = {
            "current_plan": extracted_plans[plan_name],
            "missing_fields": [
                {
                    "key": key,
                    "number": key_to_category.get(key, {}).get("number", ""),
                    "name": key_to_category.get(key, {}).get("name", key),
                    "hint": key_to_category.get(key, {}).get("hint", ""),
                }
                for key in request["missing_keys"]
            ],
            "semantic_candidate_rows": request["semantic_candidate_rows"],
        }

    verify_system = """You are re-checking an insurance quote after the full document has already been extracted.

<task>
- Review the full document once.
- Re-check only the listed missing or incomplete fields for each plan.
- Return updates grouped by plan name.
</task>

<output_shape>
- Return JSON only.
- Schema:
  {
    "plans": {
      "<plan_name>": {
        "<field_key>": string
      }
    }
  }
- For each plan, return only the missing field keys that were requested for that plan.
</output_shape>

<grounding_rules>
- Use exact wording from the document.
- Extract the value only, never the row label.
- Map semantically equivalent source labels to the closest requested field even when the wording differs.
- Do not combine unrelated values from nearby sections.
- Annual limit and other overall aggregate fields must remain separate from maternity, dental, optical, consultation, diagnostic, and other benefit-specific sublimits.
- `4. TPA` must come only from administrator-style evidence. Do not use maternity, consultation, network, physician, or benefit rows to fill it.
- `14. Diagnostic Tests & Laboratory Tests` must not use hearing aids, auditory devices, cochlear items, glasses, frames, lenses, spectacles, contact lenses, or other vision-aid rows.
- `20. Life Threatening Maternity Complications` must not use routine maternity, antenatal, prenatal, postnatal, normal delivery, or standard outpatient consultation rows.
- `35. Work Related Injury` must come only from occupational, employment, workplace, workers compensation, or on-duty injury evidence.
- `40. Claims Outside Network Within Country` must come only from out-of-network or reimbursement wording within the country, not in-network, direct-billing, or cashless rows.
- `6. Pre-existing and Chronic Conditions` may contain multiple distinct lines when the document lists regional or jurisdiction-specific variants such as Abu Dhabi, Dubai, within-emirate, outside-emirate, pharmacy-related pre-existing cover, or outside-UAE wording. Keep all supported variants together.
- Use the semantic candidate rows as retrieval hints, but verify against the document before deciding the final value.
- If the source says Not Covered, return Not Covered exactly.
- If a field is still not supported, return an empty string.
- Prefer blank over uncertain.
</grounding_rules>"""

    verify_user = f"""These plans still have missing or incomplete fields after the full extraction pass.

INPUT JSON:
{json.dumps(plans_payload, indent=2, ensure_ascii=False)}

DOCUMENT:
{structured_md}

Return only the missing field updates grouped under the same plan names."""

    content, usage = llm_chat(
        verify_system,
        verify_user,
        json_mode=True,
        log_label="verification_pdf",
        task_class="verification",
    )
    tracker.record(_current_pdf, usage)
    result = safe_json_loads(content, "verification pdf pass")
    verified_updates = result.get("plans", result) if isinstance(result, dict) else {}

    verified_plans: dict[str, dict] = {}
    for plan_name, raw_fields in extracted_plans.items():
        updated = merge_plan_records(
            {},
            raw_fields,
            fallback_plan_name=raw_fields.get("0.2. Plan Name", "").strip() or plan_name,
            fallback_underwriter=raw_fields.get("0.1. Underwriter", "").strip(),
        )
        corrections = clean_raw_fields(verified_updates.get(plan_name, {}) if isinstance(verified_updates, dict) else {})
        recovered = 0
        confirmed_missing = 0
        for key in plan_requests.get(plan_name, {}).get("missing_keys", []):
            document_candidates = _collect_document_level_candidates(
                structured_md,
                plan_names=list(extracted_plans.keys()),
                target_plan_name=plan_name,
                field_key=key,
            )
            value = _normalize_verification_recovered_value(
                key,
                corrections.get(key, ""),
                candidate_entries=document_candidates,
                source="llm",
            )
            if not _is_blankish(value):
                updated = merge_plan_records(
                    updated,
                    {key: value},
                    fallback_plan_name=updated.get("0.2. Plan Name", "").strip(),
                    fallback_underwriter=updated.get("0.1. Underwriter", "").strip(),
                )
                recovered += 1
            deterministic_value = _normalize_verification_recovered_value(
                key,
                _deterministic_candidate_backfill_value(key, document_candidates, current_value=updated.get(key, "")),
                candidate_entries=document_candidates,
                source="deterministic",
            )
            if not _is_blankish(deterministic_value) and _prefer_field_value(key, updated.get(key, ""), deterministic_value) != str(updated.get(key, "") or "").strip():
                updated = merge_plan_records(
                    updated,
                    {key: deterministic_value},
                    fallback_plan_name=updated.get("0.2. Plan Name", "").strip(),
                    fallback_underwriter=updated.get("0.1. Underwriter", "").strip(),
                )
                recovered += 1
            elif _is_field_incomplete(key, updated.get(key, "")):
                confirmed_missing += 1
        if plan_name in plan_requests:
            print(f"    {plan_name}: recovered {recovered} | confirmed missing {confirmed_missing}")
        verified_plans[plan_name] = updated

    _verification_pass_cache[cache_key] = dict(verified_plans)
    return verified_plans


# ── STEP 2.5: NORMALIZATION PASS (PER-PLAN) ───────────────────────────────────
NORMALIZE_SYSTEM = """You are validating extracted insurance plans for consistency.

<task>
- Return the same JSON keys you received for every plan.
- Preserve the exact extracted wording unless the input is malformed JSON content.
</task>

<output_shape>
- Return JSON only.
- Return the same top-level plan names and field keys as the input exactly.
</output_shape>

<grounding_rules>
- Do not invent, summarize, normalize, or reinterpret values.
- Keep underwriter and plan name exactly as provided.
- Keep Not Covered, percentages, limits, waiting periods, and conditions exactly as provided.
- If a field is blank in the input, keep it blank.
</grounding_rules>"""


def run_post_extraction_normalization(verified_plans: dict[str, dict]) -> dict[str, dict]:
    verified_plans = {name: fields for name, fields in (verified_plans or {}).items() if isinstance(fields, dict)}
    cache_key = json.dumps(verified_plans, sort_keys=True, ensure_ascii=False, default=str)
    if cache_key in _normalization_pass_cache:
        return dict(_normalization_pass_cache[cache_key])

    print("\n── Post-Extraction Normalization Pass ──")
    print(f"  Normalizing all plans in one PDF-level call ({len(verified_plans)} plans)")

    user_prompt = f"""Normalize these extracted plans in one PDF-level pass.
Return a JSON object with EXACTLY the same top-level plan names and field keys as the input. Do not add, remove, or rename any keys.

DATA:
{json.dumps(verified_plans, indent=2, ensure_ascii=False)}
"""
    try:
        content, usage = llm_chat(
            NORMALIZE_SYSTEM,
            user_prompt,
            json_mode=True,
            log_label="normalization_pdf",
            task_class="normalization",
        )
        tracker.record(_current_pdf, usage)
        result = safe_json_loads(content, "normalization")
        normalized_result = result.get("plans", result) if isinstance(result, dict) else {}

        normalized_plans: dict[str, dict] = {}
        generic_names = {"plan", "category", "n/a", "none", "", "not mentioned", "unknown", "plan name"}
        for plan_name, raw_fields in verified_plans.items():
            normalized = normalized_result.get(plan_name, raw_fields) if isinstance(normalized_result, dict) else raw_fields
            if not normalized:
                normalized = raw_fields

            normalized = dict(normalized)
            auth_name = raw_fields.get("0.2. Plan Name", "").strip() or str(plan_name).strip()
            post_name = normalized.get("0.2. Plan Name", "").strip()
            if not post_name or post_name.lower() in generic_names:
                normalized["0.2. Plan Name"] = auth_name

            non_empty = sum(1 for value in normalized.values() if value and str(value).strip())
            print(f"  ✓ {plan_name}: {non_empty} non-empty fields")
            normalized_plans[plan_name] = normalized

        _normalization_pass_cache[cache_key] = dict(normalized_plans)
        return normalized_plans
    except Exception as e:
        print(f"  ⚠️  PDF-level normalization failed: {e} — using verified data")
        return verified_plans


# ── DISABLED COSTLY PASSES ───────────────────────────────────────────────────
# Comparative best-marking, per-plan scoring, and underwriting conclusions are
# intentionally removed from the active Bahrain runner to avoid unnecessary LLM
# calls. The current main flow stores normalized plans directly.


def generate_overall_uw_summary_one_liner(final_results: dict, ordered_keys: list[str]) -> str:
    if not final_results:
        return ""

    summary_input = []
    for unique_key, data in final_results.items():
        normalized = data.get("normalized", {}) or {}
        summary_input.append(
            {
                "unique_key": unique_key,
                "underwriter": str(normalized.get("0.1. Underwriter", "") or "").strip(),
                "plan_name": str(normalized.get("0.2. Plan Name", "") or "").strip(),
                "benefits": {
                    key: str(normalized.get(key, "") or "").strip()
                    for key in ordered_keys
                    if key not in {"0.1. Underwriter", "0.2. Plan Name"} and str(normalized.get(key, "") or "").strip()
                },
            }
        )

    system_prompt = """You write a one-line executive underwriter summary for an insurance comparison.

<task>
- Review all plans and underwriters together.
- Return exactly one concise line of text.
</task>

<writing_rules>
- Maximum 35 words.
- If multiple underwriters are present, identify the strongest underwriter or say the comparison is too close to call.
- If only one underwriter is present, say that underwriter is the only quote and recommend comparing plans within it.
- Keep the wording business-ready and decision-oriented.
- Do not use bullets, labels, or line breaks.
- Do not invent facts beyond the provided data.
</writing_rules>"""

    user_prompt = f"""Provide a single-line executive UW summary for this full comparison set.

INPUT JSON:
{json.dumps(summary_input, indent=2, ensure_ascii=False)}

Return one line only."""

    try:
        content, usage = llm_chat(
            system_prompt,
            user_prompt,
            json_mode=False,
            log_label="overall_uw_summary",
            task_class="uw_summary",
            chunk_size=len(summary_input),
            log_context={"plan_count": len(summary_input)},
        )
        tracker.record(_current_pdf, usage)
        return _normalize_summary_text(content, max_length=240)
    except Exception as exc:
        print(f"  ⚠️  Overall UW summary generation failed: {exc} — using deterministic fallback")

    underwriter_keys, grouped_underwriters = _group_plans_by_underwriter(final_results)
    underwriters = [grouped_underwriters[key]["display_label"] for key in underwriter_keys]
    if len(underwriters) <= 1:
        only_uw = underwriters[0] if underwriters else "the available underwriter"
        return f"{only_uw} is the only quoted underwriter, so the decision should focus on plan-level differences within that quote."
    return "Multiple underwriters were compared, and the final recommendation should be confirmed against the extracted plan-level benefit differences."


SUMMARY_CELL_MAX_LENGTH = 420


def _normalize_summary_text(text: str, max_length: int | None = 180) -> str:
    raw_text = str(text or "").replace("**[[BEST]]**", "").replace("\r", "\n")
    raw_lines = [line for line in raw_text.split("\n") if str(line).strip()]
    if not raw_lines:
        raw_lines = [raw_text]

    cleaned_lines: list[str] = []
    seen_lines: set[str] = set()
    for line in raw_lines:
        cleaned = re.sub(r"[ \t]+", " ", str(line or "")).strip(" |;,-")
        if not cleaned:
            continue
        dedupe_key = _normalize_match_text(cleaned)
        if dedupe_key and dedupe_key in seen_lines:
            continue
        if dedupe_key:
            seen_lines.add(dedupe_key)
        cleaned_lines.append(cleaned)

    cleaned_text = "\n".join(cleaned_lines).strip()
    if max_length is not None and max_length > 0 and len(cleaned_text) > max_length:
        cleaned_text = cleaned_text[: max_length - 3].rstrip() + "..."
    return cleaned_text


def _strip_prefixed_label(text: str, label: str) -> str:
    escaped_label = re.escape(str(label or "").strip())
    return re.sub(rf"^{escaped_label}\s*:\s*", "", str(text or "").strip(), flags=re.IGNORECASE)


def _coerce_summary_candidate(value: Any) -> str:
    if isinstance(value, dict):
        for key in ("text", "summary", "value", "result"):
            candidate = value.get(key)
            if candidate not in {None, ""}:
                return _coerce_summary_candidate(candidate)
        return ""
    if isinstance(value, list):
        parts = [_coerce_summary_candidate(item) for item in value]
        return " | ".join(part for part in parts if part)
    return str(value or "")


def _has_material_benefit_detail(text: str) -> bool:
    lowered = str(text or "").lower()
    if not lowered.strip():
        return False
    if re.search(r"\b(?:aed|usd|bhd|sar|qar|omr)\b", lowered):
        return True
    if re.search(r"\d+\s*%", lowered):
        return True
    if re.search(r"\b\d[\d,]*(?:\.\d+)?\b", lowered) and any(token in lowered for token in ["limit", "annual", "day", "visit", "claim", "aggregate", "room", "session"]):
        return True
    detail_tokens = [
        "covered",
        "not covered",
        "copay",
        "co-pay",
        "coinsurance",
        "co-insurance",
        "deductible",
        "waiting",
        "subject to",
        "prior approval",
        "pre-approval",
        "approval",
        "limit",
        "sublimit",
        "sub limit",
        "reimbursement",
        "actuals",
        "restricted",
        "excluding",
        "emergency",
        "network",
    ]
    return any(token in lowered for token in detail_tokens)


def _split_summary_segments(text: str) -> list[str]:
    raw_text = str(text or "").replace("\r", "\n")
    raw_segments = re.split(r"\s*(?:\n+|\||;|•)\s*", raw_text)
    segments: list[str] = []
    seen_segments: set[str] = set()
    for raw_segment in raw_segments:
        cleaned = _normalize_summary_text(raw_segment, max_length=None)
        cleaned = re.sub(r"^[\-:,. ]+", "", cleaned).strip()
        if not cleaned:
            continue
        dedupe_key = _normalize_match_text(cleaned)
        if dedupe_key and dedupe_key in seen_segments:
            continue
        if dedupe_key:
            seen_segments.add(dedupe_key)
        segments.append(cleaned)
    return segments


def _looks_like_amount_or_limit(segment: str) -> bool:
    lowered = str(segment or "").lower()
    if not lowered:
        return False
    if re.search(r"\b(?:aed|usd|bhd|sar|qar|omr)\b", lowered):
        return True
    if any(token in lowered for token in ["annual limit", "limit", "sublimit", "sub limit", "aggregate", "per year", "per visit", "per claim", "up to", "maximum", "max "]):
        return True
    return bool(re.search(r"\b\d[\d,]*(?:\.\d+)?\b", lowered) and not any(token in lowered for token in ["day", "days", "month", "months"]))


def _looks_like_condition(segment: str) -> bool:
    lowered = str(segment or "").lower()
    if not lowered:
        return False
    condition_tokens = [
        "copay",
        "co-pay",
        "coinsurance",
        "co-insurance",
        "deductible",
        "waiting",
        "subject to",
        "approval",
        "pre-approval",
        "prior approval",
        "restricted",
        "excluding",
        "excluded",
        "emergency",
        "network",
        "reimbursement",
        "actuals",
        "room",
        "icu",
        "hospital",
        "pharmacy",
        "services",
        "consultation",
        "out-patient",
        "out patient",
        "in-patient",
        "in patient",
        "covered",
        "not covered",
    ]
    return any(token in lowered for token in condition_tokens)


def _cleanup_segment(segment: str) -> str:
    cleaned = _normalize_summary_text(segment, max_length=None)
    cleaned = re.sub(r"^(?:covered|not covered)\s*[:|-]?\s*", "", cleaned, flags=re.IGNORECASE).strip(" |;,-")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _simplify_summary_segment(segment: str) -> str:
    cleaned = _cleanup_segment(segment)
    if not cleaned:
        return ""

    simplified = cleaned
    replacement_rules = [
        (r"(?i)subject to providing (?:a )?continuity of coverage \(coc\) certificate in the uae", "CoC required in UAE"),
        (r"(?i)subject to continuity of coverage proof", "CoC proof required"),
        (r"(?i)waiting period\s*(?:of)?\s*6\s*months?", "6-month waiting period"),
        (r"(?i)waiting period\s*(?:of)?\s*12\s*months?", "12-month waiting period"),
        (r"(?i)emergency within 6 month exclusion period covered up to annual aggregate limit", "Emergency during waiting covered to annual aggregate limit"),
        (r"(?i)waiting period does not apply for previously insured members", "Waiting waived for previously insured members"),
        (r"(?i)subject to pre-approval", "Pre-approval required"),
        (r"(?i)subject to prior approval", "Prior approval required"),
        (r"(?i)actuals capped? at", "Actuals capped at"),
        (r"(?i)up to the aml", "Up to AML"),
        (r"(?i)limit:\s*upto", "Up to"),
        (r"(?i)co-?insurance", "Coinsurance"),
        (r"(?i)co\s*pay", "Copay"),
        (r"(?i)co-?pay", "Copay"),
    ]
    for pattern, replacement in replacement_rules:
        simplified = re.sub(pattern, replacement, simplified)

    simplified = re.sub(r"(?i)\bcontinuity of coverage\b", "CoC", simplified)
    simplified = re.sub(r"\s+", " ", simplified).strip(" |;,-")
    return simplified


def _select_priority_parts(parts: list[str], limit: int) -> list[str]:
    selected: list[str] = []
    for part in parts:
        if part and part not in selected:
            selected.append(part)
        if len(selected) >= limit:
            break
    return selected


def _is_deductible_segment(segment: str) -> bool:
    lowered = str(segment or "").lower()
    return "deductible" in lowered


def _is_coinsurance_segment(segment: str) -> bool:
    lowered = str(segment or "").lower()
    return any(token in lowered for token in ["coinsurance", "co-insurance"])


def _is_copay_segment(segment: str) -> bool:
    lowered = str(segment or "").lower()
    return any(token in lowered for token in ["copay", "co-pay", "co pay"])


def _is_subjectivity_segment(segment: str) -> bool:
    lowered = str(segment or "").lower()
    subjectivity_tokens = [
        "subject to",
        "approval",
        "pre-approval",
        "prior approval",
        "restricted",
        "excluding",
        "excluded",
        "emergency",
        "reimbursement",
        "actuals",
        "network",
        "room",
        "icu",
        "services",
        "consultation",
        "only",
        "applicable",
        "conditions",
    ]
    return any(token in lowered for token in subjectivity_tokens)


def _format_coverage_summary_text(summary_text: str = "", raw_value: str = "", max_length: int = SUMMARY_CELL_MAX_LENGTH) -> str:
    source_text = _coerce_summary_candidate(summary_text) or _coerce_summary_candidate(raw_value)
    normalized_source = _normalize_summary_text(source_text, max_length=max_length)
    if not normalized_source:
        return ""

    lowered_source = normalized_source.lower()
    if lowered_source in {"not mentioned", "n/a", "na", "none", "unknown"}:
        return normalized_source

    combined_candidates = "\n".join(
        _coerce_summary_candidate(candidate)
        for candidate in (summary_text, raw_value)
        if _coerce_summary_candidate(candidate)
    )
    lowered_candidates = combined_candidates.lower().strip()

    status = ""
    if re.match(r"^\s*not covered\b", lowered_candidates):
        status = "Not Covered"
    elif re.match(r"^\s*covered\b", lowered_candidates):
        status = "Covered"
    else:
        has_not_covered = "not covered" in lowered_candidates
        has_covered = "covered" in lowered_candidates
        if has_covered and not has_not_covered:
            status = "Covered"
        elif has_not_covered and not has_covered:
            status = "Not Covered"
        elif has_covered and has_not_covered:
            status = "Covered"

    if not status and normalized_source and _has_material_benefit_detail(normalized_source):
        status = "Covered"

    if not status and not _has_material_benefit_detail(normalized_source):
        return normalized_source

    segments = _split_summary_segments(source_text)
    amount_parts: list[str] = []
    deductible_parts: list[str] = []
    coinsurance_parts: list[str] = []
    copay_parts: list[str] = []
    subjectivity_parts: list[str] = []
    other_parts: list[str] = []
    seen_parts: set[str] = set()
    for segment in segments:
        cleaned = _simplify_summary_segment(segment)
        if not cleaned:
            continue
        if status and _normalize_match_text(cleaned) == _normalize_match_text(status):
            continue
        dedupe_key = _normalize_match_text(cleaned)
        if dedupe_key and dedupe_key in seen_parts:
            continue
        if dedupe_key:
            seen_parts.add(dedupe_key)
        if _looks_like_amount_or_limit(cleaned):
            amount_parts.append(cleaned)
        elif _is_deductible_segment(cleaned):
            deductible_parts.append(cleaned)
        elif _is_coinsurance_segment(cleaned):
            coinsurance_parts.append(cleaned)
        elif _is_copay_segment(cleaned):
            copay_parts.append(cleaned)
        elif _is_subjectivity_segment(cleaned) or _looks_like_condition(cleaned):
            subjectivity_parts.append(cleaned)
        else:
            other_parts.append(cleaned)

    amount_parts = _select_priority_parts(amount_parts, 1)
    deductible_parts = _select_priority_parts(deductible_parts, 1)
    coinsurance_parts = _select_priority_parts(coinsurance_parts, 1)
    copay_parts = _select_priority_parts(copay_parts, 1)
    subjectivity_parts = _select_priority_parts(subjectivity_parts, 1)
    other_parts = _select_priority_parts(other_parts, 0)

    financial_parts = deductible_parts + coinsurance_parts + copay_parts
    financial_parts = _select_priority_parts(financial_parts, 1)

    if not status:
        ordered_parts = amount_parts + financial_parts + subjectivity_parts + other_parts
        return _normalize_summary_text(" | ".join(ordered_parts) if ordered_parts else normalized_source, max_length=max_length)

    ordered_parts = amount_parts + financial_parts + subjectivity_parts + other_parts
    if not ordered_parts:
        return status
    return _normalize_summary_text(f"{status} | {' | '.join(ordered_parts)}", max_length=max_length)


def _format_planwise_summary(plan_items: list[dict[str, Any]], summary_text: str = "", max_length: int = SUMMARY_CELL_MAX_LENGTH) -> str:
    if not plan_items:
        llm_value = _coerce_nonempty_summary_text(summary_text, max_length=max_length)
        return llm_value or _format_coverage_summary_text(summary_text, summary_text, max_length=max_length)

    normalized_summary = _normalize_summary_text(summary_text, max_length=max_length)
    summary_lines = [line.strip() for line in normalized_summary.split("\n") if line.strip()]
    mapped_lines: dict[str, str] = {}
    for line in summary_lines:
        for item in plan_items:
            plan_name = str(item.get("plan_name", "") or "").strip()
            if not plan_name:
                continue
            if re.match(rf"^{re.escape(plan_name)}\s*:", line, flags=re.IGNORECASE):
                mapped_lines[plan_name] = _strip_prefixed_label(line, plan_name)
                break

    formatted_lines: list[str] = []
    is_single_plan = len(plan_items) == 1
    for item in plan_items:
        plan_name = str(item.get("plan_name", "") or "").strip()
        raw_input = str(item.get("value", "") or "")
        raw_value = _format_coverage_summary_text(raw_input, raw_input, max_length=max_length)
        selected_value = mapped_lines.get(plan_name) or _coerce_nonempty_summary_text(summary_text, max_length=max_length) or raw_value
        selected_value = _coerce_nonempty_summary_text(_strip_prefixed_label(selected_value, plan_name), max_length=max_length) or raw_value
        if plan_name and selected_value:
            if is_single_plan:
                formatted_lines.append(selected_value)
            else:
                formatted_lines.append(f"{plan_name}: {selected_value}")
        elif selected_value:
            formatted_lines.append(selected_value)

    if not formatted_lines and normalized_summary:
        return normalized_summary
    return _normalize_summary_text("\n".join(formatted_lines), max_length=max_length)


def _format_benefit_summary_text(summary_text: str = "", raw_value: str = "", max_length: int = SUMMARY_CELL_MAX_LENGTH) -> str:
    return _coerce_nonempty_summary_text(summary_text, max_length=max_length)


def _coerce_nonempty_summary_text(value: Any, max_length: int = SUMMARY_CELL_MAX_LENGTH) -> str:
    return _normalize_summary_text(_coerce_summary_candidate(value), max_length=max_length)


def _canonical_underwriter_key(value: Any) -> str:
    raw = str(value or "").strip()
    return _normalize_match_text(raw) or raw.lower()


_CATEGORY_SCOPED_FIELD_SEPARATOR = " ||| "


def _category_label_for_index(index: int) -> str:
    safe_index = max(0, int(index or 0))
    letters: list[str] = []
    while True:
        safe_index, remainder = divmod(safe_index, 26)
        letters.append(chr(ord("A") + remainder))
        if safe_index == 0:
            break
        safe_index -= 1
    return f"CAT {''.join(reversed(letters))}"


def _make_category_scoped_field_key(field_key: str, category_label: str = "") -> str:
    cleaned_category = str(category_label or "").strip()
    return f"{field_key}{_CATEGORY_SCOPED_FIELD_SEPARATOR}{cleaned_category}" if cleaned_category else str(field_key or "")


def _split_category_scoped_field_key(field_key: str) -> tuple[str, str]:
    raw_key = str(field_key or "")
    if _CATEGORY_SCOPED_FIELD_SEPARATOR not in raw_key:
        return raw_key, ""
    base_key, category_label = raw_key.split(_CATEGORY_SCOPED_FIELD_SEPARATOR, 1)
    return base_key.strip(), category_label.strip()


def _format_report_row_label(field_key: str) -> str:
    base_key, category_label = _split_category_scoped_field_key(field_key)
    label = display_benefit_label(base_key)
    return f"{label} ({category_label})" if category_label else label


def _format_category_display_label(category_label: str) -> str:
    cleaned = str(category_label or "").strip()
    if not cleaned:
        return ""
    return cleaned.title()


def _derive_source_group_label(pdf_name: str, underwriter_label: str = "") -> str:
    stem = Path(str(pdf_name or "").strip()).stem
    cleaned = re.sub(r"[_\-]+", " ", stem)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    underwriter_tokens = [token for token in re.split(r"\W+", str(underwriter_label or "")) if token]
    for token in sorted(underwriter_tokens, key=len, reverse=True):
        cleaned = re.sub(rf"(?i)\b{re.escape(token)}\b", " ", cleaned)
    cleaned = re.sub(r"\bpdf\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" _-")
    return _smart_title_case(cleaned) if cleaned else ""


def _collapse_category_scoped_summary(
    summary: dict[str, dict[str, str]],
    best_uw: dict[str, set[str]],
    worst_uw: dict[str, set[str]],
    underwriters: list[str],
    ordered_base_keys: list[str],
    category_labels: list[str],
    category_source_lookup: dict[tuple[str, str], str] | None = None,
) -> tuple[
    dict[str, dict[str, str]],
    dict[str, set[str]],
    dict[str, set[str]],
    list[str],
    dict[str, dict[str, list[dict[str, str]]]],
]:
    def _line_status(best_labels: set[str], worst_labels: set[str], underwriter_label: str) -> str:
        in_best = underwriter_label in best_labels
        in_worst = underwriter_label in worst_labels
        if in_best and not in_worst:
            return "best"
        if in_worst and not in_best:
            return "worst"
        return ""

    if not category_labels:
        row_keys = list(ordered_base_keys)
        normalized_best = {field_key: set(best_uw.get(field_key, set())) for field_key in row_keys}
        normalized_worst = {field_key: set(worst_uw.get(field_key, set())) for field_key in row_keys}
        line_highlights: dict[str, dict[str, list[dict[str, str]]]] = {field_key: {} for field_key in row_keys}
        for field_key in row_keys:
            for uw in underwriters:
                value = _format_raw_cell_text(summary.get(field_key, {}).get(uw, ""))
                if not value:
                    line_highlights[field_key][uw] = []
                    continue
                status = _line_status(normalized_best.get(field_key, set()), normalized_worst.get(field_key, set()), uw)
                line_highlights[field_key][uw] = [{"category_label": "", "status": status}]
        return summary, normalized_best, normalized_worst, row_keys, line_highlights

    collapsed_summary: dict[str, dict[str, str]] = {
        field_key: {uw: "" for uw in underwriters}
        for field_key in ordered_base_keys
    }
    collapsed_best: dict[str, set[str]] = {field_key: set() for field_key in ordered_base_keys}
    collapsed_worst: dict[str, set[str]] = {field_key: set() for field_key in ordered_base_keys}
    collapsed_line_highlights: dict[str, dict[str, list[dict[str, str]]]] = {
        field_key: {uw: [] for uw in underwriters}
        for field_key in ordered_base_keys
    }

    for field_key in ordered_base_keys:
        for uw in underwriters:
            category_entries: list[tuple[int, str, str, str]] = []
            line_entries: list[dict[str, str]] = []
            for category_order, category_label in enumerate(category_labels):
                scoped_key = _make_category_scoped_field_key(field_key, category_label)
                value = _coerce_nonempty_summary_text(summary.get(scoped_key, {}).get(uw, ""), max_length=SUMMARY_CELL_MAX_LENGTH)
                if value:
                    source_label = str((category_source_lookup or {}).get((uw, category_label), "") or "").strip()
                    category_entries.append((category_order, category_label, value, source_label))
                    display_category_label = _format_category_display_label(category_label)
                    line_entries.append(
                        {
                            "category_label": display_category_label,
                            "status": _line_status(best_uw.get(scoped_key, set()), worst_uw.get(scoped_key, set()), uw),
                        }
                    )

            if not category_entries:
                collapsed_summary[field_key][uw] = ""
                collapsed_line_highlights[field_key][uw] = []
                continue

            if len(category_entries) == 1:
                _, category_label, value, _ = category_entries[0]
                if category_label:
                    collapsed_summary[field_key][uw] = f"{_format_category_display_label(category_label)} : {value}"
                else:
                    collapsed_summary[field_key][uw] = value
                collapsed_line_highlights[field_key][uw] = line_entries
                continue

            collapsed_summary[field_key][uw] = "\n".join(
                f"{_format_category_display_label(category_label)} : {value}"
                for _, category_label, value, _ in category_entries
            ).strip()
            collapsed_line_highlights[field_key][uw] = line_entries

        best_union: set[str] = set()
        worst_union: set[str] = set()
        for category_label in category_labels:
            scoped_key = _make_category_scoped_field_key(field_key, category_label)
            best_union.update(best_uw.get(scoped_key, set()))
            worst_union.update(worst_uw.get(scoped_key, set()))

        collapsed_best[field_key] = best_union
        collapsed_worst[field_key] = worst_union

    return collapsed_summary, collapsed_best, collapsed_worst, list(ordered_base_keys), collapsed_line_highlights


def _preferred_underwriter_label(labels: list[str]) -> str:
    cleaned = [str(label).strip() for label in labels if str(label).strip()]
    if not cleaned:
        return ""

    counts: dict[str, int] = {}
    for label in cleaned:
        counts[label] = counts.get(label, 0) + 1

    return max(cleaned, key=lambda label: (counts[label], len(label), label))


def _extract_pdf_name_from_unique_key(unique_key: str) -> str:
    return str(unique_key or "").split("|")[0].strip()


def _extract_plan_name_from_unique_key(unique_key: str) -> str:
    return str(unique_key or "").split("|")[-1].strip()


def _normalize_pdf_underwriter_overrides(overrides: dict[str, Any] | None) -> dict[str, str]:
    normalized: dict[str, str] = {}
    for source_key, label in (overrides or {}).items():
        cleaned_label = str(label or "").strip()
        if not cleaned_label:
            continue
        cleaned_key = str(source_key or "").strip()
        if not cleaned_key:
            continue
        normalized[cleaned_key] = cleaned_label
    return normalized


def _first_present_value(source: dict[str, Any], keys: list[str]) -> Any:
    for key in keys:
        if key in source and source[key] is not None:
            return source[key]
    return None


def _normalize_optional_text(value: Any) -> str:
    return str(value or "").strip()


def _normalize_bool_flag(value: Any, field_name: str) -> bool:
    if isinstance(value, bool):
        return value

    normalized_value = str(value or "").strip().lower()
    if normalized_value in {"", "0", "false", "no", "n", "new", "candidate"}:
        return False
    if normalized_value in {"1", "true", "yes", "y", "existing", "baseline"}:
        return True
    raise ValueError(f"Invalid {field_name} value '{value}'. Use true/false, yes/no, or 1/0.")


def _normalize_pdf_file_metadata(file_metadata: dict[str, Any] | None) -> dict[str, dict[str, Any]]:
    normalized: dict[str, dict[str, Any]] = {}
    for source_key, raw_value in (file_metadata or {}).items():
        cleaned_key = os.path.basename(str(source_key or "").strip())
        if not cleaned_key:
            continue

        raw_dict = dict(raw_value or {}) if isinstance(raw_value, dict) else {}
        existing_value = _first_present_value(raw_dict, ["existing_benefit", "existing", "baseline", "is_existing"])

        normalized_entry = {
            "existing_benefit": _normalize_bool_flag(existing_value, "existing benefit") if existing_value is not None else False,
        }
        if normalized_entry["existing_benefit"]:
            normalized[cleaned_key] = normalized_entry
    return normalized


def _normalize_report_metadata(metadata: dict[str, Any] | None) -> dict[str, Any]:
    metadata = dict(metadata or {})
    pdf_underwriter_overrides = _normalize_pdf_underwriter_overrides(
        metadata.get("pdf_underwriter_overrides")
        or metadata.get("underwriter_overrides")
        or metadata.get("uw_overrides")
    )
    pdf_file_metadata = _normalize_pdf_file_metadata(
        metadata.get("pdf_file_metadata")
        or metadata.get("file_metadata")
        or metadata.get("pdf_metadata")
    )
    baseline_pdf = os.path.basename(str(metadata.get("baseline_pdf") or "").strip())
    baseline_underwriter = str(metadata.get("baseline_underwriter") or "").strip()
    return {
        "pdf_underwriter_overrides": pdf_underwriter_overrides,
        "pdf_file_metadata": pdf_file_metadata,
        "baseline_pdf": baseline_pdf,
        "baseline_underwriter": baseline_underwriter,
    }


def _normalize_file_entries(file_entries: list[Any] | None) -> list[dict[str, Any]]:
    normalized_entries: list[dict[str, Any]] = []
    for entry in file_entries or []:
        pdf_path = ""
        underwriter_name = ""
        existing_benefit = None

        if isinstance(entry, str):
            pdf_path = entry.strip()
        elif isinstance(entry, (tuple, list)):
            if entry:
                pdf_path = str(entry[0] or "").strip()
            if len(entry) > 1:
                underwriter_name = str(entry[1] or "").strip()
            if len(entry) > 2 and entry[2] is not None:
                existing_benefit = _normalize_bool_flag(entry[2], "existing benefit")
        elif isinstance(entry, dict):
            pdf_path = str(
                entry.get("path")
                or entry.get("file_path")
                or entry.get("pdf_path")
                or entry.get("pdf")
                or entry.get("file")
                or ""
            ).strip()
            underwriter_name = str(
                entry.get("uw_name")
                or entry.get("underwriter")
                or entry.get("underwriter_name")
                or ""
            ).strip()
            existing_value = _first_present_value(entry, ["existing_benefit", "existing", "baseline", "is_existing"])
            if existing_value is not None:
                existing_benefit = _normalize_bool_flag(existing_value, "existing benefit")

        if pdf_path:
            normalized_entries.append({
                "path": pdf_path,
                "uw_name": underwriter_name,
                "existing_benefit": existing_benefit,
            })

    return normalized_entries


def _merge_pdf_file_metadata(base_metadata: dict[str, dict[str, Any]] | None, override_metadata: dict[str, dict[str, Any]] | None) -> dict[str, dict[str, Any]]:
    merged: dict[str, dict[str, Any]] = {}
    for metadata_map in (base_metadata or {}, override_metadata or {}):
        for pdf_name, raw_metadata in metadata_map.items():
            cleaned_pdf_name = os.path.basename(str(pdf_name or "").strip())
            if not cleaned_pdf_name:
                continue
            target = merged.setdefault(
                cleaned_pdf_name,
                {"existing_benefit": False},
            )
            source = dict(raw_metadata or {})
            if source.get("existing_benefit") is True:
                target["existing_benefit"] = True
    return merged


def _build_entry_file_metadata(file_entries: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    metadata: dict[str, dict[str, Any]] = {}
    for entry in file_entries:
        pdf_name = os.path.basename(str(entry.get("path") or "").strip())
        if not pdf_name:
            continue
        entry_metadata = {
            "existing_benefit": bool(entry.get("existing_benefit") is True),
        }
        if entry.get("existing_benefit") is not None:
            metadata[pdf_name] = entry_metadata
    return metadata


def _validate_baseline_selection(pdf_paths: list[str], report_metadata: dict[str, Any]) -> dict[str, Any]:
    pdf_file_metadata = report_metadata.get("pdf_file_metadata", {}) or {}
    active_pdf_names = [os.path.basename(str(pdf_path or "").strip()) for pdf_path in pdf_paths if str(pdf_path or "").strip()]
    baseline_pdfs = [
        pdf_name for pdf_name in active_pdf_names
        if bool((pdf_file_metadata.get(pdf_name, {}) or {}).get("existing_benefit"))
    ]
    if len(active_pdf_names) > 1 and not baseline_pdfs:
        raise SystemExit("Mark at least one input PDF as the existing baseline using --existing-benefit true/false for each PDF.")

    report_metadata["baseline_pdf"] = baseline_pdfs[0] if baseline_pdfs else ""
    return report_metadata


def _get_effective_file_metadata(unique_key: str, data: dict[str, Any], report_metadata: dict[str, Any] | None = None) -> dict[str, Any]:
    pdf_name = _extract_pdf_name_from_unique_key(unique_key)
    configured_metadata = ((report_metadata or {}).get("pdf_file_metadata", {}) or {}).get(pdf_name, {})
    embedded_metadata = dict(data.get("file_metadata", {}) or {})
    merged_metadata = {
        "existing_benefit": False,
        "pdf_name": pdf_name,
    }
    for source in (configured_metadata, embedded_metadata):
        if not source:
            continue
        if "existing_benefit" in source and source.get("existing_benefit") is not None:
            merged_metadata["existing_benefit"] = bool(source.get("existing_benefit"))
    return merged_metadata


def _resolve_baseline_context(
    final_results: dict[str, Any],
    underwriter_overrides: dict[str, str] | None = None,
    report_metadata: dict[str, Any] | None = None,
) -> dict[str, str]:
    baseline_pdf = os.path.basename(str((report_metadata or {}).get("baseline_pdf") or "").strip())
    baseline_labels: list[str] = []
    for unique_key, data in final_results.items():
        file_metadata = _get_effective_file_metadata(unique_key, data, report_metadata)
        if not file_metadata.get("existing_benefit"):
            continue
        if not baseline_pdf:
            baseline_pdf = file_metadata.get("pdf_name", "")
        resolved_label = _get_effective_underwriter_label(unique_key, data, underwriter_overrides)
        if resolved_label and resolved_label not in baseline_labels:
            baseline_labels.append(resolved_label)

    baseline_underwriter = baseline_labels[0] if len(baseline_labels) == 1 else str((report_metadata or {}).get("baseline_underwriter") or "").strip()
    return {
        "baseline_pdf": baseline_pdf,
        "baseline_underwriter": baseline_underwriter,
    }


def _resolve_underwriter_override(unique_key: str, overrides: dict[str, str] | None) -> str:
    overrides = overrides or {}
    return str(
        overrides.get(unique_key)
        or overrides.get(_extract_pdf_name_from_unique_key(unique_key))
        or ""
    ).strip()


def _get_effective_underwriter_label(unique_key: str, data: dict[str, Any], underwriter_overrides: dict[str, str] | None = None) -> str:
    norm = data.get("normalized", {}) or {}
    return (
        _resolve_underwriter_override(unique_key, underwriter_overrides)
        or str(norm.get("0.1. Underwriter", "") or "").strip()
        or _extract_pdf_name_from_unique_key(unique_key)
    )


def _get_plan_display_name(unique_key: str, data: dict[str, Any], underwriter_override: str = "") -> str:
    norm = data.get("normalized", {})
    underwriter = str(underwriter_override or norm.get("0.1. Underwriter", "") or "").strip()
    plan_name = str(norm.get("0.2. Plan Name", "") or "").strip() or _extract_plan_name_from_unique_key(unique_key)
    if underwriter and plan_name:
        return f"{underwriter} – {plan_name}"
    return plan_name or unique_key


def _build_underwriter_category_assignments(
    final_results: dict,
    underwriter_overrides: dict[str, str] | None = None,
) -> dict[str, dict[str, Any]]:
    assignments: dict[str, dict[str, Any]] = {}
    ordered_keys, grouped = _group_plans_by_underwriter(final_results, underwriter_overrides=underwriter_overrides)

    for uw_key in ordered_keys:
        group = grouped.get(uw_key, {})
        display_underwriter = str(group.get("display_label", "") or "").strip()
        for category_order, (plan_name, _, unique_key, data) in enumerate(group.get("plans", [])):
            resolved_underwriter = _get_effective_underwriter_label(unique_key, data, underwriter_overrides)
            assignments[unique_key] = {
                "underwriter": resolved_underwriter or display_underwriter,
                "category_order": category_order,
                "category_label": _category_label_for_index(category_order),
                "source_plan_name": str(plan_name or "").strip() or _extract_plan_name_from_unique_key(unique_key),
            }

    return assignments


def _get_ordered_category_labels(benefit_plan_metadata: dict[str, dict[str, Any]]) -> list[str]:
    category_rows = []
    for metadata in benefit_plan_metadata.values():
        category_label = str(metadata.get("category_label", "") or "").strip()
        if not category_label:
            continue
        category_rows.append((int(metadata.get("category_order", 0) or 0), category_label))

    ordered_labels: list[str] = []
    for _, category_label in sorted(category_rows, key=lambda item: (item[0], item[1])):
        if category_label not in ordered_labels:
            ordered_labels.append(category_label)
    return ordered_labels


def _get_existing_underwriter_by_category(benefit_plan_metadata: dict[str, dict[str, Any]]) -> dict[str, str]:
    category_underwriters: dict[str, set[str]] = {}
    for metadata in benefit_plan_metadata.values():
        if not bool(metadata.get("existing_benefit")):
            continue
        underwriter = str(metadata.get("underwriter", "") or "").strip()
        category_label = str(metadata.get("category_label", "") or "").strip()
        if not underwriter:
            continue
        category_underwriters.setdefault(category_label, set()).add(underwriter)

    resolved: dict[str, str] = {}
    for category_label, underwriters in category_underwriters.items():
        if len(underwriters) == 1:
            resolved[category_label] = next(iter(underwriters))
    return resolved


def _build_relative_to_existing_highlight_input(
    row_values: dict[str, str],
    underwriters: list[str],
    baseline_underwriter: str,
) -> dict[str, Any] | None:
    baseline_label = str(baseline_underwriter or "").strip()
    if not baseline_label or baseline_label not in underwriters:
        return None

    baseline_value = _coerce_nonempty_summary_text(
        row_values.get(baseline_label, ""),
        max_length=SUMMARY_CELL_MAX_LENGTH,
    )
    if _is_blankish(baseline_value):
        return None

    candidates: list[dict[str, str]] = []
    for underwriter_label in underwriters:
        if underwriter_label == baseline_label:
            continue
        candidate_value = _coerce_nonempty_summary_text(
            row_values.get(underwriter_label, ""),
            max_length=SUMMARY_CELL_MAX_LENGTH,
        )
        if _is_blankish(candidate_value):
            continue
        candidates.append({
            "underwriter": underwriter_label,
            "value": candidate_value,
        })

    if not candidates:
        return None

    return {
        "existing_underwriter": baseline_label,
        "existing_value": baseline_value,
        "candidate_underwriters": [candidate["underwriter"] for candidate in candidates],
        "candidates": candidates,
    }


def _normalize_relative_highlight_labels(raw_labels: Any, allowed_labels: set[str]) -> list[str]:
    labels: list[str] = []
    if isinstance(raw_labels, list):
        source_values = raw_labels
    elif isinstance(raw_labels, str):
        source_values = [raw_labels]
    else:
        source_values = []

    for value in source_values:
        label = str(value or "").strip()
        if not label or label not in allowed_labels or label in labels:
            continue
        labels.append(label)
    return labels


def _group_plans_by_underwriter(final_results: dict, underwriter_overrides: dict[str, str] | None = None) -> tuple[list[str], dict[str, dict[str, Any]]]:
    grouped: dict[str, dict[str, Any]] = {}

    for unique_key, data in final_results.items():
        norm = data.get("normalized", {})
        raw_uw = _resolve_underwriter_override(unique_key, underwriter_overrides) or norm.get("0.1. Underwriter", "").strip() or _extract_pdf_name_from_unique_key(unique_key)
        raw_plan = norm.get("0.2. Plan Name", "").strip() or _extract_plan_name_from_unique_key(unique_key)
        uw_key = _canonical_underwriter_key(raw_uw) or raw_uw

        group = grouped.setdefault(
            uw_key,
            {"labels": [], "plans": []},
        )
        group["labels"].append(raw_uw)
        group["plans"].append((raw_plan, norm, unique_key, data))

    ordered_keys: list[str] = []
    for unique_key, data in final_results.items():
        norm = data.get("normalized", {})
        raw_uw = _resolve_underwriter_override(unique_key, underwriter_overrides) or norm.get("0.1. Underwriter", "").strip() or _extract_pdf_name_from_unique_key(unique_key)
        uw_key = _canonical_underwriter_key(raw_uw) or raw_uw
        if uw_key not in ordered_keys:
            ordered_keys.append(uw_key)

    for uw_key, group in grouped.items():
        group["display_label"] = _preferred_underwriter_label(group["labels"])

    return ordered_keys, grouped


def _strip_best_marker(value: Any) -> str:
    return str(value or "").replace("**[[BEST]]**", "").strip()


def _has_best_marker(value: Any) -> bool:
    return "**[[BEST]]**" in str(value or "")


def _has_multiple_underwriters(final_results: dict, underwriter_overrides: dict[str, str] | None = None) -> bool:
    underwriter_keys, _ = _group_plans_by_underwriter(final_results, underwriter_overrides=underwriter_overrides)
    return len(underwriter_keys) > 1


def _compute_best_labels(score_map: dict[str, float], enable_highlight: bool = True) -> set[str]:
    if not enable_highlight or len(score_map) < 2:
        return set()

    max_score = max(score_map.values()) if score_map else 0.0
    if max_score <= 0:
        return set()

    winners = {label for label, score in score_map.items() if score == max_score}
    return set() if len(winners) == len(score_map) else winners


def _summarize_plan_fields_for_sheet(plan_fields: dict, use_llm_summaries: bool = True) -> dict:
    if use_llm_summaries:
        summarized = summarize_dict_values(plan_fields)
        result = dict(summarized)
    else:
        result = {
            key: (value.strip() if isinstance(value, str) else value)
            for key, value in plan_fields.items()
        }

    for field_key, raw_value in plan_fields.items():
        if field_key in {"0.1. Underwriter", "0.2. Plan Name"}:
            continue
        if _has_best_marker(raw_value) and str(result.get(field_key, "")).strip():
            result[field_key] = f"**[[BEST]]** {result[field_key]}"

    return result


def _summarize_underwriter_plan_fields_for_sheet(
    underwriter_label: str,
    plan_entries: list[dict[str, Any]],
    use_llm_summaries: bool = True,
) -> dict[str, dict[str, Any]]:
    if not plan_entries:
        return {}

    normalized_entries: list[dict[str, Any]] = []
    for entry in plan_entries:
        display_name = str(entry.get("display_name", "") or "").strip()
        if not display_name:
            continue
        normalized_entries.append(
            {
                "display_name": display_name,
                "plan_name": str(entry.get("plan_name", "") or "").strip(),
                "plan_fields": dict(entry.get("plan_fields", {}) or {}),
            }
        )

    if not normalized_entries:
        return {}

    cache_key = json.dumps(
        {
            "underwriter_label": underwriter_label,
            "use_llm_summaries": use_llm_summaries,
            "plans": normalized_entries,
        },
        sort_keys=True,
        ensure_ascii=False,
        default=str,
    )
    if cache_key in _underwriter_summary_cache:
        return json.loads(json.dumps(_underwriter_summary_cache[cache_key]))

    passthrough_keys = {"0.1. Underwriter", "0.2. Plan Name"}
    field_order: list[str] = []
    for entry in normalized_entries:
        for field_key in entry["plan_fields"].keys():
            if field_key not in field_order:
                field_order.append(field_key)

    comparable_keys = [
        key for key in field_order
        if key not in passthrough_keys and any(str(entry["plan_fields"].get(key, "") or "").strip() for entry in normalized_entries)
    ]
    summarized_values: dict[str, dict[str, str]] = {
        entry["display_name"]: {
            field_key: (raw_value.strip() if isinstance(raw_value, str) else raw_value)
            for field_key, raw_value in entry["plan_fields"].items()
            if isinstance(raw_value, str)
        }
        for entry in normalized_entries
    }

    if use_llm_summaries and comparable_keys:
        system_prompt = """You summarize insurance benefit fields for an Excel comparison sheet.

<output_shape>
- Return JSON only.
- Schema:
  {
    "plans": {
      "<display_name>": {
        "<field_key>": string
      }
    }
  }
- Return only the exact plan display names and field keys provided in the input.
</output_shape>

<summary_rules>
- Treat all plans as belonging to the same underwriter and preserve plan-to-plan differences.
- For `0.1. Underwriter` and `0.2. Plan Name`, return the original value unchanged if they appear.
- For each other benefit field, produce one concise summary line using this priority: `<amount/value/main coverage phrase>` first, then only material financial or condition detail if needed.
- Use `COVERED` or `NOT COVERED` only when required to preserve meaning. If the field is already clear from the value or coverage phrase alone, do not add `COVERED`.
- If `NOT COVERED` is the key meaning, state `NOT COVERED` clearly.
- Always include the most relevant amount or value when present.
- Keep this customer-facing and compact.
- Keep only the top 1 amount/value clause, at most 1 key financial clause (deductible, coinsurance, or copay), and at most 1 short restriction or subjectivity.
- Include deductible, coinsurance, copay, waiting period, subjectivities, restrictions, approval requirements, and other material conditions only when present.
- Prefer outputs like `BHD 754,000`, `Worldwide excluding USA & CANADA`, `Private room`, `20% copay`, or `NOT COVERED` over verbose sentence fragments.
- Do not add filler wording, examples, or explanatory prose.
- Do not guess, add facts, or change the business meaning.
- Use only the provided quote text. Do not infer, normalize away, or supplement missing details.
- Keep blanks blank.
- Keep each summary concise and comparison-friendly.
</summary_rules>"""

        for field_chunk in _chunk_sequence(comparable_keys, PLAN_FIELD_SUMMARY_CHUNK_SIZE):
            chunk_payload = {
                "underwriter": underwriter_label,
                "plans": [
                    {
                        "display_name": entry["display_name"],
                        "plan_name": entry["plan_name"],
                        "fields": {
                            field_key: _strip_best_marker(entry["plan_fields"].get(field_key, ""))
                            for field_key in field_chunk
                        },
                    }
                    for entry in normalized_entries
                ],
            }
            try:
                chunk_user_prompt = f"""Summarize these insurance fields for the Benefit_Summary sheet.

INPUT JSON:
{json.dumps(chunk_payload, indent=2, ensure_ascii=False)}

Return JSON only."""
                content, usage = llm_chat(
                    system_prompt,
                    chunk_user_prompt,
                    json_mode=True,
                    log_label="benefit_dict_summary",
                    task_class="benefit_summary",
                    chunk_size=len(field_chunk),
                    log_context={
                        "underwriter": underwriter_label,
                        "plan_count": len(normalized_entries),
                        "field_count": len(field_chunk),
                        "field_keys": field_chunk,
                    },
                )
                tracker.record(_current_pdf, usage)
                chunk_summary = safe_json_loads(content, "benefit dict summary")
            except Exception as exc:
                print(f"  ⚠️  Underwriter benefit summarization chunk failed: {exc} — leaving chunk blank")
                chunk_summary = {}

            raw_plans = chunk_summary.get("plans", chunk_summary) if isinstance(chunk_summary, dict) else {}
            for entry in normalized_entries:
                display_name = entry["display_name"]
                raw_plan_result = raw_plans.get(display_name, {}) if isinstance(raw_plans, dict) else {}
                for field_key in field_chunk:
                    raw_value = str(entry["plan_fields"].get(field_key, "") or "")
                    if not raw_value.strip():
                        summarized_values.setdefault(display_name, {})[field_key] = ""
                        continue
                    candidate: Any = ""
                    if isinstance(raw_plan_result, dict):
                        candidate = raw_plan_result.get(field_key, "")
                    summarized_values.setdefault(display_name, {})[field_key] = _format_benefit_summary_text(
                        candidate,
                        raw_value,
                        max_length=SUMMARY_CELL_MAX_LENGTH,
                    )

    result: dict[str, dict[str, Any]] = {}
    for entry in normalized_entries:
        display_name = entry["display_name"]
        plan_fields = entry["plan_fields"]
        plan_result: dict[str, Any] = {}
        for field_key, raw_value in plan_fields.items():
            if not isinstance(raw_value, str):
                plan_result[field_key] = raw_value
                continue
            if field_key in passthrough_keys:
                plan_result[field_key] = raw_value.strip()
                continue

            formatted_value = str(summarized_values.get(display_name, {}).get(field_key, "") or "")
            if _has_best_marker(raw_value) and formatted_value.strip():
                formatted_value = f"**[[BEST]]** {formatted_value}"
            plan_result[field_key] = formatted_value

        result[display_name] = plan_result

    _underwriter_summary_cache[cache_key] = json.loads(json.dumps(result))
    return result


def _chunk_sequence(values: list[str], size: int) -> list[list[str]]:
    if size <= 0:
        return [values]
    return [values[idx: idx + size] for idx in range(0, len(values), size)]


def _smart_title_case(text: str) -> str:
    words = str(text or "").split()
    titled_words: list[str] = []
    for word in words:
        bare_word = re.sub(r"[^A-Za-z]", "", word)
        if bare_word.isupper() and 1 < len(bare_word) <= 4:
            titled_words.append(word)
        elif bare_word.isupper():
            titled_words.append(word.title())
        else:
            titled_words.append(word)
    return " ".join(titled_words)


def _format_raw_cell_text(value: Any) -> str:
    text = str(value or "")
    if not text.strip():
        return ""

    normalized_text = text.replace("\r\n", "\n").replace("\r", "\n")
    formatted_lines: list[str] = []
    for raw_line in normalized_text.split("\n"):
        compact_line = re.sub(r"[ \t]+", " ", str(raw_line or "")).strip()
        if compact_line:
            formatted_lines.append(compact_line)

    return "\n".join(formatted_lines)


def _apply_rich_text_emphasis(cell):
    cell_value = _format_raw_cell_text(cell.value)
    if not cell_value:
        cell.value = ""
        return

    # Excel is stricter than openpyxl about some inline rich-text payloads.
    # Store plain text here to avoid workbook repair prompts on open.
    cell.value = cell_value


def _split_cat_prefix(line: str) -> tuple[str, str]:
    text = str(line or "")
    match = re.match(r"^\s*((?:CAT|Cat)\s+[A-Z])\s*(:\s*.*)?$", text)
    if not match:
        return "", text
    prefix = match.group(1)
    suffix = match.group(2) or ""
    return prefix, suffix


def _apply_insurer_summary_rich_text(cell):
    cell_value = _format_raw_cell_text(cell.value)
    cell.value = cell_value if cell_value else ""


def _normalize_insurer_summary_line_highlights(
    line_highlights: Any,
) -> dict[str, list[dict[str, str]]]:
    normalized: dict[str, list[dict[str, str]]] = {}
    if not isinstance(line_highlights, dict):
        return normalized

    for underwriter, entries in line_highlights.items():
        cleaned_underwriter = str(underwriter or "").strip()
        if not cleaned_underwriter:
            continue
        normalized_entries: list[dict[str, str]] = []
        if isinstance(entries, list):
            for entry in entries:
                if not isinstance(entry, dict):
                    continue
                normalized_entries.append(
                    {
                        "category_label": str(entry.get("category_label", "") or "").strip(),
                        "status": str(entry.get("status", "") or "").strip().lower(),
                    }
                )
        normalized[cleaned_underwriter] = normalized_entries
    return normalized


def _build_insurer_summary_line_render_entries(
    text: Any,
    line_highlights: list[dict[str, str]] | None = None,
) -> list[dict[str, str]]:
    formatted_text = _format_raw_cell_text(text)
    if not formatted_text:
        return []

    status_by_category: dict[str, str] = {}
    fallback_status = ""
    for entry in list(line_highlights or []):
        category_label = str(entry.get("category_label", "") or "").strip()
        status = str(entry.get("status", "") or "").strip().lower()
        normalized_category = _normalize_match_text(category_label)
        if normalized_category:
            status_by_category[normalized_category] = status
        elif status:
            fallback_status = status

    rendered_lines: list[dict[str, str]] = []
    for raw_line in formatted_text.split("\n"):
        line = str(raw_line or "").strip()
        if not line:
            continue
        prefix, _ = _split_cat_prefix(line)
        normalized_category = _normalize_match_text(prefix)
        rendered_lines.append(
            {
                "text": line,
                "category_label": _format_category_display_label(prefix),
                "status": status_by_category.get(normalized_category, fallback_status if not normalized_category else ""),
            }
        )
    return rendered_lines


def _build_excel_insurer_summary_display_text(
    text: Any,
    line_highlights: list[dict[str, str]] | None = None,
)-> str:
    del line_highlights
    return _format_raw_cell_text(text)


def _build_excel_insurer_summary_rich_runs(
    text: Any,
    line_highlights: list[dict[str, str]] | None = None,
) -> list[dict[str, Any]]:
    rendered_lines = _build_insurer_summary_line_render_entries(text, line_highlights)
    if not rendered_lines:
        return []

    runs: list[dict[str, Any]] = []
    color_by_status = {
        "best": "FF006100",
        "worst": "FF9C0006",
    }
    for idx, entry in enumerate(rendered_lines):
        line_text = entry.get("text", "")
        status = entry.get("status", "")
        color = color_by_status.get(status, "")
        prefix, suffix = _split_cat_prefix(line_text)
        if prefix:
            runs.append({"text": prefix, "bold": True, "color": color})
            if suffix:
                runs.append({"text": suffix, "bold": False, "color": color})
        else:
            runs.append({"text": line_text, "bold": False, "color": color})
        if idx < len(rendered_lines) - 1:
            runs.append({"text": "\n", "bold": False, "color": ""})
    return runs


def _xlsx_sheet_path_for_name(xlsx_path: str | Path, sheet_name: str) -> str:
    import zipfile

    workbook_ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    office_rel_ns = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    package_rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    with zipfile.ZipFile(xlsx_path, "r") as archive:
        workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
        rels_root = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))

    target_by_id = {
        rel.attrib.get("Id", ""): rel.attrib.get("Target", "")
        for rel in rels_root.findall(f"{{{package_rel_ns}}}Relationship")
    }

    sheets_parent = workbook_root.find(f"{{{workbook_ns}}}sheets")
    if sheets_parent is None:
        raise ValueError("Workbook sheets section not found")

    rel_attr = f"{{{office_rel_ns}}}id"
    for sheet in sheets_parent.findall(f"{{{workbook_ns}}}sheet"):
        if sheet.attrib.get("name") != sheet_name:
            continue
        rel_id = sheet.attrib.get(rel_attr, "")
        target = target_by_id.get(rel_id, "")
        if not target:
            break
        normalized_target = target.lstrip("/")
        return normalized_target if normalized_target.startswith("xl/") else f"xl/{normalized_target}"

    raise ValueError(f"Worksheet XML target not found for sheet: {sheet_name}")


def _apply_ooxml_rich_text_to_sheet_cells(
    xlsx_path: str | Path,
    sheet_name: str,
    cell_runs_map: dict[str, list[dict[str, Any]]],
):
    import tempfile
    import zipfile

    if not cell_runs_map:
        return

    workbook_ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    xml_space_ns = "http://www.w3.org/XML/1998/namespace"

    sheet_path = _xlsx_sheet_path_for_name(xlsx_path, sheet_name)
    ET.register_namespace("", workbook_ns)

    with zipfile.ZipFile(xlsx_path, "r") as source_zip:
        original_sheet_xml = source_zip.read(sheet_path)
        original_entries = [
            (zip_info, source_zip.read(zip_info.filename))
            for zip_info in source_zip.infolist()
        ]

    sheet_root = ET.fromstring(original_sheet_xml)
    cell_by_ref = {
        cell.attrib.get("r", ""): cell
        for cell in sheet_root.findall(f".//{{{workbook_ns}}}c")
    }

    for cell_ref, runs in cell_runs_map.items():
        cell = cell_by_ref.get(cell_ref)
        if cell is None:
            continue
        for child in list(cell):
            cell.remove(child)
        cell.set("t", "inlineStr")

        inline_string = ET.SubElement(cell, f"{{{workbook_ns}}}is")
        for run_spec in runs:
            text = str(run_spec.get("text", "") or "")
            if not text:
                continue
            run = ET.SubElement(inline_string, f"{{{workbook_ns}}}r")
            run_properties = ET.SubElement(run, f"{{{workbook_ns}}}rPr")
            font = ET.SubElement(run_properties, f"{{{workbook_ns}}}rFont")
            font.set("val", "Arial")
            if run_spec.get("bold"):
                ET.SubElement(run_properties, f"{{{workbook_ns}}}b")
            color = str(run_spec.get("color", "") or "").strip()
            if color:
                color_element = ET.SubElement(run_properties, f"{{{workbook_ns}}}color")
                color_element.set("rgb", color)
            text_element = ET.SubElement(run, f"{{{workbook_ns}}}t")
            if text != text.strip() or "\n" in text or text.isspace():
                text_element.set(f"{{{xml_space_ns}}}space", "preserve")
            text_element.text = text

    modified_sheet_xml = ET.tostring(sheet_root, encoding="utf-8", xml_declaration=False)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
        temp_path = Path(temp_file.name)

    try:
        with zipfile.ZipFile(temp_path, "w") as target_zip:
            for zip_info, entry_data in original_entries:
                data = modified_sheet_xml if zip_info.filename == sheet_path else entry_data
                target_zip.writestr(zip_info, data)
        os.replace(temp_path, xlsx_path)
    finally:
        if temp_path.exists():
            try:
                temp_path.unlink()
            except OSError:
                pass


def _estimate_excel_wrapped_line_count(value: Any, column_width: float | int | None) -> int:
    text = _format_raw_cell_text(value)
    if not text:
        return 1

    usable_width = max(int(float(column_width or 0)) - 2, 8)
    visual_lines = 0
    for raw_line in text.split("\n"):
        line = str(raw_line or "")
        if not line:
            visual_lines += 1
            continue
        visual_lines += max(1, (len(line) - 1) // usable_width + 1)
    return max(1, visual_lines)


def _apply_word_insurer_summary_line_style(run, status: str):
    cleaned_status = str(status or "").strip().lower()
    if cleaned_status == "best":
        run.bold = True
        run.font.color.rgb = RGBColor(0, 97, 0)
    elif cleaned_status == "worst":
        run.font.color.rgb = RGBColor(156, 0, 6)


def _write_word_insurer_summary_value(cell, text: str, line_highlights: list[dict[str, str]] | None = None):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    rendered_lines = _build_insurer_summary_line_render_entries(text, line_highlights)
    if not rendered_lines:
        return

    for idx, entry in enumerate(rendered_lines):
        line = entry.get("text", "")
        status = entry.get("status", "")
        prefix, suffix = _split_cat_prefix(line)
        if prefix:
            run = paragraph.add_run(prefix)
            run.bold = True
            _apply_word_insurer_summary_line_style(run, status)
            if suffix:
                suffix_run = paragraph.add_run(suffix)
                _apply_word_insurer_summary_line_style(suffix_run, status)
        else:
            run = paragraph.add_run(line)
            _apply_word_insurer_summary_line_style(run, status)
        if idx < len(rendered_lines) - 1:
            paragraph.add_run("\n")


def _build_insurer_chunk_payload_from_benefit_summary(
    benefit_summary: dict[str, dict[str, Any]],
    benefit_plan_metadata: dict[str, dict[str, Any]],
    underwriters: list[str],
    field_keys: list[str],
    category_label: str | None = None,
) -> dict[str, Any]:
    grouped_entries: dict[str, list[dict[str, Any]]] = {underwriter: [] for underwriter in underwriters}
    for display_name, metadata in benefit_plan_metadata.items():
        underwriter = str(metadata.get("underwriter", "") or "").strip()
        if underwriter not in grouped_entries:
            continue
        if category_label and str(metadata.get("category_label", "") or "").strip() != str(category_label or "").strip():
            continue
        grouped_entries[underwriter].append(
            {
                "display_name": display_name,
                "plan_name": str(metadata.get("plan_name", "") or "").strip(),
                "order": int(metadata.get("order", 0) or 0),
            }
        )

    for underwriter in grouped_entries:
        grouped_entries[underwriter].sort(key=lambda item: item["order"])

    payload: dict[str, Any] = {}
    for field_key in field_keys:
        field_payload: dict[str, Any] = {}
        for underwriter in underwriters:
            plan_items = []
            for item in grouped_entries.get(underwriter, []):
                raw_val = _strip_best_marker(benefit_summary.get(item["display_name"], {}).get(field_key, ""))
                if raw_val:
                    plan_items.append({
                        "plan_name": item["plan_name"],
                        "value": raw_val,
                    })
            field_payload[underwriter] = plan_items
        payload[field_key] = field_payload
    return payload


def _strip_single_plan_prefixes(
    chunk_summary: dict[str, Any],
    chunk_payload: dict[str, Any],
) -> dict[str, Any]:
    cleaned_summary: dict[str, Any] = {}
    for field_key, uw_values in (chunk_summary or {}).items():
        cleaned_summary[field_key] = {}
        field_payload = chunk_payload.get(field_key, {}) if isinstance(chunk_payload, dict) else {}
        for uw_label, value in (uw_values or {}).items():
            cleaned_value = str(value or "").strip()
            plan_items = field_payload.get(uw_label, []) if isinstance(field_payload, dict) else []
            if len(plan_items) == 1 and cleaned_value:
                plan_name = str(plan_items[0].get("plan_name", "") or "").strip()
                cleaned_lines = [
                    _strip_prefixed_label(str(line or "").strip(), plan_name)
                    for line in cleaned_value.split("\n")
                    if str(line or "").strip()
                ]
                cleaned_value = "\n".join(cleaned_lines).strip()
            cleaned_summary[field_key][uw_label] = cleaned_value
    return cleaned_summary


def _build_insurer_chunk_summary_from_benefit_summary(
    chunk_payload: dict[str, Any],
    underwriters: list[str],
) -> dict[str, dict[str, str]]:
    chunk_summary: dict[str, dict[str, str]] = {}
    for field_key, uw_values in chunk_payload.items():
        chunk_summary[field_key] = {}
        for uw_label in underwriters:
            plan_items = uw_values.get(uw_label, []) if isinstance(uw_values, dict) else []
            if not plan_items:
                chunk_summary[field_key][uw_label] = ""
                continue

            if len(plan_items) == 1:
                chunk_summary[field_key][uw_label] = _coerce_nonempty_summary_text(
                    plan_items[0].get("value", ""),
                    max_length=SUMMARY_CELL_MAX_LENGTH,
                )
                continue

            lines: list[str] = []
            for item in plan_items:
                plan_name = str(item.get("plan_name", "") or "").strip()
                value = _coerce_nonempty_summary_text(item.get("value", ""), max_length=SUMMARY_CELL_MAX_LENGTH)
                if not value:
                    continue
                if plan_name:
                    lines.append(f"{plan_name} : {value}")
                else:
                    lines.append(value)
            chunk_summary[field_key][uw_label] = "\n".join(lines).strip()
    return chunk_summary


def _build_insurer_chunk_payload(
    uw_groups: dict[str, dict[str, Any]],
    underwriter_keys: list[str],
    field_keys: list[str],
) -> dict[str, Any]:
    payload: dict[str, Any] = {}
    for field_key in field_keys:
        field_payload: dict[str, Any] = {}
        for uw_key in underwriter_keys:
            group = uw_groups[uw_key]
            uw_label = group["display_label"]
            plan_items = []
            for plan_name, norm, _, _ in group["plans"]:
                raw_val = _strip_best_marker(norm.get(field_key, ""))
                if raw_val:
                    plan_items.append({"plan_name": plan_name, "value": raw_val})
            field_payload[uw_label] = plan_items
        payload[field_key] = field_payload
    return payload


def _summarize_insurer_chunk(chunk_payload: dict[str, Any], underwriters: list[str]) -> dict[str, dict[str, str]]:
    if not chunk_payload:
        return {}

    system_prompt = """You summarize insurance benefits for an underwriter comparison table.

<output_shape>
- Return JSON only.
- Return the exact same top-level benefit keys provided in the input.
- For each benefit key, return an object with the exact underwriter names provided in the input.
</output_shape>

<summary_rules>
- Summarize each underwriter's benefit position using the provided Benefit Summary plan values only.
- Keep the wording concise, comparison-friendly, and business-accurate.
- For each underwriter, return multi-line text with one line per plan in the same order as the input.
- Each line must use this format: <plan_name>: <main amount/value/coverage phrase> with only essential extra detail appended when needed.
- Use `COVERED` or `NOT COVERED` only when required to preserve meaning. If the value or coverage phrase already conveys the benefit clearly, omit `COVERED`.
- If a plan is not covered, state `NOT COVERED` clearly.
- Always include the most relevant amount or value when present.
- Keep this customer-facing and compact.
- Keep only the top 1 amount/value clause, at most 1 key financial clause (deductible, coinsurance, or copay), and at most 1 short restriction or subjectivity per plan.
- Compress long wording into short business phrases such as `CoC required`, `6-month waiting period`, or `Prior approval required` when the source supports that wording.
- Include deductible, coinsurance, copay, waiting period, subjectivities, restrictions, approval requirements, and other material conditions only when present.
- Prefer outputs like `CAT A: BHD 754,000`, `CAT B: Bahrain, GCC, ME, ISC & SE ASIA`, `CAT C: Bahrain Only`, or `CAT B: NOT COVERED` over verbose sentence fragments.
- Do not add filler wording, examples, or explanatory prose.
- If an underwriter has multiple plans with meaningful differences, preserve those differences briefly in the summary.
- Do not invent facts, make assumptions, or quote anything not present in the provided plan values.
- Keep blank results blank.
</summary_rules>"""

    user_prompt = f"""Summarize these insurer comparison rows for the Insurer_Summary sheet.

UNDERWRITERS:
{json.dumps(underwriters, indent=2, ensure_ascii=False)}

INPUT JSON:
{json.dumps(chunk_payload, indent=2, ensure_ascii=False)}

Return JSON only."""

    content, usage = llm_chat(
        system_prompt,
        user_prompt,
        json_mode=True,
        log_label="insurer_summary_chunk",
        task_class="insurer_summary",
        chunk_size=len(chunk_payload),
        log_context={"underwriters": underwriters, "field_keys": list(chunk_payload.keys())},
    )
    tracker.record(_current_pdf, usage)
    result = safe_json_loads(content, "insurer summary chunk")

    normalized: dict[str, dict[str, str]] = {}
    for field_key, uw_values in chunk_payload.items():
        field_result = result.get(field_key, {}) if isinstance(result, dict) else {}
        normalized[field_key] = {}
        for uw_label in underwriters:
            value: Any = ""
            if isinstance(field_result, dict):
                value = field_result.get(uw_label, "")
            normalized[field_key][uw_label] = _format_planwise_summary(uw_values.get(uw_label, []), value, max_length=SUMMARY_CELL_MAX_LENGTH)
    return normalized


def _analyze_insurer_highlight_chunk(
    chunk_summary: dict[str, dict[str, str]],
    underwriters: list[str],
    baseline_underwriter: str = "",
) -> dict[str, dict[str, Any]]:
    if not chunk_summary:
        return {}

    baseline_label = str(baseline_underwriter or "").strip()
    if not baseline_label or baseline_label not in underwriters:
        return {
            field_key: {"best_underwriters": [], "worst_underwriters": []}
            for field_key in chunk_summary.keys()
        }

    relative_input: dict[str, dict[str, Any]] = {}
    for field_key, row_values in chunk_summary.items():
        prepared_row = _build_relative_to_existing_highlight_input(row_values, underwriters, baseline_label)
        if prepared_row:
            relative_input[field_key] = prepared_row

    if not relative_input:
        return {
            field_key: {"best_underwriters": [], "worst_underwriters": []}
            for field_key in chunk_summary.keys()
        }

    system_prompt = """You analyze insurer comparison rows relative to an existing underwriter.

<output_shape>
- Return JSON only.
- Return the exact same top-level benefit keys provided in the input.
- For each benefit key return:
  {
    "better_than_existing": string[],
    "worse_than_existing": string[],
    "reason": string
  }
</output_shape>

<decision_rules>
- Each input row is an independent category-level comparison row.
- Compare each candidate underwriter only against the provided existing underwriter for that same row.
- Do not rank candidates against each other and do not choose a global best or worst across all underwriters.
- The existing underwriter must never appear in either output list.
- Add a candidate to `better_than_existing` only when the evidence clearly supports that the candidate is better than the existing value.
- Add a candidate to `worse_than_existing` only when the evidence clearly supports that the candidate is worse than the existing value.
- If the comparison is tied, ambiguous, weak, or not meaningfully comparable, leave that candidate out of both lists.
- Prefer no result over a guessed result.
</decision_rules>"""

    user_prompt = f"""Compare these insurer comparison rows against the existing underwriter only.

UNDERWRITERS:
{json.dumps(underwriters, indent=2, ensure_ascii=False)}

EXISTING UNDERWRITER:
{json.dumps(baseline_label, ensure_ascii=False)}

INPUT JSON:
{json.dumps(relative_input, indent=2, ensure_ascii=False)}

Return JSON only."""

    content, usage = llm_chat(
        system_prompt,
        user_prompt,
        json_mode=True,
        log_label="insurer_highlighting_chunk",
        task_class="insurer_highlighting",
        chunk_size=len(chunk_summary),
        log_context={"underwriters": underwriters, "field_keys": list(chunk_summary.keys())},
    )
    tracker.record(_current_pdf, usage)
    result = safe_json_loads(content, "insurer highlighting chunk")

    normalized: dict[str, dict[str, Any]] = {}
    for field_key in chunk_summary.keys():
        prepared_row = relative_input.get(field_key)
        if not prepared_row:
            normalized[field_key] = {
                "best_underwriters": [],
                "worst_underwriters": [],
            }
            continue

        field_result = result.get(field_key, {}) if isinstance(result, dict) else {}
        allowed_labels = set(prepared_row.get("candidate_underwriters", []))
        better_labels = _normalize_relative_highlight_labels(
            field_result.get("better_than_existing", []),
            allowed_labels,
        )
        worse_labels = _normalize_relative_highlight_labels(
            field_result.get("worse_than_existing", []),
            allowed_labels,
        )
        overlapping_labels = set(better_labels) & set(worse_labels)
        if overlapping_labels:
            better_labels = [label for label in better_labels if label not in overlapping_labels]
            worse_labels = [label for label in worse_labels if label not in overlapping_labels]
        normalized[field_key] = {
            "best_underwriters": better_labels,
            "worst_underwriters": worse_labels,
        }
    return normalized


def _build_insurer_summary_view(
    final_results: dict,
    keys: list[str],
    use_llm_summaries: bool = True,
    underwriter_overrides: dict[str, str] | None = None,
    benefit_summary: dict[str, dict[str, Any]] | None = None,
    benefit_plan_metadata: dict[str, dict[str, Any]] | None = None,
    baseline_underwriter: str | None = None,
) -> dict[str, Any]:
    underwriter_keys, uw_groups = _group_plans_by_underwriter(final_results, underwriter_overrides=underwriter_overrides)
    underwriters = [uw_groups[uw_key]["display_label"] for uw_key in underwriter_keys]
    multiple_underwriters = len(underwriters) > 1

    uw_score_stats = {}
    scored_underwriters: dict[str, float] = {}
    for uw_key in underwriter_keys:
        group = uw_groups[uw_key]
        uw_label = group["display_label"]
        uw_total_scores = []
        for _, _, _, data in group["plans"]:
            total_score = data.get("total_score", 0)
            if isinstance(total_score, (int, float)):
                uw_total_scores.append(float(total_score))

        avg_score = round(sum(uw_total_scores) / len(uw_total_scores), 1) if uw_total_scores else ""
        max_score = round(max(uw_total_scores), 1) if uw_total_scores else ""
        uw_score_stats[uw_label] = {
            "avg": avg_score,
            "max": max_score,
            "count": len(group["plans"]),
        }
        if uw_total_scores:
            scored_underwriters[uw_label] = float(avg_score)

    effective_benefit_summary = benefit_summary or {}
    effective_benefit_plan_metadata = benefit_plan_metadata or {}
    existing_underwriter_by_category = _get_existing_underwriter_by_category(effective_benefit_plan_metadata)
    comparable_keys = [key for key in keys if key not in {"0.1. Underwriter", "0.2. Plan Name"}]
    category_labels = _get_ordered_category_labels(effective_benefit_plan_metadata)
    category_source_lookup = {
        (
            str(metadata.get("underwriter", "") or "").strip(),
            str(metadata.get("category_label", "") or "").strip(),
        ): str(metadata.get("source_group_label", "") or "").strip()
        for metadata in effective_benefit_plan_metadata.values()
        if str(metadata.get("underwriter", "") or "").strip() and str(metadata.get("category_label", "") or "").strip()
    }
    summary_row_keys = list(comparable_keys)
    if category_labels:
        summary_row_keys = [
            _make_category_scoped_field_key(field_key, category_label)
            for field_key in comparable_keys
            for category_label in category_labels
        ]

    summary = {field_key: {uw: "" for uw in underwriters} for field_key in summary_row_keys}
    best_uw: dict[str, set[str]] = {field_key: set() for field_key in summary_row_keys}
    worst_uw: dict[str, set[str]] = {field_key: set() for field_key in summary_row_keys}

    if multiple_underwriters and comparable_keys:
        category_scope = category_labels or [""]
        for category_label in category_scope:
            category_baseline_underwriter = ""
            if category_labels:
                category_baseline_underwriter = str(existing_underwriter_by_category.get(str(category_label or "").strip(), "") or "").strip()
            else:
                category_baseline_underwriter = str(baseline_underwriter or "").strip()
            for field_chunk in _chunk_sequence(comparable_keys, INSURER_COMPARISON_CHUNK_SIZE):
                if effective_benefit_summary and effective_benefit_plan_metadata:
                    chunk_payload = _build_insurer_chunk_payload_from_benefit_summary(
                        effective_benefit_summary,
                        effective_benefit_plan_metadata,
                        underwriters,
                        field_chunk,
                        category_label=category_label or None,
                    )
                else:
                    chunk_payload = _build_insurer_chunk_payload(uw_groups, underwriter_keys, field_chunk)

                if not any(plan_items for field_values in chunk_payload.values() for plan_items in field_values.values()):
                    continue

                if effective_benefit_summary and effective_benefit_plan_metadata:
                    chunk_summary = _build_insurer_chunk_summary_from_benefit_summary(chunk_payload, underwriters)
                else:
                    try:
                        chunk_summary = _summarize_insurer_chunk(chunk_payload, underwriters)
                    except Exception as exc:
                        print(f"  ⚠️  Insurer summary chunk failed: {exc} — using raw joined values")
                        chunk_summary = _build_insurer_chunk_summary_from_benefit_summary(chunk_payload, underwriters)

                for field_key, uw_values in chunk_summary.items():
                    summary_key = _make_category_scoped_field_key(field_key, category_label)
                    if summary_key not in summary:
                        continue
                    for uw_label, value in uw_values.items():
                        summary[summary_key][uw_label] = value

                try:
                    chunk_highlights = _analyze_insurer_highlight_chunk(
                        chunk_summary,
                        underwriters,
                        baseline_underwriter=category_baseline_underwriter,
                    )
                except Exception as exc:
                    print(f"  ⚠️  Insurer highlighting chunk failed: {exc} — leaving chunk unhighlighted")
                    chunk_highlights = {}

                for field_key, highlight in chunk_highlights.items():
                    summary_key = _make_category_scoped_field_key(field_key, category_label)
                    if summary_key not in best_uw:
                        continue
                    best_labels = highlight.get("best_underwriters") if isinstance(highlight.get("best_underwriters"), list) else None
                    worst_labels = highlight.get("worst_underwriters") if isinstance(highlight.get("worst_underwriters"), list) else None
                    if best_labels is None:
                        best_label = str(highlight.get("best_underwriter", "") or "").strip()
                        best_labels = [best_label] if best_label else []
                    if worst_labels is None:
                        worst_label = str(highlight.get("worst_underwriter", "") or "").strip()
                        worst_labels = [worst_label] if worst_label else []
                    for best_label in best_labels:
                        if best_label:
                            best_uw[summary_key].add(best_label)
                    for worst_label in worst_labels:
                        if worst_label:
                            worst_uw[summary_key].add(worst_label)

    summary, best_uw, worst_uw, summary_row_keys, line_highlights = _collapse_category_scoped_summary(
        summary,
        best_uw,
        worst_uw,
        underwriters,
        comparable_keys,
        category_labels,
        category_source_lookup=category_source_lookup,
    )

    best_avg_score = max(scored_underwriters.values(), default=0.0)
    best_avg_uws = _compute_best_labels(
        scored_underwriters,
        enable_highlight=multiple_underwriters and bool(scored_underwriters) and best_avg_score > 0,
    )

    return {
        "underwriters": underwriters,
        "summary": summary,
        "row_keys": summary_row_keys,
        "category_labels": category_labels,
        "best_uw": best_uw,
        "worst_uw": worst_uw,
        "line_highlights": line_highlights,
        "uw_score_stats": uw_score_stats,
        "best_avg_uws": best_avg_uws,
        "multiple_underwriters": multiple_underwriters,
        "baseline_underwriter": str(baseline_underwriter or "").strip(),
    }


def build_report_payload(final_results: dict, keys: list[str], use_llm_summaries: bool = True, report_metadata: dict[str, Any] | None = None) -> dict[str, Any]:
    def clean_report_value(value: Any):
        if isinstance(value, (dict, list)):
            return clean_excel_string(_coerce_summary_candidate(value))
        if isinstance(value, str):
            return clean_excel_string(value)
        return value

    def clean_dict_values(d):
        return {k: clean_report_value(v) for k, v in d.items()}

    report_metadata = _normalize_report_metadata(report_metadata)
    keys = list(keys or [])
    for field_key in _report_ordered_keys():
        if field_key not in keys:
            keys.append(field_key)
    underwriter_overrides = report_metadata.get("pdf_underwriter_overrides", {})
    report_metadata.update(_resolve_baseline_context(final_results, underwriter_overrides, report_metadata))

    plan_category_assignments = _build_underwriter_category_assignments(
        final_results,
        underwriter_overrides=underwriter_overrides,
    )

    def resolve_plan_display_name(unique_key: str, data: dict[str, Any]) -> str:
        category_info = plan_category_assignments.get(unique_key, {})
        category_label = str(category_info.get("category_label", "") or "").strip()
        underwriter_label = str(
            category_info.get("underwriter", "")
            or _resolve_underwriter_override(unique_key, underwriter_overrides)
            or data.get("normalized", {}).get("0.1. Underwriter", "")
            or ""
        ).strip()
        if underwriter_label and category_label:
            return f"{underwriter_label} – {category_label}"
        return clean_excel_string(_get_plan_display_name(unique_key, data, underwriter_override=underwriter_label))

    display_names = {k: clean_excel_string(resolve_plan_display_name(k, v)) for k, v in final_results.items()}
    word_display_names = {k: clean_excel_string(resolve_plan_display_name(k, v)) for k, v in final_results.items()}

    underwriter_keys, uw_groups = _group_plans_by_underwriter(final_results, underwriter_overrides=underwriter_overrides)
    benefit_summary: dict[str, dict[str, Any]] = {}
    benefit_plan_metadata: dict[str, dict[str, Any]] = {}
    word_column_labels: dict[str, str] = {}
    plan_order_counter = 0
    for uw_key in underwriter_keys:
        group = uw_groups[uw_key]
        group_plan_entries = [
            {
                "display_name": display_names[unique_key],
                "plan_name": str(plan_category_assignments.get(unique_key, {}).get("category_label", "") or plan_name or "").strip(),
                "plan_fields": dict(norm or {}),
            }
            for plan_name, norm, unique_key, _ in group["plans"]
        ]
        group_summary = _summarize_underwriter_plan_fields_for_sheet(
            group["display_label"],
            group_plan_entries,
            use_llm_summaries=use_llm_summaries,
        )
        for display_name, values in group_summary.items():
            benefit_summary[display_name] = clean_dict_values(values)
        for plan_name, _, unique_key, data in group["plans"]:
            display_name = display_names[unique_key]
            file_metadata = _get_effective_file_metadata(unique_key, data, report_metadata)
            category_info = plan_category_assignments.get(unique_key, {})
            resolved_underwriter = str(category_info.get("underwriter", "") or _get_effective_underwriter_label(unique_key, data, underwriter_overrides)).strip()
            benefit_plan_metadata[display_name] = {
                "underwriter": resolved_underwriter,
                "plan_name": str(category_info.get("category_label", "") or plan_name or "").strip() or _extract_plan_name_from_unique_key(unique_key),
                "source_plan_name": str(category_info.get("source_plan_name", "") or plan_name or "").strip() or _extract_plan_name_from_unique_key(unique_key),
                "category_label": str(category_info.get("category_label", "") or "").strip(),
                "category_order": int(category_info.get("category_order", 0) or 0),
                "source_group_label": _derive_source_group_label(file_metadata.get("pdf_name", ""), resolved_underwriter),
                "order": plan_order_counter,
                "pdf_name": file_metadata.get("pdf_name", ""),
                "existing_benefit": bool(file_metadata.get("existing_benefit")),
            }
            word_column_labels[display_name] = word_display_names.get(unique_key, display_name)
            plan_order_counter += 1
    actual_raw_data = {
        display_names[k]: clean_dict_values(deepcopy(v["raw"]))
        for k, v in final_results.items()
    }

    word_plans = []
    for unique_key, data in final_results.items():
        normalized_copy = {
            key: clean_report_value(value)
            for key, value in data.get("normalized", {}).items()
        }
        category_info = plan_category_assignments.get(unique_key, {})
        if str(category_info.get("category_label", "") or "").strip():
            normalized_copy["0.2. Plan Name"] = str(category_info.get("category_label", "") or "").strip()
        if str(category_info.get("underwriter", "") or "").strip():
            normalized_copy["0.1. Underwriter"] = str(category_info.get("underwriter", "") or "").strip()
        word_plans.append(
            {
                "unique_key": unique_key,
                "display_name": word_display_names[unique_key],
                "normalized": normalized_copy,
            }
        )

    insurer_summary_view = _build_insurer_summary_view(
        final_results,
        keys,
        use_llm_summaries=use_llm_summaries,
        underwriter_overrides=underwriter_overrides,
        benefit_summary=benefit_summary,
        benefit_plan_metadata=benefit_plan_metadata,
        baseline_underwriter=report_metadata.get("baseline_underwriter", ""),
    )
    insurer_summary_payload = {
        "underwriters": insurer_summary_view["underwriters"],
        "summary": insurer_summary_view["summary"],
        "row_keys": list(insurer_summary_view.get("row_keys", [])),
        "category_labels": list(insurer_summary_view.get("category_labels", [])),
        "best_uw": {
            field_key: sorted(list(labels))
            for field_key, labels in insurer_summary_view["best_uw"].items()
        },
        "worst_uw": {
            field_key: sorted(list(labels))
            for field_key, labels in insurer_summary_view["worst_uw"].items()
        },
        "line_highlights": insurer_summary_view.get("line_highlights", {}),
        "uw_score_stats": insurer_summary_view["uw_score_stats"],
        "best_avg_uws": sorted(list(insurer_summary_view["best_avg_uws"])),
        "multiple_underwriters": insurer_summary_view["multiple_underwriters"],
        "baseline_underwriter": insurer_summary_view.get("baseline_underwriter", ""),
    }

    word_payload = {
        "overall_uw_summary": "",
        "plans": word_plans,
        "column_labels": word_column_labels,
    }

    report_payload = {
        "keys": list(keys),
        "display_names": display_names,
        "benefit_summary": benefit_summary,
        "benefit_plan_metadata": benefit_plan_metadata,
        "actual_raw_data": actual_raw_data,
        "insurer_summary": insurer_summary_payload,
        "word": word_payload,
        "metadata": report_metadata,
    }
    return _reconcile_insurer_summary_from_benefit_summary(report_payload)


def _normalize_single_plan_insurer_summary_payload(report_payload: dict[str, Any]) -> dict[str, Any]:
    insurer_summary = report_payload.get("insurer_summary", {}) or {}
    summary = insurer_summary.get("summary", {}) or {}
    word_plans = report_payload.get("word", {}).get("plans", []) or []
    if not summary or not word_plans:
        return report_payload

    underwriter_plan_names: dict[str, set[str]] = {}
    for plan in word_plans:
        normalized = plan.get("normalized", {}) or {}
        underwriter = str(normalized.get("0.1. Underwriter", "") or "").strip()
        plan_name = str(normalized.get("0.2. Plan Name", "") or "").strip()
        if not underwriter or not plan_name:
            continue
        underwriter_plan_names.setdefault(underwriter, set()).add(plan_name)

    single_plan_underwriters = {
        underwriter: next(iter(plan_names))
        for underwriter, plan_names in underwriter_plan_names.items()
        if len(plan_names) == 1
    }
    if not single_plan_underwriters:
        return report_payload

    normalized_payload = deepcopy(report_payload)
    normalized_summary = normalized_payload.get("insurer_summary", {}).get("summary", {}) or {}
    for field_key, uw_values in normalized_summary.items():
        if not isinstance(uw_values, dict):
            continue
        for underwriter, value in list(uw_values.items()):
            plan_name = single_plan_underwriters.get(str(underwriter or "").strip())
            if not plan_name:
                continue
            text = str(value or "").strip()
            if not text:
                continue
            cleaned_lines = []
            for line in text.split("\n"):
                stripped_line = str(line or "").strip()
                if not stripped_line:
                    continue
                cleaned_lines.append(_strip_prefixed_label(stripped_line, plan_name))
            uw_values[underwriter] = "\n".join(cleaned_lines).strip()

    return normalized_payload


def _build_lossless_insurer_summary_rows(
    benefit_summary: dict[str, dict[str, Any]],
    benefit_plan_metadata: dict[str, dict[str, Any]],
    underwriters: list[str],
    row_keys: list[str],
) -> tuple[list[str], list[str], dict[str, dict[str, str]]]:
    normalized_underwriters = [str(underwriter or "").strip() for underwriter in underwriters if str(underwriter or "").strip()]
    if not normalized_underwriters:
        for metadata in benefit_plan_metadata.values():
            underwriter = str(metadata.get("underwriter", "") or "").strip()
            if underwriter and underwriter not in normalized_underwriters:
                normalized_underwriters.append(underwriter)

    normalized_row_keys: list[str] = []
    excluded_keys = {"0.1. Underwriter", "0.2. Plan Name"}
    preferred_row_order = []
    for source_key in list(_report_ordered_keys()) + list(row_keys or []):
        cleaned_key = str(source_key or "").strip()
        if cleaned_key and cleaned_key not in excluded_keys and cleaned_key not in preferred_row_order:
            preferred_row_order.append(cleaned_key)

    benefit_row_keys: list[str] = []
    for fields in benefit_summary.values():
        for source_key in (fields or {}).keys():
            cleaned_key = str(source_key or "").strip()
            if cleaned_key and cleaned_key not in excluded_keys and cleaned_key not in benefit_row_keys:
                benefit_row_keys.append(cleaned_key)

    for field_key in preferred_row_order:
        if field_key in benefit_row_keys and field_key not in normalized_row_keys:
            normalized_row_keys.append(field_key)
    for field_key in benefit_row_keys:
        if field_key not in normalized_row_keys:
            normalized_row_keys.append(field_key)

    grouped_entries: dict[str, list[dict[str, Any]]] = {underwriter: [] for underwriter in normalized_underwriters}
    for display_name, metadata in benefit_plan_metadata.items():
        underwriter = str(metadata.get("underwriter", "") or "").strip()
        if not underwriter:
            continue
        if underwriter not in grouped_entries:
            grouped_entries[underwriter] = []
            normalized_underwriters.append(underwriter)
        grouped_entries[underwriter].append(
            {
                "display_name": str(display_name or "").strip(),
                "category_label": str(metadata.get("category_label", "") or "").strip(),
                "category_order": int(metadata.get("category_order", 0) or 0),
                "order": int(metadata.get("order", 0) or 0),
            }
        )

    for underwriter in grouped_entries:
        grouped_entries[underwriter].sort(
            key=lambda item: (item["category_order"], item["order"], item["display_name"])
        )

    summary: dict[str, dict[str, str]] = {
        field_key: {underwriter: "" for underwriter in normalized_underwriters}
        for field_key in normalized_row_keys
    }

    for field_key in normalized_row_keys:
        for underwriter in normalized_underwriters:
            lines: list[str] = []
            for item in grouped_entries.get(underwriter, []):
                raw_value = benefit_summary.get(item["display_name"], {}).get(field_key, "")
                value = str(raw_value or "").strip()
                if not value:
                    continue
                category_label = str(item.get("category_label", "") or "").strip()
                if category_label:
                    lines.append(f"{_format_category_display_label(category_label)} : {value}")
                else:
                    lines.append(value)
            summary[field_key][underwriter] = "\n".join(lines).strip()

    return normalized_underwriters, normalized_row_keys, summary


def _reconcile_insurer_summary_from_benefit_summary(report_payload: dict[str, Any]) -> dict[str, Any]:
    benefit_summary = report_payload.get("benefit_summary", {}) or {}
    benefit_plan_metadata = report_payload.get("benefit_plan_metadata", {}) or {}
    insurer_summary = report_payload.get("insurer_summary", {}) or {}
    if not benefit_summary or not benefit_plan_metadata or not insurer_summary:
        return report_payload

    underwriters = list(insurer_summary.get("underwriters", []) or [])
    row_keys = [
        field_key
        for field_key in list(insurer_summary.get("row_keys", []) or report_payload.get("keys", []) or [])
        if field_key not in {"0.1. Underwriter", "0.2. Plan Name"}
    ]
    underwriters, row_keys, rebuilt_summary = _build_lossless_insurer_summary_rows(
        benefit_summary,
        benefit_plan_metadata,
        underwriters,
        row_keys,
    )
    if not underwriters or not row_keys:
        return report_payload

    normalized_payload = deepcopy(report_payload)
    normalized_insurer_summary = normalized_payload.setdefault("insurer_summary", {})
    normalized_insurer_summary["underwriters"] = underwriters
    normalized_insurer_summary["row_keys"] = row_keys
    normalized_insurer_summary["summary"] = rebuilt_summary
    if "line_highlights" not in normalized_insurer_summary:
        normalized_insurer_summary["line_highlights"] = {}

    return normalized_payload


def _collect_insurer_summary_expected_entries(report_payload: dict[str, Any]) -> list[dict[str, str]]:
    benefit_summary = report_payload.get("benefit_summary", {}) or {}
    benefit_plan_metadata = report_payload.get("benefit_plan_metadata", {}) or {}
    expected_entries: list[dict[str, str]] = []
    for display_name, fields in benefit_summary.items():
        metadata = benefit_plan_metadata.get(display_name, {}) or {}
        underwriter = str(metadata.get("underwriter", "") or "").strip()
        category_label = str(metadata.get("category_label", "") or "").strip()
        if not underwriter or not category_label:
            continue
        formatted_category_label = _format_category_display_label(category_label)
        for field_key, raw_value in (fields or {}).items():
            if field_key in {"0.1. Underwriter", "0.2. Plan Name"}:
                continue
            value = str(raw_value or "").strip()
            if not value:
                continue
            expected_entries.append(
                {
                    "display_name": str(display_name or "").strip(),
                    "underwriter": underwriter,
                    "category_label": formatted_category_label,
                    "field_key": str(field_key or "").strip(),
                    "row_label": _format_report_row_label(str(field_key or "").strip()),
                    "value": value,
                    "expected_line": f"{formatted_category_label} : {value}",
                }
            )
    return expected_entries


def _normalize_insurer_summary_line_match_text(value: Any) -> str:
    formatted = _format_raw_cell_text(value)
    if not formatted:
        return ""
    return re.sub(r"\b((?:CAT|Cat)\s+[A-Z])\s*:\s*", r"\1: ", formatted)


def _validate_insurer_summary_payload_completeness(report_payload: dict[str, Any]) -> list[dict[str, str]]:
    summary = (report_payload.get("insurer_summary", {}) or {}).get("summary", {}) or {}
    mismatches: list[dict[str, str]] = []
    for entry in _collect_insurer_summary_expected_entries(report_payload):
        summary_value = str((summary.get(entry["field_key"], {}) or {}).get(entry["underwriter"], "") or "").strip()
        normalized_summary_value = _normalize_insurer_summary_line_match_text(summary_value)
        normalized_expected_line = _normalize_insurer_summary_line_match_text(entry["expected_line"])
        if normalized_expected_line not in normalized_summary_value:
            mismatches.append(
                {
                    **entry,
                    "actual_value": summary_value,
                    "location": "payload",
                }
            )
    return mismatches


def _validate_insurer_summary_workbook_completeness(workbook_path: str | Path, report_payload: dict[str, Any]) -> list[dict[str, str]]:
    workbook_path = Path(workbook_path)
    if not workbook_path.exists():
        return [
            {
                "location": "workbook",
                "field_key": "",
                "row_label": "",
                "underwriter": "",
                "category_label": "",
                "display_name": "",
                "value": "",
                "expected_line": "",
                "actual_value": f"Workbook not found: {workbook_path}",
            }
        ]

    workbook = load_workbook(workbook_path, data_only=True)
    if "Insurer_Summary" not in workbook.sheetnames:
        return [
            {
                "location": "workbook",
                "field_key": "",
                "row_label": "",
                "underwriter": "",
                "category_label": "",
                "display_name": "",
                "value": "",
                "expected_line": "",
                "actual_value": "Insurer_Summary sheet not found",
            }
        ]

    worksheet = workbook["Insurer_Summary"]
    headers = {
        str(worksheet.cell(row=1, column=column_index).value or "").strip(): column_index
        for column_index in range(1, worksheet.max_column + 1)
        if str(worksheet.cell(row=1, column=column_index).value or "").strip()
    }
    row_lookup = {
        str(worksheet.cell(row=row_index, column=1).value or "").strip(): row_index
        for row_index in range(2, worksheet.max_row + 1)
        if str(worksheet.cell(row=row_index, column=1).value or "").strip()
    }

    mismatches: list[dict[str, str]] = []
    for entry in _collect_insurer_summary_expected_entries(report_payload):
        column_index = headers.get(entry["underwriter"])
        row_index = row_lookup.get(entry["row_label"])
        actual_value = ""
        if column_index and row_index:
            actual_value = str(worksheet.cell(row=row_index, column=column_index).value or "").strip()
        normalized_actual_value = _normalize_insurer_summary_line_match_text(actual_value)
        normalized_expected_line = _normalize_insurer_summary_line_match_text(entry["expected_line"])
        if normalized_expected_line not in normalized_actual_value:
            mismatches.append(
                {
                    **entry,
                    "actual_value": actual_value,
                    "location": "workbook",
                }
            )
    return mismatches


def _raise_if_insurer_summary_incomplete(
    report_payload: dict[str, Any],
    workbook_path: str | Path | None = None,
):
    payload_mismatches = _validate_insurer_summary_payload_completeness(report_payload)
    workbook_mismatches = _validate_insurer_summary_workbook_completeness(workbook_path, report_payload) if workbook_path else []
    all_mismatches = payload_mismatches + workbook_mismatches
    if not all_mismatches:
        payload_count = len(payload_mismatches)
        workbook_count = len(workbook_mismatches)
        print(f"✓ Insurer Summary completeness validated (payload mismatches: {payload_count}, workbook mismatches: {workbook_count})")
        return

    preview_lines = []
    for mismatch in all_mismatches[:8]:
        preview_lines.append(
            f"[{mismatch.get('location', '')}] {mismatch.get('underwriter', '')} | {mismatch.get('category_label', '')} | {mismatch.get('field_key', '')} | expected {mismatch.get('expected_line', '')!r} | actual {mismatch.get('actual_value', '')!r}"
        )
    raise ValueError(
        "Insurer Summary is missing non-empty Benefit Summary values:\n" + "\n".join(preview_lines)
    )


def summarize_dict_values(d: dict) -> dict:
    if not d:
        return {}

    cache_key = json.dumps(d, sort_keys=True, ensure_ascii=False, default=str)
    if cache_key in _summary_dict_cache:
        return dict(_summary_dict_cache[cache_key])

    passthrough_keys = {"0.1. Underwriter", "0.2. Plan Name"}
    result = {}
    for key, value in d.items():
        if not isinstance(value, str):
            result[key] = value
            continue
        if key in passthrough_keys:
            result[key] = value.strip()
            continue

        result[key] = _format_benefit_summary_text("", value, max_length=SUMMARY_CELL_MAX_LENGTH) if value else ""

    _summary_dict_cache[cache_key] = dict(result)
    return result


# ── STEP 3: EXCEL OUTPUT ──────────────────────────────────────────────────────
def _write_insurer_summary_sheet_from_payload(output_path: str, report_payload: dict[str, Any]):
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter

    insurer_summary = report_payload.get("insurer_summary", {})
    if not insurer_summary.get("multiple_underwriters"):
        return

    underwriters = list(insurer_summary.get("underwriters", []))
    summary = insurer_summary.get("summary", {})
    line_highlights = {
        field_key: _normalize_insurer_summary_line_highlights(entries)
        for field_key, entries in (insurer_summary.get("line_highlights", {}) or {}).items()
    }
    row_keys = list(insurer_summary.get("row_keys", []) or report_payload.get("keys", []))

    wb = load_workbook(output_path)
    if "Insurer_Summary" in wb.sheetnames:
        del wb["Insurer_Summary"]
    ws = wb.create_sheet("Insurer_Summary")

    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    burgundy_fill  = PatternFill(start_color="A52054", end_color="A52054", fill_type="solid")
    white_font     = Font(color="FFFFFF", bold=True, name="Arial")
    bold_font      = Font(bold=True, name="Arial")
    normal_font    = Font(name="Arial")
    side           = Side(style="thin")
    thin_border    = Border(left=side, right=side, top=side, bottom=side)
    wrap_top       = Alignment(wrap_text=True, vertical="top")
    center_mid     = Alignment(horizontal="center", vertical="center", wrap_text=True)

    header = ["Benefit Category"] + [clean_excel_string(uw) for uw in underwriters]
    for col_idx, h in enumerate(header, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.fill = burgundy_fill
        cell.font = white_font
        cell.alignment = center_mid
        cell.border = thin_border

    data_row_start = 2
    cell_runs_map: dict[str, list[dict[str, Any]]] = {}
    for row_idx, field_key in enumerate(row_keys, data_row_start):
        label = _format_report_row_label(field_key)
        label_cell = ws.cell(row=row_idx, column=1, value=clean_excel_string(label))
        label_cell.font = bold_font
        label_cell.alignment = wrap_top
        label_cell.border = thin_border

        for col_idx, uw in enumerate(underwriters, 2):
            val = clean_excel_string(summary.get(field_key, {}).get(uw, ""))
            rendered_line_highlights = line_highlights.get(field_key, {}).get(uw, [])
            cell = ws.cell(
                row=row_idx,
                column=col_idx,
                value=_build_excel_insurer_summary_display_text(
                    val,
                    rendered_line_highlights,
                ),
            )
            cell.alignment = wrap_top
            cell.border = thin_border
            cell.font = normal_font
            if val:
                cell_ref = f"{get_column_letter(col_idx)}{row_idx}"
                cell_runs_map[cell_ref] = _build_excel_insurer_summary_rich_runs(val, rendered_line_highlights)
            if not val:
                _apply_insurer_summary_rich_text(cell)

    ws.column_dimensions["A"].width = 40
    col_letters = ["B", "C", "D", "E", "F", "G", "H"]
    for i, _ in enumerate(underwriters):
        if i < len(col_letters):
            ws.column_dimensions[col_letters[i]].width = 55

    for row_idx, field_key in enumerate(row_keys, data_row_start):
        max_lines = 1
        max_lines = max(max_lines, _estimate_excel_wrapped_line_count(_format_report_row_label(field_key), 40))
        for uw in underwriters:
            val = summary.get(field_key, {}).get(uw, "")
            lines = _estimate_excel_wrapped_line_count(val, 55)
            max_lines = max(max_lines, lines)
        ws.row_dimensions[row_idx].height = max(30, max_lines * 18)

    wb.save(output_path)
    _apply_ooxml_rich_text_to_sheet_cells(output_path, "Insurer_Summary", cell_runs_map)
    highlighted = sum(
        1
        for field_key in row_keys
        if any(
            entry.get("status")
            for uw in underwriters
            for entry in line_highlights.get(field_key, {}).get(uw, [])
        )
    )
    print(f"  ✓ Sheet 3 'Insurer_Summary' added ({len(underwriters)} underwriters, {len(row_keys)} rows, {highlighted} rows highlighted)")


def save_final_analysis_from_payload(report_payload: dict[str, Any], output_path: str | Path | None = None):
    output_path = str(output_path or (ensure_output_dir() / "Insurance_Technical_Analysis.xlsx"))

    burgundy_fill  = PatternFill(start_color="A52054", end_color="A52054", fill_type="solid")
    highlight_fill = PatternFill(start_color="DFF0D8", end_color="DFF0D8", fill_type="solid")
    raw_label_fill = PatternFill(start_color="E8ECEF", end_color="E8ECEF", fill_type="solid")
    raw_odd_fill   = PatternFill(start_color="FAFBFC", end_color="FAFBFC", fill_type="solid")
    raw_even_fill  = PatternFill(start_color="F3F6F9", end_color="F3F6F9", fill_type="solid")
    raw_empty_fill = PatternFill(start_color="F6F6F6", end_color="F6F6F6", fill_type="solid")
    white_font     = Font(color="FFFFFF", bold=True, name="Arial")
    normal_font    = Font(name="Arial")
    side_style     = Side(style='thin')
    thin_border    = Border(left=side_style, right=side_style, top=side_style, bottom=side_style)
    NOT_FOUND      = {"not mentioned", "not found", "n/a", "none", "not applicable"}

    keys = list(report_payload.get("keys", []))
    summary_map = report_payload.get("benefit_summary", {})
    raw_map = report_payload.get("actual_raw_data", {})
    multiple_underwriters = bool(report_payload.get("insurer_summary", {}).get("multiple_underwriters"))

    print(f"\n  Column headers in Excel: {list(summary_map.keys())}")

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        benefit_df = pd.DataFrame(summary_map).reindex(keys)
        benefit_df.index = [display_benefit_label(key) for key in benefit_df.index]
        benefit_df.to_excel(writer, sheet_name="Benefit_Summary")

        raw_df = pd.DataFrame(raw_map).reindex(keys)
        raw_df.index = [display_benefit_label(key) for key in raw_df.index]
        raw_df.to_excel(writer, sheet_name="Actual_Raw_Data")

        benefit_ws = writer.sheets["Benefit_Summary"]
        raw_ws = writer.sheets["Actual_Raw_Data"]

        for cell in benefit_ws[1]:
            cell.fill = burgundy_fill
            cell.font = white_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border

        benefit_ws.column_dimensions['A'].width = 45
        from openpyxl.utils import get_column_letter
        for col_idx in range(2, benefit_ws.max_column + 1):
            col_letter = get_column_letter(col_idx)
            benefit_ws.column_dimensions[col_letter].width = 55

        for row in benefit_ws.iter_rows(min_row=2, max_row=benefit_ws.max_row, min_col=1, max_col=benefit_ws.max_column):
            for cell in row:
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                cell.border = thin_border
                if cell.column == 1:
                    cell.fill = burgundy_fill
                    cell.font = white_font
                    cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
                    continue

                val_str = str(cell.value).strip() if cell.value else ""
                if "**[[BEST]]**" in val_str:
                    if multiple_underwriters:
                        cell.fill = highlight_fill
                        cell.font = Font(bold=True, name="Arial")
                    else:
                        cell.font = normal_font
                    cell.value = _strip_best_marker(val_str)
                elif any(m in val_str.lower() for m in NOT_FOUND) or not val_str:
                    cell.value = ""
                    cell.font = normal_font
                else:
                    cell.font = normal_font

        for row_idx in range(2, benefit_ws.max_row + 1):
            max_lines = 1
            row_label_value = benefit_ws.cell(row=row_idx, column=1).value
            max_lines = max(max_lines, _estimate_excel_wrapped_line_count(row_label_value, 45))
            for col_idx in range(2, benefit_ws.max_column + 1):
                cell_value = benefit_ws.cell(row=row_idx, column=col_idx).value
                line_count = _estimate_excel_wrapped_line_count(cell_value, 55)
                max_lines = max(max_lines, line_count)
            benefit_ws.row_dimensions[row_idx].height = max(24, max_lines * 18)

        for cell in raw_ws[1]:
            cell.fill = burgundy_fill
            cell.font = white_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = thin_border
        raw_ws.row_dimensions[1].height = 34
        raw_ws.freeze_panes = "B2"
        raw_ws.auto_filter.ref = raw_ws.dimensions
        raw_ws.column_dimensions['A'].width = 42
        for col_idx in range(2, raw_ws.max_column + 1):
            col_letter = get_column_letter(col_idx)
            raw_ws.column_dimensions[col_letter].width = 72

        for row_idx in range(2, raw_ws.max_row + 1):
            data_fill = raw_even_fill if row_idx % 2 == 0 else raw_odd_fill
            for col_idx in range(1, raw_ws.max_column + 1):
                cell = raw_ws.cell(row=row_idx, column=col_idx)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="top")

                if col_idx == 1:
                    cell.font = Font(bold=True, name="Arial", size=11)
                    cell.alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
                    cell.fill = raw_label_fill
                    continue

                val_str = str(cell.value).strip() if cell.value else ""
                if not val_str:
                    cell.value = ""
                    cell.font = normal_font
                    cell.fill = raw_empty_fill
                    continue

                cell.font = normal_font
                cell.fill = data_fill
                _apply_rich_text_emphasis(cell)

        for row_idx in range(2, raw_ws.max_row + 1):
            max_lines = 1
            for col_idx in range(2, raw_ws.max_column + 1):
                cell_value = raw_ws.cell(row=row_idx, column=col_idx).value
                text = _format_raw_cell_text(cell_value)
                line_count = text.count("\n") + 1 if text else 1
                max_lines = max(max_lines, line_count)
            raw_ws.row_dimensions[row_idx].height = max(36, max_lines * 20 + 6)

    _write_insurer_summary_sheet_from_payload(output_path, report_payload)
    print(f"\n✓ Excel saved: {output_path}")


def save_final_analysis(final_results: dict, keys: list, output_path: str | Path | None = None, report_json_path: str | Path | None = None):
    report_payload = build_report_payload(final_results, keys)
    if report_json_path is not None:
        write_json_file(report_json_path, report_payload)
        print(f"✓ Report payload JSON saved: {report_json_path}")
    save_final_analysis_from_payload(report_payload, output_path=output_path)


# ── STEP 4: WORD OUTPUT ───────────────────────────────────────────────────────
WORD_HEADER_FILL = "C1245A"
WORD_ROW_FILL_LIGHT = "FFFFFF"
WORD_ROW_FILL_DARK = "C9CDD3"
WORD_ROW_FILL_WHITE = "FFFFFF"
WORD_HEADER_FONT_NAME = "Arial"
WORD_HEADER_FONT_SIZE = 10
WORD_BENEFIT_SECTION_HEADERS = {
    "7. In-Patient Benefits": "INPATIENT SERVICES",
    "12. Out Patient Deductible on Consultation": "OUT-PATIENT SERVICES",
    "18. Maternity In Patient Services": "MATERNITY BENEFITS",
    "22. Dental Benefit": "ADDITIONAL BENEFITS",
}


def _set_run_font_name(run, font_name: str):
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:eastAsia'), font_name)
    r_fonts.set(qn('w:cs'), font_name)


def _style_word_header_run(run):
    _set_run_font_name(run, WORD_HEADER_FONT_NAME)
    run.font.size = Pt(WORD_HEADER_FONT_SIZE)
    run.bold = True


def _style_word_heading(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not paragraph.runs:
        paragraph.add_run()
    for run in paragraph.runs:
        _style_word_header_run(run)


def set_burgundy_and_center(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), WORD_HEADER_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.color.rgb = RGBColor(255, 255, 255)
    _style_word_header_run(run)


def _get_word_usable_width(section):
    return Emu(int(section.page_width) - int(section.left_margin) - int(section.right_margin))


def add_centered_burgundy_header(table, text, total_columns, section=None):
    row = table.add_row()
    merged_cell = row.cells[0].merge(row.cells[total_columns - 1])
    merged_cell.text = text
    set_burgundy_and_center(merged_cell)
    if section is not None:
        _set_word_cell_width(merged_cell, _get_word_usable_width(section))


def _maybe_add_word_section_subheader(table, benefit_key: str, total_columns: int, section):
    subheader_text = WORD_BENEFIT_SECTION_HEADERS.get(str(benefit_key or "").strip(), "")
    if not subheader_text:
        return
    add_centered_burgundy_header(table, subheader_text, total_columns, section=section)


def _set_word_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def _set_word_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width.twips)))
    tc_w.set(qn('w:type'), 'dxa')


def _apply_word_table_column_widths(table, section, serial_width=Mm(12), benefit_width=Mm(70)):
    table.autofit = False
    total_columns = len(table.rows[0].cells) if table.rows else 0
    plan_columns = max(total_columns - 2, 0)
    available_width = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    fixed_width = int(serial_width) + int(benefit_width)
    remaining_width = max(available_width - fixed_width, 0)
    plan_width = Emu(remaining_width // plan_columns) if plan_columns else Emu(0)

    for row in table.rows:
        if len(row.cells) >= 2:
            _set_word_cell_width(row.cells[0], serial_width)
            _set_word_cell_width(row.cells[1], benefit_width)
            for idx in range(2, len(row.cells)):
                _set_word_cell_width(row.cells[idx], plan_width)


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn('w:tblHeader'))
    if tbl_header is None:
        tbl_header = OxmlElement('w:tblHeader')
        tr_pr.append(tbl_header)
    tbl_header.set(qn('w:val'), 'true')


def _apply_word_row_fill(row_cells, fill: str):
    for cell in row_cells:
        _set_word_cell_fill(cell, fill)


def _add_word_insurer_summary_table(doc: Document, final_results: dict, ordered_keys: list[str]):
    summary_view = _build_insurer_summary_view(final_results, ordered_keys)
    if not summary_view["multiple_underwriters"]:
        return

    underwriters = summary_view["underwriters"]
    summary = summary_view["summary"]
    line_highlights = {
        field_key: _normalize_insurer_summary_line_highlights(entries)
        for field_key, entries in (summary_view.get("line_highlights", {}) or {}).items()
    }

    heading = doc.add_heading('Insurer Summary', level=1)
    _style_word_heading(heading)

    table = doc.add_table(rows=1, cols=len(underwriters) + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Benefit Category'
    set_burgundy_and_center(hdr_cells[0])

    for idx, underwriter in enumerate(underwriters, 1):
        hdr_cells[idx].text = clean_excel_string(underwriter)
        set_burgundy_and_center(hdr_cells[idx])

    for field_key in ordered_keys:
        row_cells = table.add_row().cells
        label = field_key.split('. ', 1)[-1] if '. ' in field_key else field_key
        row_cells[0].text = clean_excel_string(label)
        if row_cells[0].paragraphs[0].runs:
            row_cells[0].paragraphs[0].runs[0].bold = True

        for idx, underwriter in enumerate(underwriters, 1):
            row_cells[idx].text = ""
            _write_word_insurer_summary_value(
                row_cells[idx],
                clean_excel_string(summary[field_key].get(underwriter, '')),
                line_highlights.get(field_key, {}).get(underwriter, []),
            )

    doc.add_page_break()


def _add_word_insurer_summary_table_from_payload(doc: Document, report_payload: dict[str, Any]):
    insurer_summary = report_payload.get("insurer_summary", {})
    if not insurer_summary.get("multiple_underwriters"):
        return

    underwriters = list(insurer_summary.get("underwriters", []))
    summary = insurer_summary.get("summary", {})
    line_highlights = {
        field_key: _normalize_insurer_summary_line_highlights(entries)
        for field_key, entries in (insurer_summary.get("line_highlights", {}) or {}).items()
    }
    ordered_keys = list(report_payload.get("keys", []))

    heading = doc.add_heading('Insurer Summary', level=1)
    _style_word_heading(heading)

    table = doc.add_table(rows=1, cols=len(underwriters) + 1)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Benefit Category'
    set_burgundy_and_center(hdr_cells[0])

    for idx, underwriter in enumerate(underwriters, 1):
        hdr_cells[idx].text = clean_excel_string(underwriter)
        set_burgundy_and_center(hdr_cells[idx])

    for field_key in ordered_keys:
        row_cells = table.add_row().cells
        label = field_key.split('. ', 1)[-1] if '. ' in field_key else field_key
        row_cells[0].text = clean_excel_string(label)
        if row_cells[0].paragraphs[0].runs:
            row_cells[0].paragraphs[0].runs[0].bold = True

        for idx, underwriter in enumerate(underwriters, 1):
            row_cells[idx].text = ""
            _write_word_insurer_summary_value(
                row_cells[idx],
                clean_excel_string(summary.get(field_key, {}).get(underwriter, '')),
                line_highlights.get(field_key, {}).get(underwriter, []),
            )


    doc.add_page_break()


def build_overall_uw_word_summary(final_results: dict, ordered_keys: list[str]) -> str:
    return ""


def _configure_word_page_a3_landscape(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(420)
    section.page_height = Mm(297)


def save_to_word(final_results: dict, ordered_keys: list, output_path: str | Path | None = None):
    doc = Document()
    _configure_word_page_a3_landscape(doc)

    title_head = doc.add_heading('Corporate Health Insurance Comparison', 0)
    _style_word_heading(title_head)
    summary_head = doc.add_heading('Executive Decision Summary', level=1)
    _style_word_heading(summary_head)

    overall_uw_summary = build_overall_uw_word_summary(final_results, ordered_keys)
    if overall_uw_summary:
        for line in overall_uw_summary.split("\n"):
            if line.strip():
                doc.add_paragraph(clean_excel_string(line.strip()))

    plans = list(final_results.keys())
    table = doc.add_table(rows=1, cols=len(plans) + 2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sr No'
    set_burgundy_and_center(hdr_cells[0])
    hdr_cells[1].text = 'Benefit Categories'
    set_burgundy_and_center(hdr_cells[1])
    _set_repeat_table_header(table.rows[0])
    _apply_word_table_column_widths(table, doc.sections[0])

    def get_display_name(unique_key, data):
        norm = data.get("normalized", {})
        underwriter = norm.get("0.1. Underwriter", "").strip()
        plan_name   = norm.get("0.2. Plan Name", "").strip()
        if not plan_name:
            plan_name = unique_key.split("|")[-1].strip()
        if underwriter and plan_name:
            return f"{underwriter} – {plan_name}"
        return plan_name or unique_key

    for i, plan_key in enumerate(plans):
        display_name = clean_excel_string(get_display_name(plan_key, final_results[plan_key]))
        hdr_cells[i + 2].text = display_name
        set_burgundy_and_center(hdr_cells[i + 2])

    for row_number, key in enumerate(ordered_keys, 1):
        _maybe_add_word_section_subheader(table, key, len(plans) + 2, doc.sections[0])
        benefit_label = clean_excel_string(display_benefit_label(key))

        row_cells = table.add_row().cells
        row_fill = WORD_ROW_FILL_DARK if row_number % 2 else WORD_ROW_FILL_LIGHT
        _apply_word_row_fill(row_cells, row_fill)
        row_cells[0].text = f"{row_number:02d}."
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_cells[0].paragraphs[0].runs:
            row_cells[0].paragraphs[0].runs[0].bold = True

        row_cells[1].text = benefit_label
        if row_cells[1].paragraphs[0].runs:
            row_cells[1].paragraphs[0].runs[0].bold = True

        for i, plan in enumerate(plans):
            val = clean_excel_string(str(final_results[plan]["normalized"].get(key, "")))
            cell = row_cells[i + 2]
            if not val.strip() or val.lower() in ["not mentioned", "not found"]:
                cell.text = ""
            else:
                is_best = "**[[BEST]]**" in val
                clean_val = val.replace("**[[BEST]]**", "").strip()
                cell.text = ""
                run = cell.paragraphs[0].add_run(clean_val)
                if is_best:
                    run.bold = True
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    _apply_word_table_column_widths(table, doc.sections[0])

    output_path = str(output_path or (ensure_output_dir() / "Corporate_Comparison_Final.docx"))
    doc.save(output_path)
    print(f"✓ Word doc saved: {output_path}")


def _resolve_word_table_source(report_payload: dict[str, Any]) -> tuple[
    list[tuple[str, str]],
    dict[str, dict[str, Any]],
    str,
    dict[str, set[str]],
    dict[str, set[str]],
    dict[str, dict[str, list[dict[str, str]]]],
]:
    insurer_summary = report_payload.get("insurer_summary", {}) or {}
    insurer_underwriters = list(insurer_summary.get("underwriters", []) or [])
    insurer_summary_map = insurer_summary.get("summary", {}) or {}
    if insurer_summary.get("multiple_underwriters") and insurer_underwriters and insurer_summary_map:
        has_nonblank_summary = any(
            str(value).strip()
            for field_values in insurer_summary_map.values()
            for value in (field_values or {}).values()
        )
        if has_nonblank_summary:
            transposed_summary = {
                underwriter: {
                    field_key: str((field_values or {}).get(underwriter, ""))
                    for field_key, field_values in insurer_summary_map.items()
                }
                for underwriter in insurer_underwriters
            }
            best_uw = {
                field_key: set(labels)
                for field_key, labels in (insurer_summary.get("best_uw", {}) or {}).items()
            }
            worst_uw = {
                field_key: set(labels)
                for field_key, labels in (insurer_summary.get("worst_uw", {}) or {}).items()
            }
            line_highlights = {
                field_key: _normalize_insurer_summary_line_highlights(entries)
                for field_key, entries in (insurer_summary.get("line_highlights", {}) or {}).items()
            }
            return (
                [(clean_excel_string(name), name) for name in insurer_underwriters],
                transposed_summary,
                "insurer_summary",
                best_uw,
                worst_uw,
                line_highlights,
            )

    benefit_summary = report_payload.get("benefit_summary", {}) or {}
    if benefit_summary:
        word_column_labels = report_payload.get("word", {}).get("column_labels", {}) or {}
        return (
            [
                (clean_excel_string(str(word_column_labels.get(name, name) or name)), name)
                for name in benefit_summary.keys()
            ],
            benefit_summary,
            "benefit_summary",
            {},
            {},
            {},
        )

    word_plans = report_payload.get("word", {}).get("plans", []) or []
    if word_plans:
        word_plan_map = {
            str(plan.get("unique_key", "")): dict(plan.get("normalized", {}) or {})
            for plan in word_plans
            if str(plan.get("unique_key", "")).strip()
        }
        word_columns = [
            (
                clean_excel_string(str(plan.get("display_name", "") or plan.get("unique_key", ""))),
                str(plan.get("unique_key", "")),
            )
            for plan in word_plans
            if str(plan.get("unique_key", "")).strip()
        ]
        if word_columns and word_plan_map:
            return word_columns, word_plan_map, "word_plans", {}, {}, {}

    return [], {}, "empty", {}, {}, {}


def save_to_word_from_payload(report_payload: dict[str, Any], output_path: str | Path | None = None):
    doc = Document()
    _configure_word_page_a3_landscape(doc)

    title_head = doc.add_heading('Corporate Health Insurance Comparison', 0)
    _style_word_heading(title_head)
    summary_head = doc.add_heading('Executive Decision Summary', level=1)
    _style_word_heading(summary_head)

    overall_uw_summary = str(report_payload.get("word", {}).get("overall_uw_summary", "")).strip()
    if overall_uw_summary:
        for line in overall_uw_summary.split("\n"):
            if line.strip():
                doc.add_paragraph(clean_excel_string(line.strip()))

    word_columns, word_source_map, word_source_kind, best_uw_map, worst_uw_map, line_highlights_map = _resolve_word_table_source(report_payload)

    ordered_keys = list(
        report_payload.get("insurer_summary", {}).get("row_keys", [])
        if word_source_kind == "insurer_summary"
        else report_payload.get("keys", [])
    )
    table = doc.add_table(rows=1, cols=len(word_columns) + 2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sr No'
    set_burgundy_and_center(hdr_cells[0])
    hdr_cells[1].text = 'Benefit Categories'
    set_burgundy_and_center(hdr_cells[1])
    _set_repeat_table_header(table.rows[0])
    _apply_word_table_column_widths(table, doc.sections[0])

    for idx, (display_name, _) in enumerate(word_columns, 1):
        hdr_cells[idx + 1].text = display_name
        set_burgundy_and_center(hdr_cells[idx + 1])

    for row_number, key in enumerate(ordered_keys, 1):
        _maybe_add_word_section_subheader(table, key, len(word_columns) + 2, doc.sections[0])
        benefit_label = clean_excel_string(_format_report_row_label(key))

        row_cells = table.add_row().cells
        row_fill = WORD_ROW_FILL_DARK if row_number % 2 else WORD_ROW_FILL_LIGHT
        _apply_word_row_fill(row_cells, row_fill)
        row_cells[0].text = f"{row_number:02d}."
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_cells[0].paragraphs[0].runs:
            row_cells[0].paragraphs[0].runs[0].bold = True

        row_cells[1].text = benefit_label
        if row_cells[1].paragraphs[0].runs:
            row_cells[1].paragraphs[0].runs[0].bold = True

        for idx, (_, source_key) in enumerate(word_columns, 1):
            raw_value = str(word_source_map.get(source_key, {}).get(key, ""))
            val = clean_excel_string(raw_value)
            cell = row_cells[idx + 1]
            if not val.strip() or val.lower() in ["not mentioned", "not found"]:
                cell.text = ""
            else:
                is_best = _has_best_marker(raw_value)
                clean_val = _strip_best_marker(val)
                cell.text = ""
                if word_source_kind == "insurer_summary":
                    _write_word_insurer_summary_value(cell, clean_val, line_highlights_map.get(key, {}).get(source_key, []))
                    runs = [run for paragraph in cell.paragraphs for run in paragraph.runs]
                else:
                    run = cell.paragraphs[0].add_run(clean_val)
                    runs = [run]
                if word_source_kind != "insurer_summary" and is_best:
                    for word_run in runs:
                        word_run.bold = True
                        word_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    _apply_word_table_column_widths(table, doc.sections[0])

    output_path = str(output_path or (ensure_output_dir() / "Corporate_Comparison_Final.docx"))
    doc.save(output_path)
    print(f"✓ Word doc saved: {output_path}")


def generate_reports_from_payload(
    report_payload: dict[str, Any],
    output_dir: str | Path | None = None,
    report_json_path: str | Path | None = None,
) -> dict[str, str]:
    output_dir = ensure_output_dir(output_dir)
    report_payload = _reconcile_insurer_summary_from_benefit_summary(report_payload)
    report_payload = _normalize_single_plan_insurer_summary_payload(report_payload)
    if report_json_path is not None:
        write_json_file(report_json_path, report_payload)
        print(f"✓ Report payload JSON saved: {report_json_path}")
    excel_path = output_dir / "Insurance_Technical_Analysis.xlsx"
    word_path = output_dir / "Corporate_Comparison_Final.docx"
    save_final_analysis_from_payload(report_payload, output_path=excel_path)
    save_to_word_from_payload(report_payload, output_path=word_path)
    _raise_if_insurer_summary_incomplete(report_payload, workbook_path=excel_path)
    return {
        "excel": str(excel_path),
        "word": str(word_path),
    }


def generate_reports_from_final_results(final_results: dict[str, Any], output_dir: str | Path | None = None, report_json_path: str | Path | None = None) -> dict[str, str]:
    ordered_keys = _report_ordered_keys()
    metadata_path = None
    report_metadata = {}
    if report_json_path is not None:
        metadata_path = Path(report_json_path).with_name("quote_compare_run_metadata.json")
    elif output_dir is not None:
        metadata_path = ensure_output_dir(output_dir) / "quote_compare_run_metadata.json"
    if metadata_path is not None and Path(metadata_path).exists():
        report_metadata = _normalize_report_metadata(read_json_file(metadata_path))

    report_payload = build_report_payload(final_results, ordered_keys, use_llm_summaries=True, report_metadata=report_metadata)
    if report_json_path is not None:
        write_json_file(report_json_path, report_payload)
        print(f"✓ Report payload JSON saved: {report_json_path}")
    outputs = generate_reports_from_payload(report_payload, output_dir=output_dir)
    outputs["report_json"] = str(report_json_path) if report_json_path else ""
    return outputs


def resume_from_raw_extractions(
    raw_extractions: dict[str, Any],
    metadata: dict[str, Any] | None = None,
    pdf_paths: list[str] | None = None,
    output_dir: str | Path | None = None,
) -> dict[str, str]:
    output_dir = ensure_output_dir(output_dir)
    report_metadata = _normalize_report_metadata(metadata)
    pdf_path_map = {
        os.path.basename(str(pdf_path or "").strip()): str(pdf_path or "").strip()
        for pdf_path in (pdf_paths or [])
        if str(pdf_path or "").strip()
    }
    pdf_filters = {
        os.path.basename(str(pdf_path or "").strip())
        for pdf_path in (pdf_paths or [])
        if str(pdf_path or "").strip()
    }

    filtered_raw: dict[str, dict[str, Any]] = {}
    for unique_key, raw_fields in (raw_extractions or {}).items():
        if not isinstance(raw_fields, dict):
            continue
        pdf_name = _extract_pdf_name_from_unique_key(unique_key)
        if pdf_filters and pdf_name not in pdf_filters:
            continue
        filtered_raw[unique_key] = deepcopy(raw_fields)

    if not filtered_raw:
        if pdf_filters:
            raise SystemExit(
                "No matching plans found in the raw extraction JSON for the requested PDFs: "
                + ", ".join(sorted(pdf_filters))
            )
        raise SystemExit("The raw extraction JSON is empty or does not contain any usable plan records.")

    active_pdfs = list(dict.fromkeys(_extract_pdf_name_from_unique_key(unique_key) for unique_key in filtered_raw.keys()))
    report_metadata = _validate_baseline_selection(active_pdfs, report_metadata)

    print(f"\n{'='*60}")
    print("Resume Mode: normalization and report generation from raw extractions")
    print(f"Processing PDFs: {active_pdfs}")
    print(f"{'='*60}")

    all_normalized_data: dict[str, dict[str, Any]] = {}
    final_results: dict[str, dict[str, Any]] = {}
    ordered_keys = _report_ordered_keys()

    for pdf_name in active_pdfs:
        global _current_pdf
        _current_pdf = pdf_name

        pdf_unique_keys = [unique_key for unique_key in filtered_raw.keys() if _extract_pdf_name_from_unique_key(unique_key) == pdf_name]
        verified_plans: dict[str, dict[str, Any]] = {}

        for unique_key in pdf_unique_keys:
            plan_name = _extract_plan_name_from_unique_key(unique_key)
            raw_fields = deepcopy(filtered_raw[unique_key])
            if not str(raw_fields.get("0.2. Plan Name", "") or "").strip():
                raw_fields["0.2. Plan Name"] = plan_name
            verified_plans[plan_name] = raw_fields

        print(f"\n{'-'*60}")
        print(f"Resuming PDF: {pdf_name}")
        print(f"Plans in raw snapshot: {list(verified_plans.keys())}")
        verified_plans = _backfill_annual_limit_from_cached_markdown(pdf_path_map.get(pdf_name, ""), verified_plans)
        verified_plans = _backfill_tpa_from_cached_markdown(pdf_path_map.get(pdf_name, ""), verified_plans)
        for unique_key in pdf_unique_keys:
            plan_name = _extract_plan_name_from_unique_key(unique_key)
            if plan_name in verified_plans:
                filtered_raw[unique_key] = deepcopy(verified_plans[plan_name])
        normalized_plans = run_post_extraction_normalization(verified_plans)

        for unique_key in pdf_unique_keys:
            plan_name = _extract_plan_name_from_unique_key(unique_key)
            normalized_fields = deepcopy(normalized_plans.get(plan_name, verified_plans.get(plan_name, {})))
            all_normalized_data[unique_key] = normalized_fields
            file_metadata = _get_effective_file_metadata(
                unique_key,
                {"file_metadata": report_metadata.get("pdf_file_metadata", {}).get(pdf_name, {})},
                report_metadata,
            )
            final_results[unique_key] = {
                "raw": deepcopy(filtered_raw[unique_key]),
                "normalized": normalized_fields,
                "scores": {},
                "total_score": None,
                "uw_conclusion": "",
                "file_metadata": file_metadata,
            }
            print(f"  ✓ {unique_key}: normalized and finalized without page extraction")

    raw_snapshot_path = output_dir / "quote_compare_raw_extractions.json"
    final_results_path = output_dir / "quote_compare_final_results.json"
    report_payload_path = output_dir / "quote_compare_report_payload.json"
    metadata_path = output_dir / "quote_compare_run_metadata.json"

    write_json_file(raw_snapshot_path, filtered_raw)
    write_json_file(final_results_path, final_results)
    write_json_file(metadata_path, report_metadata)
    print(f"✓ Raw extraction JSON saved: {raw_snapshot_path}")
    print(f"✓ Final results JSON saved: {final_results_path}")
    print(f"✓ Run metadata JSON saved: {metadata_path}")

    report_payload = build_report_payload(final_results, ordered_keys, report_metadata=report_metadata)
    write_json_file(report_payload_path, report_payload)
    print(f"✓ Report payload JSON saved: {report_payload_path}")
    generate_reports_from_payload(report_payload, output_dir=output_dir)
    _save_token_report(tracker, output_path=output_dir / "Token_Usage_Report.xlsx")

    return {
        "raw_json": str(raw_snapshot_path),
        "final_json": str(final_results_path),
        "report_json": str(report_payload_path),
        "metadata_json": str(metadata_path),
        "excel": str(output_dir / "Insurance_Technical_Analysis.xlsx"),
        "word": str(output_dir / "Corporate_Comparison_Final.docx"),
        "token_report": str(output_dir / "Token_Usage_Report.xlsx"),
    }


# ── STEP 5: MAIN ──────────────────────────────────────────────────────────────
def _save_token_report(tracker: TokenTracker, output_path: str | Path | None = None):
    summary = tracker.summary()
    rows = []
    for label, stats in summary.items():
        rows.append({"PDF / Source": label, **stats})

    df = pd.DataFrame(rows).set_index("PDF / Source")
    call_df = pd.DataFrame(tracker.call_details())
    output_path = str(output_path or (ensure_output_dir() / "Token_Usage_Report.xlsx"))

    try:
        with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Token_Usage")
            ws = writer.sheets["Token_Usage"]

            from openpyxl.styles import PatternFill, Font, Alignment
            burgundy = PatternFill(start_color="A52054", end_color="A52054", fill_type="solid")
            white_b  = Font(color="FFFFFF", bold=True, name="Arial")
            total_fill = PatternFill(start_color="FFE0E6", end_color="FFE0E6", fill_type="solid")
            bold_f   = Font(bold=True, name="Arial")

            for cell in ws[1]:
                cell.fill = burgundy
                cell.font = white_b
                cell.alignment = Alignment(horizontal="center")

            for cell in ws[ws.max_row]:
                cell.fill = total_fill
                cell.font = bold_f

            ws.column_dimensions["A"].width = 45
            for col in ["B", "C", "D", "E", "F"]:
                ws.column_dimensions[col].width = 18

            cost_col = ws.max_column
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=cost_col, max_col=cost_col):
                for cell in row:
                    cell.number_format = '"$"#,##0.0000'

            if not call_df.empty:
                call_df.to_excel(writer, sheet_name="Call_Details", index=False)
                call_ws = writer.sheets["Call_Details"]

                for cell in call_ws[1]:
                    cell.fill = burgundy
                    cell.font = white_b
                    cell.alignment = Alignment(horizontal="center")

                detail_widths = {
                    "A": 10,
                    "B": 40,
                    "C": 28,
                    "D": 28,
                    "E": 18,
                    "F": 18,
                    "G": 14,
                    "H": 14,
                    "I": 14,
                    "J": 14,
                }
                for col, width in detail_widths.items():
                    call_ws.column_dimensions[col].width = width

                detail_cost_col = call_ws.max_column
                for row in call_ws.iter_rows(min_row=2, max_row=call_ws.max_row, min_col=detail_cost_col, max_col=detail_cost_col):
                    for cell in row:
                        cell.number_format = '"$"#,##0.000000'

        print(f"✓ Token report saved: {output_path}")
    except PermissionError as exc:
        print(f"⚠️  Token report not saved: {exc}")


def main(metadata: dict = None, pdf_paths: list = None, output_dir: str | Path | None = None):
    metadata = metadata or {}
    pdf_paths = pdf_paths or []
    output_dir = ensure_output_dir(output_dir)
    report_metadata = _normalize_report_metadata(metadata)

    print("\n" + "="*60)
    print("STARTUP DIAGNOSTICS")
    print("="*60)
    print(f"  AZURE_DOC_ENDPOINT : {AZURE_ENDPOINT!r}")
    key_display = (AZURE_KEY[:6] + '...' + AZURE_KEY[-4:]) if AZURE_KEY and len(AZURE_KEY) > 10 else repr(AZURE_KEY)
    print(f"  AZURE_DOC_KEY      : {key_display}")
    premium_key_display = (OPENAI_API_KEY[:6] + '...' + OPENAI_API_KEY[-4:]) if OPENAI_API_KEY and len(OPENAI_API_KEY) > 10 else repr(OPENAI_API_KEY)
    standard_key_display = (OPENAI_STANDARD_API_KEY[:6] + '...' + OPENAI_STANDARD_API_KEY[-4:]) if OPENAI_STANDARD_API_KEY and len(OPENAI_STANDARD_API_KEY) > 10 else repr(OPENAI_STANDARD_API_KEY)
    print(f"  OPENAI_ENDPOINT    : {openai_endpoint!r}")
    print(f"  OPENAI_KEY         : {premium_key_display}")
    print(f"  OPENAI_DEPLOY      : {openai_deployment!r}")
    print(f"  41_ENDPOINT        : {openai_standard_endpoint!r}")
    print(f"  41_KEY             : {standard_key_display}")
    print(f"  41_DEPLOY          : {openai_standard_deployment!r}")
    print(f"  REASONING_EFFORT   : {OPENAI_REASONING_EFFORT!r}")
    print(f"  FORCE_REFRESH      : {FORCE_REFRESH}")
    print(f"  OUTPUT_DIR         : {str(output_dir)!r}")
    print("="*60 + "\n")

    if not AZURE_ENDPOINT or not AZURE_ENDPOINT.startswith("http"):
        print("❌ AZURE_DOC_ENDPOINT is missing or invalid — check your .env file")
        return
    if not AZURE_KEY or len(AZURE_KEY) < 10:
        print("❌ AZURE_DOC_KEY is missing or too short — check your .env file")
        return
    if not openai_endpoint or not OPENAI_API_KEY or not openai_deployment:
        print("❌ Premium OpenAI env vars missing — check AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_API_KEY, AZURE_OPENAI_CHAT_DEPLOYMENT_NAME")
        return
    if not openai_standard_endpoint or not OPENAI_STANDARD_API_KEY or not openai_standard_deployment:
        print("❌ Standard OpenAI env vars missing — check AZURE_OPENAI_STANDARD_ENDPOINT/AZURE_OPENAI_41_ENDPOINT, AZURE_OPENAI_STANDARD_API_KEY/AZURE_OPENAI_41_KEY, AZURE_OPENAI_STANDARD_DEPLOYMENT_NAME/AZURE_OPENAI_41_DEPLOYMENT_NAME")
        return

    files = pdf_paths or [
        # Specify each input as one metadata record in this single list.
        # Recommended format:
        # {
        #     "file_path": r"FILE_PATH_1.pdf",
        #     "uw_name": "UW_NAME_1",
        #     "existing_benefit": True,
        # },
        # {
        #     "file_path": r"FILE_PATH_2.pdf",
        #     "uw_name": "UW_NAME_2",
        #     "existing_benefit": False,
        # },
        {
            "file_path": r"data\UAE\SANCTUARY\TAKAFUL-EXISTING\SANCTUARY-TAKAFUL-A.pdf",
            "uw_name": "Takaful",
            "existing_benefit": True,
        },
        {
            "file_path": r"data\UAE\SANCTUARY\TAKAFUL-EXISTING\SANCTUARY-TAKAFUL-B.pdf",
            "uw_name": "Takaful",
            "existing_benefit": True,
        },
         {
            "file_path": r"data\UAE\SANCTUARY\DAMAN\SACTUARY-DAMAN-1.pdf",
            "uw_name": "DAMAN",
            "existing_benefit": False,
        },
         {
            "file_path": r"data\UAE\SANCTUARY\DAMAN\SACTUARY-DAMAN-2.pdf",
            "uw_name": "DAMAN",
            "existing_benefit": False,
        },
         {
            "file_path": r"data\UAE\SANCTUARY\GIG\SANCTUARY-GIG-1.pdf",
            "uw_name": "GIG",
            "existing_benefit": False,
        },
         {
            "file_path": r"data\UAE\SANCTUARY\GIG\SANCTUARY-GIG-2.pdf",
            "uw_name": "GIG",
            "existing_benefit": False,
          },
        
         {
            "file_path": r"data\UAE\SANCTUARY\ORIENT\SANCTUARY-ORIENT-1.pdf",
            "uw_name": "ORIENT",
            "existing_benefit": False,
        },
         {
            "file_path": r"data\UAE\SANCTUARY\ORIENT\SANCTUARY-ORIENT-2.pdf",
            "uw_name": "ORIENT",
            "existing_benefit": False,
        }
    ]
    file_entries = _normalize_file_entries(files)
    files = [entry["path"] for entry in file_entries]

    entry_overrides = {
        os.path.basename(entry["path"]): entry["uw_name"]
        for entry in file_entries
        if entry["uw_name"]
    }
    if entry_overrides:
        report_metadata["pdf_underwriter_overrides"] = {
            **entry_overrides,
            **report_metadata.get("pdf_underwriter_overrides", {}),
        }

    report_metadata["pdf_file_metadata"] = _merge_pdf_file_metadata(
        report_metadata.get("pdf_file_metadata", {}),
        _build_entry_file_metadata(file_entries),
    )
    report_metadata = _validate_baseline_selection(files, report_metadata)

    configured_overrides = report_metadata.get("pdf_underwriter_overrides", {})
    if configured_overrides:
        print("  ℹ️  Applying user-provided underwriter overrides for report grouping and Word output:")
        for pdf_name, underwriter_label in configured_overrides.items():
            print(f"     - {pdf_name}: {underwriter_label}")

    configured_file_metadata = report_metadata.get("pdf_file_metadata", {})
    if configured_file_metadata:
        print("  ℹ️  Applying per-file baseline metadata:")
        for pdf_name, file_metadata in configured_file_metadata.items():
            print(f"     - {pdf_name}: existing={bool(file_metadata.get('existing_benefit'))}")
    
    if len(files) == 1:
        print(f"ℹ️  Single-file mode enabled for: {files[0]}")
    
    all_raw_data        = {}
    raw_export_data     = {}
    all_normalized_data = {}
    final_results       = {}
    page_extraction_audit = {}
    ordered_keys = _report_ordered_keys()

    for pdf in files:
        pdf_name = os.path.basename(pdf)
        print(f"\n{'='*60}\nProcessing: {pdf_name}\n{'='*60}")
        global _current_pdf
        _current_pdf = pdf_name
        try:
            if FORCE_REFRESH:
                print(f"  FORCE_REFRESH=True — clearing all caches for {pdf_name}")
                clear_cache(pdf)

            _, structured_md = get_ocr_text(pdf)
            print(f"  Markdown ready ({len(structured_md):,} chars) — running extraction...")

            extracted_data, page_audit = run_extraction_pass(structured_md)
            if not isinstance(extracted_data, dict):
                print(f"  ❌ Extraction returned non-dict for {pdf_name}")
                continue

            page_extraction_audit[pdf_name] = page_audit
            extracted_data = run_post_extraction_verification(structured_md, extracted_data)
            extracted_data = run_document_evidence_enrichment(structured_md, extracted_data)
            normalized_extracted_data = run_post_extraction_normalization(extracted_data)

            print(f"  Plans found: {list(extracted_data.keys())}")

            # Words the LLM returns when it cannot identify the plan name
            GENERIC_PLAN_NAMES = {"plan", "category", "n/a", "none", "",
                                   "not mentioned", "unknown", "plan name"}

            for plan_name, raw_fields in extracted_data.items():
                if not isinstance(raw_fields, dict):
                    print(f"  ⚠️  Skipping '{plan_name}' — not a dict")
                    continue

                plan_name_key = "0.2. Plan Name"
                # The JSON top-level key (plan_name) is set by the LLM directly
                # from the document column header and is the most reliable source.
                # Only use the inner field value if it is actually meaningful.
                extracted_pn = raw_fields.get(plan_name_key, "").strip()
                if extracted_pn.lower() in GENERIC_PLAN_NAMES:
                    raw_fields[plan_name_key] = plan_name
                    authoritative_name = plan_name
                else:
                    authoritative_name = extracted_pn

                verified_fields = raw_fields
                normalized_fields = normalized_extracted_data.get(plan_name, verified_fields)

                # Hard-lock the name after verification too
                if verified_fields.get(plan_name_key, "").strip().lower() in GENERIC_PLAN_NAMES:
                    verified_fields[plan_name_key] = authoritative_name
                if normalized_fields.get(plan_name_key, "").strip().lower() in GENERIC_PLAN_NAMES:
                    normalized_fields = dict(normalized_fields)
                    normalized_fields[plan_name_key] = authoritative_name

                unique_key = f"{pdf_name} | {authoritative_name}"
                all_raw_data[unique_key] = verified_fields
                raw_export_data[unique_key] = deepcopy(verified_fields)
                all_normalized_data[unique_key] = deepcopy(normalized_fields)
                print(f"  ✓ Stored: {unique_key}")

        except Exception as e:
            import traceback
            print(f"  ❌ Error processing {pdf_name}: {e}")
            traceback.print_exc()

    if not all_raw_data:
        print("❌ No plans extracted.")
        return

    raw_snapshot_path = output_dir / "quote_compare_raw_extractions.json"
    page_audit_path = output_dir / "quote_compare_page_audit.json"
    write_json_file(raw_snapshot_path, all_raw_data)
    write_json_file(page_audit_path, page_extraction_audit)
    print(f"✓ Raw extraction JSON saved: {raw_snapshot_path}")
    print(f"✓ Page audit JSON saved: {page_audit_path}")

    print(f"\n{'='*60}")
    print(f"Total plans extracted: {len(all_raw_data)}")
    print(f"Plans: {list(all_raw_data.keys())}")

    print("\n── Normalization Pass (per-plan) ──")
    print("  ℹ️  Normalization already completed once after page extraction for each plan")

    print("\n── Comparative Best-Marking Pass ──")
    print("  ℹ️  Comparative best-marking disabled — using normalized data as-is")
    compared_data = all_normalized_data

    print("\n── Scoring & Conclusion Pass ──")
    print("  ℹ️  UW scoring and per-plan UW conclusions disabled — storing normalized plans directly")
    for unique_key in all_raw_data.keys():
        norm_data = compared_data.get(unique_key, all_normalized_data.get(unique_key, {}))
        if not norm_data:
            continue
        try:
            print(f"  Finalizing: {unique_key}")
            file_metadata = _get_effective_file_metadata(unique_key, {"file_metadata": report_metadata.get("pdf_file_metadata", {}).get(_extract_pdf_name_from_unique_key(unique_key), {})}, report_metadata)
            final_results[unique_key] = {
                "raw":           deepcopy(raw_export_data.get(unique_key, all_raw_data[unique_key])),
                "normalized":    norm_data,
                "scores":        {},
                "total_score":   None,
                "uw_conclusion": "",
                "file_metadata": file_metadata,
            }
            print(f"  ✓ {unique_key}: finalized without UW scoring")
        except Exception as e:
            print(f"  ❌ Finalization error for {unique_key}: {e}")
            file_metadata = _get_effective_file_metadata(unique_key, {"file_metadata": report_metadata.get("pdf_file_metadata", {}).get(_extract_pdf_name_from_unique_key(unique_key), {})}, report_metadata)
            final_results[unique_key] = {
                "raw":           deepcopy(raw_export_data.get(unique_key, all_raw_data[unique_key])),
                "normalized":    norm_data,
                "scores":        {},
                "total_score":   None,
                "uw_conclusion": "",
                "file_metadata": file_metadata,
            }

    if not final_results:
        print("❌ final_results is empty — nothing to save.")
        return

    final_results_path = output_dir / "quote_compare_final_results.json"
    report_payload_path = output_dir / "quote_compare_report_payload.json"
    metadata_path = output_dir / "quote_compare_run_metadata.json"
    write_json_file(final_results_path, final_results)
    write_json_file(metadata_path, report_metadata)
    reloaded_final_results = read_json_file(final_results_path)
    print(f"✓ Final results JSON saved: {final_results_path}")
    print(f"✓ Run metadata JSON saved: {metadata_path}")

    print(f"\n── Saving outputs ({len(final_results)} plans) ──")
    report_payload = build_report_payload(reloaded_final_results, ordered_keys, report_metadata=report_metadata)
    write_json_file(report_payload_path, report_payload)
    print(f"✓ Report payload JSON saved: {report_payload_path}")
    generate_reports_from_payload(report_payload, output_dir=output_dir)
    _save_token_report(tracker, output_path=output_dir / "Token_Usage_Report.xlsx")
    print("\n✅ Done!")

    return {
        "raw_json": str(raw_snapshot_path),
        "page_audit_json": str(page_audit_path),
        "final_json": str(final_results_path),
        "report_json": str(report_payload_path),
        "metadata_json": str(metadata_path),
        "excel": str(output_dir / "Insurance_Technical_Analysis.xlsx"),
        "word": str(output_dir / "Corporate_Comparison_Final.docx"),
        "token_report": str(output_dir / "Token_Usage_Report.xlsx"),
    }


def parse_cli_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run QuoteCompare extraction for one or more PDF files.")
    parser.add_argument("pdfs", nargs="*", help="One or more PDF paths to process.")
    parser.add_argument("--output-dir", dest="output_dir", default=None, help="Optional output folder for JSON, Excel, Word, and logs.")
    parser.add_argument("--report-only-json", dest="report_only_json", default=None, help="Optional path to quote_compare_report_payload.json or quote_compare_final_results.json to regenerate Excel and Word without rerunning extraction.")
    parser.add_argument("--resume-from-raw-json", dest="resume_from_raw_json", default=None, help="Optional path to quote_compare_raw_extractions.json to resume from normalization/report generation without rerunning page extraction.")
    parser.add_argument("--uw-name", dest="uw_names", action="append", default=None, help="Optional user-provided underwriter name. Repeat once per PDF in the same order as the PDF arguments, or provide a single value to apply to all PDFs.")
    parser.add_argument("--existing-benefit", dest="existing_benefits", action="append", default=None, help="Repeat once per PDF with true/false to identify the existing baseline quote.")
    return parser.parse_args()


if __name__ == "__main__":
    cli_args = parse_cli_args()
    if cli_args.report_only_json:
        target_output_dir = cli_args.output_dir or Path(cli_args.report_only_json).resolve().parent
        source_path = Path(cli_args.report_only_json)
        source_data = read_json_file(source_path)
        if source_path.name.lower() == "quote_compare_final_results.json":
            fallback_report_json = source_path.with_name("quote_compare_report_payload.json")
            generate_reports_from_final_results(
                source_data,
                output_dir=target_output_dir,
                report_json_path=fallback_report_json,
            )
        else:
            generate_reports_from_payload(
                source_data,
                output_dir=target_output_dir,
                report_json_path=source_path,
            )
    elif cli_args.resume_from_raw_json:
        target_output_dir = cli_args.output_dir or Path(cli_args.resume_from_raw_json).resolve().parent
        raw_path = Path(cli_args.resume_from_raw_json)
        raw_data = read_json_file(raw_path)

        uw_names = list(cli_args.uw_names or [])
        existing_benefits = list(cli_args.existing_benefits or [])
        pdfs = list(cli_args.pdfs or [])
        metadata = {}

        def _align_cli_values(values: list[str], flag_name: str, allow_single_repeat: bool = True) -> list[str]:
            if not values:
                return [""] * len(pdfs)
            if not pdfs:
                raise SystemExit(f"{flag_name} requires one or more PDF paths.")
            if len(values) == len(pdfs):
                return values
            if allow_single_repeat and len(values) == 1 and len(pdfs) > 1:
                return values * len(pdfs)
            raise SystemExit(f"Provide either one {flag_name} for all PDFs or one {flag_name} per PDF in the same order.")

        if uw_names:
            if not pdfs:
                raise SystemExit("--uw-name requires one or more PDF paths when using --resume-from-raw-json.")
            if len(uw_names) not in {1, len(pdfs)}:
                raise SystemExit("Provide either one --uw-name for all PDFs or one --uw-name per PDF in the same order.")
            if len(uw_names) == 1 and len(pdfs) > 1:
                uw_names = uw_names * len(pdfs)
            metadata = {
                "pdf_underwriter_overrides": {
                    os.path.basename(pdf_path): str(uw_name or "").strip()
                    for pdf_path, uw_name in zip(pdfs, uw_names)
                    if str(uw_name or "").strip()
                }
            }
        if pdfs:
            aligned_existing_benefits = _align_cli_values(existing_benefits, "--existing-benefit", allow_single_repeat=False)
            file_metadata = {}
            for pdf_path, existing_value in zip(pdfs, aligned_existing_benefits):
                pdf_name = os.path.basename(pdf_path)
                file_metadata[pdf_name] = {
                    "existing_benefit": _normalize_bool_flag(existing_value, "existing benefit") if str(existing_value).strip() else False,
                }
            metadata["pdf_file_metadata"] = file_metadata

        resume_from_raw_extractions(raw_data, metadata=metadata, pdf_paths=pdfs or None, output_dir=target_output_dir)
    else:
        uw_names = list(cli_args.uw_names or [])
        existing_benefits = list(cli_args.existing_benefits or [])
        pdfs = list(cli_args.pdfs or [])
        metadata = {}

        def _align_cli_values(values: list[str], flag_name: str, allow_single_repeat: bool = True) -> list[str]:
            if not values:
                return [""] * len(pdfs)
            if not pdfs:
                raise SystemExit(f"{flag_name} requires one or more PDF paths.")
            if len(values) == len(pdfs):
                return values
            if allow_single_repeat and len(values) == 1 and len(pdfs) > 1:
                return values * len(pdfs)
            raise SystemExit(f"Provide either one {flag_name} for all PDFs or one {flag_name} per PDF in the same order.")

        if uw_names:
            if not pdfs:
                raise SystemExit("--uw-name requires one or more PDF paths.")
            if len(uw_names) not in {1, len(pdfs)}:
                raise SystemExit("Provide either one --uw-name for all PDFs or one --uw-name per PDF in the same order.")
            if len(uw_names) == 1 and len(pdfs) > 1:
                uw_names = uw_names * len(pdfs)
            metadata = {
                "pdf_underwriter_overrides": {
                    os.path.basename(pdf_path): str(uw_name or "").strip()
                    for pdf_path, uw_name in zip(pdfs, uw_names)
                    if str(uw_name or "").strip()
                }
            }
        if pdfs:
            aligned_existing_benefits = _align_cli_values(existing_benefits, "--existing-benefit", allow_single_repeat=False)
            file_metadata = {}
            for pdf_path, existing_value in zip(pdfs, aligned_existing_benefits):
                pdf_name = os.path.basename(pdf_path)
                file_metadata[pdf_name] = {
                    "existing_benefit": _normalize_bool_flag(existing_value, "existing benefit") if str(existing_value).strip() else False,
                }
            metadata["pdf_file_metadata"] = file_metadata
        main(metadata=metadata, pdf_paths=pdfs or None, output_dir=cli_args.output_dir)